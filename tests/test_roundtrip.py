#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""Tests for md2docx.py / docx2md.py.

Exercises CONTRACT:C1-THEME-SCHEMA.1.1,
CONTRACT:C2-PROVENANCE.2.0, and CONTRACT:C3-ROUNDTRIP.1.2.

No test framework required — this is a plain script:
  python3 tests/test_roundtrip.py

(pytest also picks up the test_* functions if you have it installed and
prefer `pytest tests/`, but that is not a requirement of this project.)
"""
import base64
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
import traceback
import uuid
import warnings
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

import md2docx  # noqa: E402
import docx2md  # noqa: E402

FIXTURES = Path(__file__).resolve().parent / "fixtures"
THEMES = REPO_ROOT / "themes"


def _rewrite_docx(source, destination, mutate):
    """Rewrite a temporary test package while retaining existing ZIP metadata."""
    with zipfile.ZipFile(source, "r") as package:
        infos = package.infolist()
        parts = {
            info.filename: package.read(info)
            for info in infos
        }
    mutate(parts)
    original_names = {info.filename for info in infos}
    with zipfile.ZipFile(destination, "w") as package:
        for info in infos:
            package.writestr(info, parts[info.filename])
        for name in sorted(set(parts) - original_names):
            package.writestr(name, parts[name])


def _word_text_run(text):
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    return run


def _word_paragraph(text):
    paragraph = OxmlElement("w:p")
    paragraph.append(_word_text_run(text))
    return paragraph


def _assert_roundtrip_refusal(path, expected_code):
    try:
        docx2md.convert(path, report_provenance=False)
        raise AssertionError(f"expected {expected_code}")
    except docx2md.RoundtripRefusalError as exc:
        assert exc.code == expected_code, str(exc)
        return exc


# ---------------------------------------------------------------------------
# CONTRACT:C1-THEME-SCHEMA.1.1
# ---------------------------------------------------------------------------

def test_theme_deep_merge():
    base = {"a": 1, "b": {"x": 1, "y": 2}, "c": [1, 2]}
    over = {"b": {"y": 20}, "c": [9], "d": 4}
    merged = md2docx.deep_merge(base, over)

    assert merged["a"] == 1, "untouched top-level key must survive"
    assert merged["b"]["x"] == 1, "untouched nested key must survive"
    assert merged["b"]["y"] == 20, "overridden nested key must take the new value"
    assert merged["c"] == [9], "non-dict values (lists) replace wholesale, not merge element-wise"
    assert merged["d"] == 4, "new keys in `over` are added"

    # Real schema: overriding one heading level must not blow away the others.
    partial = {"headings": {"h1": {"size_pt": 24}}}
    merged = md2docx.deep_merge(md2docx.DEFAULTS, partial)
    assert merged["headings"]["h1"]["size_pt"] == 24
    assert merged["headings"]["h1"]["color"] == md2docx.DEFAULTS["headings"]["h1"]["color"]
    assert merged["headings"]["h2"] == md2docx.DEFAULTS["headings"]["h2"]
    assert merged["fonts"] == md2docx.DEFAULTS["fonts"]


def test_template_resolution_order():
    with tempfile.TemporaryDirectory() as td:
        base = Path(td)

        # 1. Nothing present anywhere -> hard-coded DEFAULTS.
        cfg, tpl_path, msg = md2docx.resolve_template(None, base_dir=base)
        assert cfg == md2docx.DEFAULTS
        assert tpl_path is None
        assert "built-in defaults" in msg

        # 2. themes/neutral.json present -> used.
        (base / "themes").mkdir()
        (base / "themes" / "neutral.json").write_text(
            json.dumps({"name": "Neutral", "fonts": {"body": "Georgia"}}))
        cfg, tpl_path, msg = md2docx.resolve_template(None, base_dir=base)
        assert cfg["fonts"]["body"] == "Georgia"
        assert tpl_path == base / "themes" / "neutral.json"

        # 3. md2docx-template.json next to the script beats themes/neutral.json.
        (base / "md2docx-template.json").write_text(
            json.dumps({"name": "Local Override", "fonts": {"body": "Verdana"}}))
        cfg, tpl_path, msg = md2docx.resolve_template(None, base_dir=base)
        assert cfg["fonts"]["body"] == "Verdana"
        assert tpl_path == base / "md2docx-template.json"

        # 4. An explicit --template flag beats everything, even a bogus one
        #    (which must fail loudly rather than silently fall back).
        explicit = base / "explicit.json"
        explicit.write_text(json.dumps({"name": "Explicit", "fonts": {"body": "Times"}}))
        cfg, tpl_path, msg = md2docx.resolve_template(str(explicit), base_dir=base)
        assert cfg["fonts"]["body"] == "Times"
        assert tpl_path == explicit

        try:
            md2docx.resolve_template(str(base / "does-not-exist.json"), base_dir=base)
            raise AssertionError("expected SystemExit for a missing explicit --template path")
        except SystemExit as e:
            assert "not found" in str(e)


def test_shipped_themes_load():
    shipped = sorted(THEMES.glob("*.json"))
    assert len(shipped) >= 3, "expected at least neutral, plum, and marked-docs"
    required_top_level = set(md2docx.DEFAULTS.keys())
    for theme_path in shipped:
        raw = json.loads(theme_path.read_text(encoding="utf8"))
        merged = md2docx.deep_merge(md2docx.DEFAULTS, raw)
        assert required_top_level.issubset(merged.keys()), f"{theme_path.name} missing schema keys after merge"
        assert "name" in raw, f"{theme_path.name} should declare a 'name' for provenance stamping"


def test_controlled_word_styles_are_fully_materialized():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "styles.md"
        output = td / "styles.docx"
        source.write_text(
            "# One\n\n## Two\n\n### Three\n\n#### Four\n\nBody.\n",
            encoding="utf8",
        )
        theme_path = THEMES / "plum.json"
        cfg = md2docx.deep_merge(
            md2docx.DEFAULTS,
            json.loads(theme_path.read_text(encoding="utf8")),
        )
        md2docx.Converter(cfg, tpl_path=theme_path).convert(source, output)

        with zipfile.ZipFile(output, "r") as package:
            styles_root = ElementTree.fromstring(
                package.read("word/styles.xml"))
        styles = {
            style.get(qn("w:styleId")): style
            for style in styles_root.findall(qn("w:style"))
        }

        def properties(style_id):
            r_pr = styles[style_id].find(qn("w:rPr"))
            assert r_pr is not None, style_id
            r_fonts = r_pr.find(qn("w:rFonts"))
            assert r_fonts is not None, style_id
            return r_pr, r_fonts

        normal_r_pr, normal_fonts = properties("Normal")
        for name in md2docx.EXPLICIT_FONT_ATTRIBUTES:
            assert normal_fonts.get(qn(f"w:{name}")) == "Arial"
        for name in md2docx.THEME_FONT_ATTRIBUTES:
            assert normal_fonts.get(qn(f"w:{name}")) is None
        assert normal_r_pr.find(qn("w:sz")).get(qn("w:val")) == "20"
        assert normal_r_pr.find(qn("w:szCs")).get(qn("w:val")) == "20"

        expected_sizes = {
            1: "36",
            2: "28",
            3: "23",
            4: "21",
        }
        for level, size in expected_sizes.items():
            paragraph_id = f"Heading{level}"
            character_id = f"Heading{level}Char"
            for style_id in (paragraph_id, character_id):
                r_pr, r_fonts = properties(style_id)
                for name in md2docx.EXPLICIT_FONT_ATTRIBUTES:
                    assert r_fonts.get(qn(f"w:{name}")) == "Arial"
                for name in md2docx.THEME_FONT_ATTRIBUTES:
                    assert r_fonts.get(qn(f"w:{name}")) is None
                assert r_pr.find(qn("w:sz")).get(qn("w:val")) == size
                assert r_pr.find(qn("w:szCs")).get(qn("w:val")) == size
                assert r_pr.find(qn("w:b")).get(qn("w:val")) == "1"
                assert r_pr.find(qn("w:bCs")).get(qn("w:val")) == "1"
                assert r_pr.find(qn("w:i")).get(qn("w:val")) == "0"
                assert r_pr.find(qn("w:iCs")).get(qn("w:val")) == "0"
            assert styles[paragraph_id].find(
                qn("w:link")).get(qn("w:val")) == character_id
            assert styles[character_id].find(
                qn("w:link")).get(qn("w:val")) == paragraph_id

        recovered, report = docx2md.convert_with_report(
            output, report_provenance=False)
        assert recovered == source.read_text(encoding="utf8")
        assert report["schema"] == "office180-docx-roundtrip-report/0.2"
        assert report["state"] == "exact-supported-profile"
        assert report["semanticNormalization"] == {
            "state": "exact",
            "events": [],
        }
        projection = report["visualStyleProjection"]
        assert (
            projection["schema"]
            == "office180-docx-visual-style-projection/0.2"
        )
        assert projection["state"] == "materialized"
        assert projection["bodyFont"] == "Arial"
        assert projection["normalizations"] == []
        assert projection["diagnostics"] == []

        nonbold_output = td / "styles-nonbold.docx"
        nonbold_cfg = md2docx.deep_merge(
            cfg, {"headings": {"bold": False}})
        md2docx.Converter(nonbold_cfg).convert(
            source, nonbold_output)
        with zipfile.ZipFile(nonbold_output, "r") as package:
            nonbold_root = ElementTree.fromstring(
                package.read("word/styles.xml"))
        nonbold_styles = {
            style.get(qn("w:styleId")): style
            for style in nonbold_root.findall(qn("w:style"))
        }
        for style_id in ("Heading1", "Heading1Char"):
            r_pr = nonbold_styles[style_id].find(qn("w:rPr"))
            assert r_pr.find(qn("w:b")).get(qn("w:val")) == "0"
            assert r_pr.find(qn("w:bCs")).get(qn("w:val")) == "0"
        _, nonbold_report = docx2md.convert_with_report(
            nonbold_output, report_provenance=False)
        assert (
            nonbold_report["visualStyleProjection"]["state"]
            == "materialized"
        )


# ---------------------------------------------------------------------------
# CONTRACT:C2-PROVENANCE.2.0
# ---------------------------------------------------------------------------

def test_provenance_stamp():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "example.md"
        src.write_text("# Title\n\nHello world.\n", encoding="utf8")
        out = td / "example.docx"

        tpl_path = THEMES / "plum.json"
        cfg = md2docx.deep_merge(md2docx.DEFAULTS, json.loads(tpl_path.read_text(encoding="utf8")))
        md2docx.Converter(cfg, tpl_path=tpl_path).convert(src, out)

        doc = docx2md.Document(str(out))
        cp = doc.core_properties
        prov = json.loads(cp.comments)

        assert set(prov.keys()) == {"t", "tpl", "tplsha", "srcsha", "gen"}
        assert prov["t"] == f"md2docx/{md2docx.TOOL_VERSION}"
        assert prov["tpl"] == "Plum"
        assert prov["tplsha"] is not None and len(prov["tplsha"]) == 16
        assert prov["tplsha"] == hashlib.sha256(tpl_path.read_bytes()).hexdigest()[:16]
        assert prov["srcsha"] == hashlib.sha256(src.read_bytes()).hexdigest()[:16]
        assert re.fullmatch(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}Z", prov["gen"])

        assert cp.keywords == "md2docx"
        assert cp.category == "Plum"
        assert cp.subject.endswith("example.md")

        embedded = docx2md.read_embedded_source(out, required=True)
        assert embedded.original_text == src.read_text(encoding="utf8")
        assert embedded.canonical_text == md2docx.canonicalize_markdown(
            embedded.original_text)
        assert embedded.original_sha256 == hashlib.sha256(
            src.read_bytes()).hexdigest()
        assert embedded.canonical_sha256 == hashlib.sha256(
            embedded.canonical_text.encode("utf8")).hexdigest()

        # No template at all -> tplsha is JSON null, tpl says so explicitly.
        out2 = td / "example-notpl.docx"
        md2docx.Converter(md2docx.DEFAULTS, tpl_path=None).convert(src, out2)
        prov2 = json.loads(docx2md.Document(str(out2)).core_properties.comments)
        assert prov2["tplsha"] is None
        assert prov2["tpl"] == "built-in defaults"


def test_embedded_source_survives_resave_and_three_way_merge():
    base_text = "# Branch\n\nFirst paragraph.\n\nSecond paragraph.\n"
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "branch.md"
        baseline_docx = td / "branch.docx"
        resaved_docx = td / "branch-resaved.docx"
        edited_docx = td / "branch-edited.docx"
        current = td / "branch-current.md"
        src.write_text(base_text, encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(src, baseline_docx)

        # python-docx is an independent package reopen/save path. The custom
        # XML merge base must remain related and hash-valid after that save.
        package = docx2md.Document(str(baseline_docx))
        package.save(resaved_docx)
        resaved = docx2md.read_embedded_source(
            resaved_docx, required=True)
        assert resaved.canonical_text == base_text
        resaved_text, resaved_report = docx2md.convert_with_report(
            resaved_docx, report_provenance=False)
        assert resaved_text == base_text
        assert (
            resaved_report["visualStyleProjection"]["state"]
            == "materialized"
        )

        edited = docx2md.Document(str(resaved_docx))
        target = next(
            paragraph for paragraph in edited.paragraphs
            if paragraph.text == "First paragraph.")
        target.text = "First paragraph edited in Word."
        edited.save(edited_docx)

        current_text = (
            "# Branch\n\nFirst paragraph.\n\n"
            "Second paragraph edited in Markdown.\n"
        )
        current.write_text(current_text, encoding="utf8")
        merged = docx2md.merge_with_current(
            edited_docx, current, report_provenance=False)
        assert not merged.conflicts
        assert merged.text == (
            "# Branch\n\nFirst paragraph edited in Word.\n\n"
            "Second paragraph edited in Markdown.\n"
        )

        current.write_text(
            "# Branch\n\nFirst paragraph edited in Markdown.\n\n"
            "Second paragraph.\n",
            encoding="utf8")
        conflicted = docx2md.merge_with_current(
            edited_docx, current, report_provenance=False)
        assert conflicted.conflicts
        assert "<<<<<<<" in conflicted.text
        assert "|||||||" in conflicted.text
        assert ">>>>>>>" in conflicted.text

        current.write_bytes(
            b"# Branch\r\n\r\nFirst paragraph.\r\n\r\n"
            b"Second paragraph.\r\n")
        try:
            docx2md.merge_with_current(
                edited_docx, current, report_provenance=False)
            raise AssertionError("expected CRLF current-source refusal")
        except docx2md.RoundtripRefusalError as exc:
            assert exc.code == "DOCX-ROUNDTRIP-CURRENT-NONCANONICAL"


def test_report_and_merge_use_one_immutable_docx_snapshot():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source_a = td / "a.md"
        source_b = td / "b.md"
        docx_a = td / "a.docx"
        docx_b = td / "b.docx"
        active = td / "active.docx"
        current = td / "current.md"
        text_a = "# A\n\nBody A.\n"
        text_b = "# B\n\nBody B.\n"
        source_a.write_text(text_a, encoding="utf8")
        source_b.write_text(text_b, encoding="utf8")
        current.write_text(text_a, encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(source_a, docx_a)
        md2docx.Converter(md2docx.DEFAULTS).convert(source_b, docx_b)
        bytes_a = docx_a.read_bytes()
        bytes_b = docx_b.read_bytes()

        real_parts_reader = docx2md._read_safe_docx_parts_bytes

        def swap_path_after_snapshot(payload):
            parts = real_parts_reader(payload)
            shutil.copy2(docx_b, active)
            return parts

        shutil.copy2(docx_a, active)
        docx2md._read_safe_docx_parts_bytes = swap_path_after_snapshot
        try:
            recovered, report = docx2md.convert_with_report(
                active, report_provenance=False)
        finally:
            docx2md._read_safe_docx_parts_bytes = real_parts_reader
        assert active.read_bytes() == bytes_b
        assert recovered == text_a
        assert report["input"]["sha256"] == hashlib.sha256(
            bytes_a).hexdigest()
        assert report["embeddedSource"]["canonicalSha256"] == hashlib.sha256(
            text_a.encode("utf8")).hexdigest()

        shutil.copy2(docx_a, active)
        docx2md._read_safe_docx_parts_bytes = swap_path_after_snapshot
        try:
            merged = docx2md.merge_with_current(
                active, current, report_provenance=False)
        finally:
            docx2md._read_safe_docx_parts_bytes = real_parts_reader
        assert active.read_bytes() == bytes_b
        assert not merged.conflicts
        assert merged.text == text_a
        assert merged.base_sha256 == hashlib.sha256(
            text_a.encode("utf8")).hexdigest()
        assert merged.edited_sha256 == merged.base_sha256


def test_three_way_merge_returns_multiple_conflict_regions():
    base_text = (
        "# Branch\n\nOne.\n\nMiddle A.\n\nMiddle B.\n\nFour.\n"
    )
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        edited_path = td / "edited.docx"
        current = td / "current.md"
        source.write_text(base_text, encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(
            source, edited_path)

        edited = docx2md.Document(str(edited_path))
        for paragraph in edited.paragraphs:
            if paragraph.text == "One.":
                paragraph.text = "One edited in Word."
            elif paragraph.text == "Four.":
                paragraph.text = "Four edited in Word."
        edited.save(edited_path)
        current.write_text(
            (
                "# Branch\n\nOne edited in Markdown.\n\nMiddle A.\n\n"
                "Middle B.\n\nFour edited in Markdown.\n"
            ),
            encoding="utf8",
        )

        merged = docx2md.merge_with_current(
            edited_path, current, report_provenance=False)
        assert merged.conflicts
        assert merged.text.count("<<<<<<<") == 2
        assert merged.text.count("|||||||") == 2
        assert merged.text.count(">>>>>>>") == 2

        output = td / "merged.md"
        base_output = td / "base.md"
        report_output = td / "merge-report.json"
        cli = subprocess.run(
            [
                sys.executable,
                str(REPO_ROOT / "docx2md.py"),
                str(edited_path),
                "--merge-current",
                str(current),
                "--out",
                str(output),
                "--base-out",
                str(base_output),
                "--report",
                str(report_output),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert cli.returncode == 1
        assert "DOCX-ROUNDTRIP-MERGE-CONFLICT" in cli.stderr
        assert output.read_text(encoding="utf8").count("<<<<<<<") == 2
        assert base_output.read_text(encoding="utf8") == base_text
        report = json.loads(report_output.read_text(encoding="utf8"))
        assert report["merge"]["state"] == "conflicted"
        assert (
            report["embeddedSource"]["state"]
            == "internally-consistent"
        )


def test_merge_file_exit_status_is_authoritative_for_conflicts():
    assert not docx2md._merge_exit_has_conflicts(0)
    assert docx2md._merge_exit_has_conflicts(1)
    assert docx2md._merge_exit_has_conflicts(127)
    assert not docx2md._merge_exit_has_conflicts(128)
    assert not docx2md._merge_exit_has_conflicts(-9)

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        docx_path = td / "source.docx"
        current = td / "current.md"
        source.write_text("# Source\n\nBody.\n", encoding="utf8")
        current.write_text("# Source\n\nBody.\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(source, docx_path)

        real_subprocess = docx2md.subprocess

        class ConflictProcess:
            @staticmethod
            def run(*args, **kwargs):
                return real_subprocess.CompletedProcess(
                    args[0],
                    2,
                    stdout=(
                        b"<<<<<<< current\ncurrent\n||||||| base\nbase\n"
                        b"=======\nedited\n>>>>>>> edited\n"
                    ),
                    # git-merge-file(1) permits a conflict warning on stderr;
                    # its positive conflict-count exit value is authoritative.
                    stderr=b"warning: two conflicts\n",
                )

        docx2md.subprocess = ConflictProcess
        try:
            merged = docx2md.merge_with_current(
                docx_path, current, report_provenance=False)
        finally:
            docx2md.subprocess = real_subprocess
        assert merged.conflicts
        assert merged.text.startswith("<<<<<<<")

        class ErrorProcess:
            @staticmethod
            def run(*args, **kwargs):
                return real_subprocess.CompletedProcess(
                    args[0],
                    128,
                    stdout=b"",
                    stderr=b"fatal: merge tool failed\n",
                )

        docx2md.subprocess = ErrorProcess
        try:
            try:
                docx2md.merge_with_current(
                    docx_path, current, report_provenance=False)
                raise AssertionError("expected merge-tool failure")
            except docx2md.RoundtripRefusalError as exc:
                assert exc.code == "DOCX-ROUNDTRIP-MERGE"
                assert "merge tool failed" in exc.message
        finally:
            docx2md.subprocess = real_subprocess


def test_embedded_source_tamper_is_refused():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "source.md"
        original = td / "source.docx"
        tampered = td / "tampered.docx"
        src.write_text("# Trusted\n\nBody.\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(src, original)
        snapshot = docx2md.read_embedded_source(original, required=True)

        with zipfile.ZipFile(original, "r") as package:
            infos = package.infolist()
            parts = {
                info.filename: package.read(info)
                for info in infos
            }
        marker = f'canonicalSha256="{snapshot.canonical_sha256}"'.encode()
        replacement = ('canonicalSha256="' + "0" * 64 + '"').encode()
        assert parts[snapshot.part].count(marker) == 1
        parts[snapshot.part] = parts[snapshot.part].replace(
            marker, replacement)
        with zipfile.ZipFile(tampered, "w") as package:
            for info in infos:
                package.writestr(info, parts[info.filename])

        try:
            docx2md.read_embedded_source(tampered, required=True)
            raise AssertionError("expected embedded-source tamper refusal")
        except docx2md.RoundtripRefusalError as exc:
            assert exc.code == "DOCX-ROUNDTRIP-PROVENANCE"


def test_embedded_source_requires_internal_consistency():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        original = td / "source.docx"
        source.write_text("# Base\n\nBody.\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(source, original)
        snapshot = docx2md.read_embedded_source(original, required=True)
        item_number = re.search(
            r"item([1-9]\d*)\.xml$", snapshot.part).group(1)
        props_name = f"customXml/itemProps{item_number}.xml"

        _, report = docx2md.convert_with_report(
            original, report_provenance=False)
        assert report["embeddedSource"]["state"] == "internally-consistent"

        writer_candidate = td / "writer-candidate.docx"
        shutil.copy2(original, writer_candidate)
        writer_before = writer_candidate.read_bytes()
        try:
            md2docx.embed_source_snapshot(
                writer_candidate,
                source.read_bytes(),
                "# Unrelated canonical text\n",
            )
            raise AssertionError("expected writer canonical relation refusal")
        except ValueError as exc:
            assert "canonicalize(original)" in str(exc)
        assert writer_candidate.read_bytes() == writer_before

        wrong_uuid = td / "wrong-uuid.docx"

        def replace_uuid(parts):
            parts[props_name] = re.sub(
                br'ds:itemID="\{[A-F0-9-]+\}"',
                b'ds:itemID="{00000000-0000-0000-0000-000000000001}"',
                parts[props_name],
            )

        _rewrite_docx(original, wrong_uuid, replace_uuid)
        exc = _assert_roundtrip_refusal(
            wrong_uuid, "DOCX-ROUNDTRIP-PROVENANCE")
        assert "datastore identity" in exc.message

        forged = td / "forged-canonical.docx"
        forged_text = "# Forged\n\nBody.\n"
        forged_bytes = forged_text.encode("utf8")
        forged_sha = hashlib.sha256(forged_bytes).hexdigest()
        forged_uuid = uuid.uuid5(
            uuid.NAMESPACE_URL,
            (
                f"{docx2md.EMBEDDED_SOURCE_SCHEMA}:"
                f"{snapshot.original_sha256}:{forged_sha}"
            ),
        )

        def replace_canonical(parts):
            item = parts[snapshot.part]
            item = re.sub(
                br'canonicalSha256="[a-f0-9]{64}"',
                f'canonicalSha256="{forged_sha}"'.encode(),
                item,
            )
            item = re.sub(
                br'(<o:canonical encoding="base64">)[^<]*(</o:canonical>)',
                (
                    rb"\1"
                    + base64.b64encode(forged_bytes)
                    + rb"\2"
                ),
                item,
            )
            parts[snapshot.part] = item
            parts[props_name] = re.sub(
                br'ds:itemID="\{[A-F0-9-]+\}"',
                (
                    'ds:itemID="{'
                    + str(forged_uuid).upper()
                    + '}"'
                ).encode(),
                parts[props_name],
            )

        _rewrite_docx(original, forged, replace_canonical)
        exc = _assert_roundtrip_refusal(
            forged, "DOCX-ROUNDTRIP-PROVENANCE")
        assert "canonicalize(original)" in exc.message

        external = td / "external-relationship.docx"

        def make_relationship_external(parts):
            rels_name = "word/_rels/document.xml.rels"
            marker = (
                f'Type="{docx2md.CUSTOM_XML_RELATIONSHIP}" '
                f'Target="../{snapshot.part}"'
            ).encode()
            replacement = (
                f'Type="{docx2md.CUSTOM_XML_RELATIONSHIP}" '
                f'TargetMode="External" Target="../{snapshot.part}"'
            ).encode()
            assert parts[rels_name].count(marker) == 1
            parts[rels_name] = parts[rels_name].replace(
                marker, replacement)

        _rewrite_docx(original, external, make_relationship_external)
        _assert_roundtrip_refusal(
            external, "DOCX-ROUNDTRIP-PROVENANCE")

        duplicate_relationship = td / "duplicate-relationship.docx"

        def duplicate_properties_relationship(parts):
            rels_name = (
                f"customXml/_rels/item{item_number}.xml.rels")
            relationship = re.search(
                br"<Relationship [^>]*/>", parts[rels_name]).group(0)
            parts[rels_name] = parts[rels_name].replace(
                b"</Relationships>",
                relationship + b"</Relationships>",
            )

        _rewrite_docx(
            original,
            duplicate_relationship,
            duplicate_properties_relationship,
        )
        _assert_roundtrip_refusal(
            duplicate_relationship, "DOCX-ROUNDTRIP-PROVENANCE")

        core_bad = td / "core-bad.docx"
        package = docx2md.Document(str(original))
        provenance = json.loads(package.core_properties.comments)
        provenance["srcsha"] = "0" * 16
        package.core_properties.comments = json.dumps(
            provenance, separators=(",", ":"))
        package.save(core_bad)
        try:
            docx2md.read_embedded_source(core_bad, required=True)
            raise AssertionError("expected direct core contradiction refusal")
        except docx2md.RoundtripRefusalError as exc:
            assert exc.code == "DOCX-ROUNDTRIP-PROVENANCE"
            assert "core-property source hash" in exc.message


def test_embedded_source_writer_refuses_self_inconsistent_publication():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        source.write_text("# Source\n\nBody.\n", encoding="utf8")

        wrong_core = td / "wrong-core.docx"
        package = docx2md.Document()
        package.add_heading("Source", level=1)
        package.add_paragraph("Body.")
        package.core_properties.comments = json.dumps(
            {"srcsha": "0" * 16}, separators=(",", ":"))
        package.save(wrong_core)
        before = wrong_core.read_bytes()
        try:
            md2docx.embed_source_snapshot(
                wrong_core,
                source.read_bytes(),
                source.read_text(encoding="utf8"),
            )
            raise AssertionError("expected final semantic validation refusal")
        except ValueError as exc:
            assert "semantic validation" in str(exc)
            assert "core-property source hash" in str(exc)
        assert wrong_core.read_bytes() == before

        already_embedded = td / "already-embedded.docx"
        md2docx.Converter(md2docx.DEFAULTS).convert(
            source, already_embedded)
        before = already_embedded.read_bytes()
        try:
            md2docx.embed_source_snapshot(
                already_embedded,
                source.read_bytes(),
                source.read_text(encoding="utf8"),
            )
            raise AssertionError("expected duplicate office180 item refusal")
        except ValueError as exc:
            assert "semantic validation" in str(exc)
            assert "more than one office180 embedded source" in str(exc)
        assert already_embedded.read_bytes() == before


def test_embedded_source_missing_and_package_limits_refuse():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        missing = td / "missing.docx"
        docx2md.Document().save(missing)
        try:
            docx2md.read_embedded_source(missing, required=True)
            raise AssertionError("expected missing merge-base refusal")
        except docx2md.RoundtripRefusalError as exc:
            assert exc.code == "DOCX-ROUNDTRIP-NO-MERGE-BASE"

        source = td / "source.md"
        original = td / "source.docx"
        source.write_text("# Source\n\nBody.\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(source, original)

        original_part_limit = docx2md.MAX_DOCX_PARTS
        docx2md.MAX_DOCX_PARTS = 1
        try:
            _assert_roundtrip_refusal(
                original, "DOCX-ROUNDTRIP-PACKAGE")
        finally:
            docx2md.MAX_DOCX_PARTS = original_part_limit

        original_size_limit = docx2md.MAX_DOCX_UNCOMPRESSED_BYTES
        docx2md.MAX_DOCX_UNCOMPRESSED_BYTES = 1
        try:
            _assert_roundtrip_refusal(
                original, "DOCX-ROUNDTRIP-PACKAGE")
        finally:
            docx2md.MAX_DOCX_UNCOMPRESSED_BYTES = original_size_limit

        original_source_limit = docx2md.MAX_EMBEDDED_SOURCE_BYTES
        docx2md.MAX_EMBEDDED_SOURCE_BYTES = 1
        try:
            _assert_roundtrip_refusal(
                original, "DOCX-ROUNDTRIP-PROVENANCE")
        finally:
            docx2md.MAX_EMBEDDED_SOURCE_BYTES = original_source_limit

        duplicate = td / "duplicate.docx"
        with zipfile.ZipFile(original, "r") as source_package:
            infos = source_package.infolist()
            payloads = {
                info.filename: source_package.read(info)
                for info in infos
            }
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            with zipfile.ZipFile(duplicate, "w") as package:
                for info in infos:
                    package.writestr(info, payloads[info.filename])
                package.writestr(
                    "word/document.xml", payloads["word/document.xml"])
        _assert_roundtrip_refusal(
            duplicate, "DOCX-ROUNDTRIP-PACKAGE")


# ---------------------------------------------------------------------------
# CONTRACT:C3-ROUNDTRIP.1.2
# ---------------------------------------------------------------------------

_STRIP_RE = re.compile(r"[`*#>|\[\]()-]")


def _word_bag(text):
    """Symmetric, structure-agnostic tokenizer: strip markdown punctuation,
    split on whitespace. Applying the SAME function to the source and the
    round-tripped output makes a multiset (Counter) diff a legitimate
    "zero lost / zero invented tokens" check — it does not require
    re-implementing a Markdown parser to be meaningful, because any
    asymmetry it detects is real: something that was words in one text and
    isn't in the other."""
    return Counter(_STRIP_RE.sub(" ", text).split())


def _normalize_source_for_comparison(text):
    """Drop constructs md2docx.py documents as intentionally, entirely
    dropped (not just reformatted): full-line HTML comments and
    horizontal rules. Everything else in the fixture is expected to
    survive the round trip."""
    kept = []
    for line in text.splitlines():
        s = line.strip()
        if s.startswith("<!--") and s.endswith("-->"):
            continue
        if re.fullmatch(r"-{3,}", s):
            continue
        kept.append(line)
    return "\n".join(kept)


def test_trailing_ascii_space_normalization_is_ordered_and_merge_safe():
    base_text = "# Branch\n\nFirst paragraph.\n\nSecond paragraph.\n"
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "branch.md"
        edited_docx = td / "branch-edited.docx"
        current = td / "branch-current.md"
        source.write_text(base_text, encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(
            source, edited_docx)

        before = docx2md.read_embedded_source(
            edited_docx, required=True)
        document = docx2md.Document(str(edited_docx))
        paragraph_indexes = {
            paragraph.text: index
            for index, paragraph in enumerate(document.paragraphs)
        }
        next(
            paragraph for paragraph in document.paragraphs
            if paragraph.text == "Branch"
        ).text = "Branch "
        next(
            paragraph for paragraph in document.paragraphs
            if paragraph.text == "First paragraph."
        ).text = "First paragraph edited in Word.  "
        document.save(edited_docx)

        after = docx2md.read_embedded_source(
            edited_docx, required=True)
        assert after == before
        recovered, report = docx2md.convert_with_report(
            edited_docx, report_provenance=False)
        assert recovered == (
            "# Branch\n\nFirst paragraph edited in Word.\n\n"
            "Second paragraph.\n"
        )
        assert report["state"] == "normalized-supported-profile"
        events = report["semanticNormalization"]["events"]
        assert [
            event["paragraphIndex"] for event in events
        ] == [
            paragraph_indexes["Branch"],
            paragraph_indexes["First paragraph."],
        ]
        assert [event["styleId"] for event in events] == [
            "Heading1",
            "Normal",
        ]
        assert [event["count"] for event in events] == [1, 2]
        assert [event["inputTextSha256"] for event in events] == [
            hashlib.sha256(b"Branch ").hexdigest(),
            hashlib.sha256(
                b"First paragraph edited in Word.  "
            ).hexdigest(),
        ]

        current.write_text(
            (
                "# Branch\n\nFirst paragraph.\n\n"
                "Second paragraph edited in Markdown.\n"
            ),
            encoding="utf8",
        )
        merged = docx2md.merge_with_current(
            edited_docx, current, report_provenance=False)
        assert not merged.conflicts
        assert merged.text == (
            "# Branch\n\nFirst paragraph edited in Word.\n\n"
            "Second paragraph edited in Markdown.\n"
        )


def test_visual_style_projection_reports_semantic_drift_only():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "styles.md"
        original = td / "styles.docx"
        metadata_only = td / "styles-metadata.docx"
        drifted = td / "styles-drifted.docx"
        source.write_text("# Title\n\nBody.\n", encoding="utf8")
        theme_path = THEMES / "plum.json"
        cfg = md2docx.deep_merge(
            md2docx.DEFAULTS,
            json.loads(theme_path.read_text(encoding="utf8")),
        )
        md2docx.Converter(cfg, tpl_path=theme_path).convert(
            source, original)

        def add_unrelated_metadata(parts):
            root = ElementTree.fromstring(parts["word/styles.xml"])
            normal = next(
                style for style in root.findall(qn("w:style"))
                if style.get(qn("w:styleId")) == "Normal"
            )
            normal.set(qn("w:rsid"), "0BADF00D")
            parts["word/styles.xml"] = ElementTree.tostring(
                root, encoding="utf8", xml_declaration=True)

        _rewrite_docx(original, metadata_only, add_unrelated_metadata)
        metadata_text, metadata_report = docx2md.convert_with_report(
            metadata_only, report_provenance=False)
        assert metadata_text == source.read_text(encoding="utf8")
        assert (
            metadata_report["visualStyleProjection"]["state"]
            == "materialized"
        )

        def replace_explicit_heading_fonts(parts):
            root = ElementTree.fromstring(parts["word/styles.xml"])
            styles = {
                style.get(qn("w:styleId")): style
                for style in root.findall(qn("w:style"))
            }
            theme_values = {
                "asciiTheme": "majorAscii",
                "hAnsiTheme": "majorHAnsi",
                "eastAsiaTheme": "majorEastAsia",
                "cstheme": "majorBidi",
            }
            for style_id in ("Heading1", "Heading1Char"):
                r_pr = styles[style_id].find(qn("w:rPr"))
                r_fonts = r_pr.find(qn("w:rFonts"))
                for name in md2docx.EXPLICIT_FONT_ATTRIBUTES:
                    r_fonts.attrib.pop(qn(f"w:{name}"), None)
                for name, value in theme_values.items():
                    r_fonts.set(qn(f"w:{name}"), value)
            parts["word/styles.xml"] = ElementTree.tostring(
                root, encoding="utf8", xml_declaration=True)

        _rewrite_docx(
            metadata_only, drifted, replace_explicit_heading_fonts)
        recovered, report = docx2md.convert_with_report(
            drifted, report_provenance=False)
        assert recovered == source.read_text(encoding="utf8")
        assert report["state"] == "exact-supported-profile"
        assert report["semanticNormalization"] == {
            "state": "exact",
            "events": [],
        }
        projection = report["visualStyleProjection"]
        assert projection["state"] == "drifted"
        assert {
            (diagnostic["styleId"], diagnostic["property"])
            for diagnostic in projection["diagnostics"]
        } == {
            ("Heading1", "explicitFonts"),
            ("Heading1", "themeFonts"),
            ("Heading1Char", "explicitFonts"),
            ("Heading1Char", "themeFonts"),
        }
        assert projection["themeFonts"]["majorLatin"] == "Calibri"
        assert projection["bodyFont"] == "Arial"
        assert all(
            diagnostic["code"]
            == "DOCX-ROUNDTRIP-VISUAL-STYLE-DRIFT"
            for diagnostic in projection["diagnostics"]
        )
        assert projection["diagnostics"] == report["diagnostics"]

        cli_output = td / "recovered.md"
        cli = subprocess.run(
            [
                sys.executable,
                str(REPO_ROOT / "docx2md.py"),
                str(drifted),
                "-o",
                str(cli_output),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert cli.returncode == 0, cli.stderr
        assert (
            "DOCX-ROUNDTRIP-VISUAL-STYLE-DRIFT"
            in cli.stderr
        )
        assert cli_output.read_text(encoding="utf8") == (
            source.read_text(encoding="utf8")
        )


def test_native_word_heading_cascade_normalization_is_proof_bounded():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "native-styles.md"
        baseline = td / "native-styles.docx"
        native_normalized = td / "native-styles-resaved.docx"
        source.write_text(
            "# One\n\n## Two\n\n### Three\n\n#### Four\n\nBody.\n",
            encoding="utf8",
        )
        theme_path = THEMES / "plum.json"
        cfg = md2docx.deep_merge(
            md2docx.DEFAULTS,
            json.loads(theme_path.read_text(encoding="utf8")),
        )
        md2docx.Converter(cfg, tpl_path=theme_path).convert(
            source,
            baseline,
        )

        def rewrite_styles(parts, mutate):
            root = ElementTree.fromstring(parts["word/styles.xml"])
            styles = {
                style.get(qn("w:styleId")): style
                for style in root.findall(qn("w:style"))
            }
            mutate(root, styles)
            parts["word/styles.xml"] = ElementTree.tostring(
                root,
                encoding="utf8",
                xml_declaration=True,
            )

        def apply_native_heading_omissions(parts):
            def mutate(_root, styles):
                for level in range(1, 5):
                    r_pr = styles[f"Heading{level}"].find(qn("w:rPr"))
                    for name in ("rFonts", "i", "iCs"):
                        element = r_pr.find(qn(f"w:{name}"))
                        if element is not None:
                            r_pr.remove(element)

            rewrite_styles(parts, mutate)

        _rewrite_docx(
            baseline,
            native_normalized,
            apply_native_heading_omissions,
        )
        baseline_text, baseline_report = docx2md.convert_with_report(
            baseline,
            report_provenance=False,
        )
        recovered, report = docx2md.convert_with_report(
            native_normalized,
            report_provenance=False,
        )
        assert baseline_text == recovered == source.read_text(encoding="utf8")
        assert (
            baseline_report["visualStyleProjection"]["state"]
            == "materialized"
        )
        assert (
            baseline_report["visualStyleProjection"]["normalizations"]
            == []
        )
        assert report["state"] == "exact-supported-profile"
        assert report["semanticNormalization"] == {
            "state": "exact",
            "events": [],
        }
        projection = report["visualStyleProjection"]
        assert (
            projection["state"]
            == "native-normalized-materialized-equivalent"
        )
        assert projection["diagnostics"] == []
        assert [
            (event["styleId"], event["property"])
            for event in projection["normalizations"]
        ] == [
            (f"Heading{level}", property_name)
            for level in range(1, 5)
            for property_name in ("explicitFonts", "italic")
        ]
        assert report["diagnostics"] == projection["normalizations"]
        assert all(
            event["code"]
            == "DOCX-ROUNDTRIP-VISUAL-STYLE-NATIVE-NORMALIZATION"
            and event["severity"] == "info"
            and event["effectiveState"] == "materialized-equivalent"
            and event["proof"]["baseStyleId"] == "Normal"
            and event["proof"]["linkedStyleId"]
            == f"{event['styleId']}Char"
            for event in projection["normalizations"]
        )
        assert all(
            event["proof"]["effectiveValues"]
            == (
                {
                    "ascii": "Arial",
                    "hAnsi": "Arial",
                    "eastAsia": "Arial",
                    "cs": "Arial",
                }
                if event["property"] == "explicitFonts"
                else {"latin": False, "complex": False}
            )
            for event in projection["normalizations"]
        )

        def ensure_r_fonts(style):
            r_pr = style.find(qn("w:rPr"))
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = ElementTree.Element(qn("w:rFonts"))
                r_pr.insert(0, r_fonts)
            return r_fonts

        def theme_reference(_root, styles):
            r_fonts = ensure_r_fonts(styles["Heading1"])
            r_fonts.set(qn("w:asciiTheme"), "majorHAnsi")

        def wrong_based_on(_root, styles):
            styles["Heading1"].find(qn("w:basedOn")).set(
                qn("w:val"),
                "Heading2",
            )

        def inherited_italic(root, _styles):
            defaults = root.find(
                f"{qn('w:docDefaults')}/"
                f"{qn('w:rPrDefault')}/{qn('w:rPr')}"
            )
            for name in ("i", "iCs"):
                defaults.append(ElementTree.Element(qn(f"w:{name}")))

        def wrong_linked_font(_root, styles):
            r_fonts = ensure_r_fonts(styles["Heading1Char"])
            for name in md2docx.EXPLICIT_FONT_ATTRIBUTES:
                r_fonts.set(qn(f"w:{name}"), "Times New Roman")

        def explicit_font_conflict(_root, styles):
            r_fonts = ensure_r_fonts(styles["Heading1"])
            for name in md2docx.EXPLICIT_FONT_ATTRIBUTES:
                r_fonts.set(qn(f"w:{name}"), "Times New Roman")

        def missing_link(_root, styles):
            link = styles["Heading1"].find(qn("w:link"))
            styles["Heading1"].remove(link)

        def partial_direct_font(_root, styles):
            ensure_r_fonts(styles["Heading1"]).set(
                qn("w:ascii"),
                "Arial",
            )

        counterexamples = (
            ("theme-reference", theme_reference, "Heading1", "themeFonts"),
            ("wrong-based-on", wrong_based_on, "Heading1", "basedOn"),
            (
                "inherited-italic",
                inherited_italic,
                "Heading1",
                "italic",
            ),
            (
                "wrong-linked-font",
                wrong_linked_font,
                "Heading1Char",
                "explicitFonts",
            ),
            (
                "explicit-font-conflict",
                explicit_font_conflict,
                "Heading1",
                "explicitFonts",
            ),
            ("missing-link", missing_link, "Heading1", "link"),
            (
                "partial-direct-font",
                partial_direct_font,
                "Heading1",
                "explicitFonts",
            ),
        )
        for name, mutation, style_id, property_name in counterexamples:
            candidate = td / f"{name}.docx"

            def mutate_candidate(parts, mutation=mutation):
                rewrite_styles(parts, mutation)

            _rewrite_docx(
                native_normalized,
                candidate,
                mutate_candidate,
            )
            candidate_text, candidate_report = (
                docx2md.convert_with_report(
                    candidate,
                    report_provenance=False,
                )
            )
            assert candidate_text == source.read_text(encoding="utf8")
            candidate_projection = (
                candidate_report["visualStyleProjection"]
            )
            assert candidate_projection["state"] == "drifted", name
            assert (
                style_id,
                property_name,
            ) in {
                (item["styleId"], item["property"])
                for item in candidate_projection["diagnostics"]
            }, name
            assert all(
                item["code"]
                == "DOCX-ROUNDTRIP-VISUAL-STYLE-DRIFT"
                for item in candidate_projection["diagnostics"]
            ), name


def test_roundtrip_kitchen_sink():
    src = FIXTURES / "kitchen-sink.md"
    source_text = src.read_text(encoding="utf8")

    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "kitchen-sink.docx"
        theme_path = THEMES / "neutral.json"
        cfg = md2docx.deep_merge(md2docx.DEFAULTS, json.loads(theme_path.read_text(encoding="utf8")))
        md2docx.Converter(cfg, tpl_path=theme_path).convert(src, out)

        roundtripped = docx2md.convert(out)

    canonical_source = md2docx.canonicalize_markdown(source_text)
    assert roundtripped == canonical_source, (
        "supported kitchen-sink profile must round-trip at exact canonical "
        "string equality")

    # Structural spot-checks (fast, readable failure messages before the
    # coarser word-bag diff below).
    assert roundtripped.startswith("**CUI//TEST**"), "banner line must round-trip as the first line"
    assert "# Kitchen Sink" in roundtripped
    assert "### Level three heading" in roundtripped
    assert "#### Level four heading" in roundtripped
    assert "- First bullet item" in roundtripped
    assert "1. First literal numbered item" in roundtripped
    assert "| Name | Kind | Notes |" in roundtripped
    assert "| --- | --- | --- |" in roundtripped
    assert "def convert(src, out_path):" in roundtripped
    assert "```" in roundtripped
    assert "> A blockquote line that spans two source lines" in roundtripped
    assert "**bold**" in roundtripped and "*italic*" in roundtripped and "`code`" in roundtripped

    source_bag = _word_bag(_normalize_source_for_comparison(source_text))
    output_bag = _word_bag(roundtripped)

    lost = source_bag - output_bag
    invented = output_bag - source_bag
    assert not lost, f"tokens present in source but missing after round trip: {dict(lost)}"
    assert not invented, f"tokens present after round trip but absent from source: {dict(invented)}"


def test_canonicalizer_is_idempotent_and_normalizes_supported_profile():
    source = """\

#  Title

* one

This paragraph soft-wraps
onto another line with [docs](https://example.com/docs).

| **Name** | Value |
| :--- | ---: |
| alpha | `one` |

> quoted on
> two lines
"""
    expected = """\
# Title

- one

This paragraph soft-wraps onto another line with docs (https://example.com/docs).

| Name | Value |
| --- | --- |
| alpha | `one` |

> quoted on two lines
"""
    canonical = md2docx.canonicalize_markdown(source)
    assert canonical == expected
    assert md2docx.canonicalize_markdown(canonical) == canonical

    recovered_expected, recovered = md2docx.check_canonical_roundtrip(
        canonical)
    assert recovered_expected == canonical
    assert recovered == canonical

    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "noncanonical.md"
        out = Path(td) / "noncanonical.docx"
        src.write_text(source, encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(src, out)
        assert docx2md.convert(
            out, report_provenance=False) == canonical


def test_canonicalizer_refuses_lossy_constructs():
    cases = [
        ("  - nested\n", "MD-CANON-NESTED-LIST"),
        ("> > nested\n", "MD-CANON-NESTED-BLOCKQUOTE"),
        ("![alt](image.png)\n", "MD-CANON-UNSUPPORTED-IMAGE"),
        ("![alt][image]\n\n[image]: image.png\n",
         "MD-CANON-UNSUPPORTED-IMAGE"),
        ('<img src="image.png" alt="alt">\n',
         "MD-CANON-UNSUPPORTED-IMAGE"),
        ("```python\nprint('x')\n```\n", "MD-CANON-FENCE-INFO"),
        ("```\nunclosed\n", "MD-CANON-UNCLOSED-FENCE"),
        ("| A | B |\n| --- | --- |\n| a\\|b | c |\n",
         "MD-CANON-ESCAPED-PIPE"),
        ("| A | B |\n| --- |\n| a | b |\n",
         "MD-CANON-TABLE-SEPARATOR"),
        ("| A | B |\n| --- | --- |\n| `a|b` | c |\n",
         "MD-CANON-ESCAPED-PIPE"),
        ("| A | B |\n| --- | --- |\n| a |\n",
         "MD-CANON-TABLE"),
    ]
    for source, code in cases:
        try:
            md2docx.canonicalize_markdown(source)
            raise AssertionError(f"expected {code} for {source!r}")
        except md2docx.MarkdownCanonicalizationError as exc:
            assert exc.code == code


def test_canonicalizer_refuses_nested_inline_and_indented_code():
    cases = [
        (
            "Nested **bold *italic* end**.\n",
            "MD-CANON-UNSUPPORTED-INLINE",
        ),
        (
            "Nested *italic **bold** end*.\n",
            "MD-CANON-UNSUPPORTED-INLINE",
        ),
        (
            "    print('semantic code')\n",
            "MD-CANON-INDENTED-CODE",
        ),
        (
            "\tprint('semantic code')\n",
            "MD-CANON-INDENTED-CODE",
        ),
    ]
    for source, code in cases:
        try:
            md2docx.canonicalize_markdown(source)
            raise AssertionError(f"expected {code} for {source!r}")
        except md2docx.MarkdownCanonicalizationError as exc:
            assert exc.code == code


def test_canonicalizer_refuses_known_markdown_semantic_aliases():
    cases = [
        ("Setext one\n===\n", "MD-CANON-SETEXT-HEADING"),
        ("Setext two\n---\n", "MD-CANON-SETEXT-HEADING"),
        ("hard break  \nnext line\n", "MD-CANON-HARD-BREAK"),
        ("hard break\\\nnext line\n", "MD-CANON-HARD-BREAK"),
        ("Use ``multi-backtick`` code.\n",
         "MD-CANON-UNSUPPORTED-CODE-SPAN"),
        ("Use `unterminated code.\n",
         "MD-CANON-UNSUPPORTED-CODE-SPAN"),
        ("Use _underscore italic_.\n", "MD-CANON-UNSUPPORTED-INLINE"),
        ("Use __underscore bold__.\n", "MD-CANON-UNSUPPORTED-INLINE"),
    ]
    for source, code in cases:
        try:
            md2docx.canonicalize_markdown(source)
            raise AssertionError(f"expected {code} for {source!r}")
        except md2docx.MarkdownCanonicalizationError as exc:
            assert exc.code == code

    assert md2docx.canonicalize_markdown(
        "snake_case and `one code span`.\n"
    ) == "snake_case and `one code span`.\n"


def test_normalize_and_check_cli():
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "sample.md"
        src.write_text("* bullet\n\nsoft\nwrap\n", encoding="utf8")
        command = [sys.executable, str(REPO_ROOT / "md2docx.py")]

        normalized = subprocess.run(
            [*command, "--normalize", str(src)],
            check=True, capture_output=True, text=True)
        assert normalized.stdout == "- bullet\n\nsoft wrap\n"
        assert normalized.stderr == ""

        noncanonical = subprocess.run(
            [*command, "--check", str(src)],
            check=False, capture_output=True, text=True)
        assert noncanonical.returncode == 1
        assert "not canonical" in noncanonical.stderr

        src.write_text(normalized.stdout, encoding="utf8")
        checked = subprocess.run(
            [*command, "--check", str(src)],
            check=False, capture_output=True, text=True)
        assert checked.returncode == 0, checked.stderr
        assert "canonical round-trip OK" in checked.stdout


def test_forward_cli_is_byte_exact_and_publishes_atomically():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        command = [sys.executable, str(REPO_ROOT / "md2docx.py")]

        crlf = td / "crlf.md"
        crlf.write_bytes(b"# Title\r\n\r\nBody.\r\n")
        checked = subprocess.run(
            [*command, "--check", str(crlf)],
            check=False,
            capture_output=True,
            text=True,
        )
        assert checked.returncode == 1
        assert "not canonical" in checked.stderr

        invalid = td / "invalid.md"
        invalid.write_bytes(b"# Invalid\n\xff\n")
        for args in (
            ["--normalize", str(invalid)],
            ["--check", str(invalid)],
            [str(invalid), "-o", str(td / "invalid.docx")],
        ):
            failed = subprocess.run(
                [*command, *args],
                check=False,
                capture_output=True,
                text=True,
            )
            assert failed.returncode == 1
            assert "MD-CANON-INVALID-UTF8" in failed.stderr
            assert "Traceback" not in failed.stderr
        assert not (td / "invalid.docx").exists()

        canonical = td / "canonical.md"
        canonical.write_text("# Canonical\n", encoding="utf8")
        alias = subprocess.run(
            [
                *command,
                "--normalize",
                str(canonical),
                "-o",
                str(canonical),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert alias.returncode == 1
        assert "MD-CANON-PATH-ALIAS" in alias.stderr
        assert canonical.read_text(encoding="utf8") == "# Canonical\n"

        target = td / "existing.docx"
        target.write_bytes(b"unchanged")
        original_limit = md2docx.MAX_EMBEDDED_SOURCE_BYTES
        md2docx.MAX_EMBEDDED_SOURCE_BYTES = 1
        try:
            try:
                md2docx.Converter(md2docx.DEFAULTS).convert(
                    canonical, target)
                raise AssertionError("expected source-size refusal")
            except md2docx.MarkdownCanonicalizationError as exc:
                assert exc.code == "MD-CANON-SOURCE-LIMIT"
        finally:
            md2docx.MAX_EMBEDDED_SOURCE_BYTES = original_limit
        assert target.read_bytes() == b"unchanged"

        original_embed = md2docx.embed_source_snapshot

        def fail_embedding(*_args, **_kwargs):
            raise ValueError("forced embedding failure")

        md2docx.embed_source_snapshot = fail_embedding
        try:
            try:
                md2docx.Converter(md2docx.DEFAULTS).convert(
                    canonical, target)
                raise AssertionError("expected forced embedding failure")
            except ValueError as exc:
                assert "forced embedding failure" in str(exc)
        finally:
            md2docx.embed_source_snapshot = original_embed
        assert target.read_bytes() == b"unchanged"


def test_clis_normalize_routine_io_and_malformed_package_failures():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        source.write_text("# Source\n", encoding="utf8")
        forward = subprocess.run(
            [
                sys.executable,
                str(REPO_ROOT / "md2docx.py"),
                str(source),
                "-o",
                str(td / "missing" / "output.docx"),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert forward.returncode == 1
        assert "MD-CANON-IO" in forward.stderr
        assert "Traceback" not in forward.stderr

        malformed = td / "malformed.docx"
        with zipfile.ZipFile(malformed, "w") as package:
            package.writestr(
                "[Content_Types].xml",
                (
                    '<?xml version="1.0"?>'
                    '<Types xmlns="http://schemas.openxmlformats.org/package/'
                    '2006/content-types">'
                    '<Default Extension="xml" ContentType="application/xml"/>'
                    "</Types>"
                ),
            )
            package.writestr(
                "word/_rels/document.xml.rels",
                (
                    '<?xml version="1.0"?>'
                    '<Relationships xmlns="http://schemas.openxmlformats.org/'
                    'package/2006/relationships"/>'
                ),
            )
            package.writestr(
                "word/document.xml",
                (
                    '<?xml version="1.0"?>'
                    '<w:document xmlns:w="http://schemas.openxmlformats.org/'
                    'wordprocessingml/2006/main"><w:body><w:p><w:r>'
                    "<w:t>body</w:t></w:r></w:p></w:body></w:document>"
                ),
            )
        reverse_output = td / "malformed.md"
        reverse = subprocess.run(
            [
                sys.executable,
                str(REPO_ROOT / "docx2md.py"),
                str(malformed),
                "-o",
                str(reverse_output),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert reverse.returncode == 1
        assert "DOCX-ROUNDTRIP-PACKAGE" in reverse.stderr
        assert "Traceback" not in reverse.stderr
        assert not reverse_output.exists()


def test_docx_reverse_refuses_unrepresentable_content():
    def assert_refusal(doc, expected_code, path):
        doc.save(path)
        try:
            docx2md.convert(path, report_provenance=False)
            raise AssertionError(f"expected {expected_code}")
        except docx2md.RoundtripRefusalError as exc:
            assert exc.code == expected_code

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)

        tracked = docx2md.Document()
        paragraph = tracked.add_paragraph("visible")
        paragraph._p.append(OxmlElement("w:ins"))
        assert_refusal(
            tracked, "DOCX-ROUNDTRIP-TRACKED-CHANGES",
            td / "tracked.docx")

        image = docx2md.Document()
        run = image.add_paragraph().add_run()
        run._r.append(OxmlElement("w:drawing"))
        assert_refusal(
            image, "DOCX-ROUNDTRIP-IMAGE", td / "image.docx")

        nested = docx2md.Document()
        nested.add_paragraph("nested", style="List Bullet 2")
        assert_refusal(
            nested, "DOCX-ROUNDTRIP-NESTED-LIST",
            td / "nested.docx")

        unknown_style = docx2md.Document()
        unknown_style.add_paragraph("styled", style="Title")
        assert_refusal(
            unknown_style, "DOCX-ROUNDTRIP-STYLE",
            td / "unknown-style.docx")


def test_docx_reverse_preflight_covers_lossy_word_constructs():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "base.md"
        baseline = td / "base.docx"
        source.write_text("before after\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(source, baseline)

        def edited(name, edit):
            path = td / f"{name}.docx"
            shutil.copy2(baseline, path)
            document = docx2md.Document(str(path))
            edit(document)
            document.save(path)
            return path

        def insert_block_wrapper(tag, hidden_text):
            def apply(document):
                wrapper = OxmlElement(tag)
                if tag == "w:sdt":
                    content = OxmlElement("w:sdtContent")
                    content.append(_word_paragraph(hidden_text))
                    wrapper.append(content)
                else:
                    wrapper.append(_word_paragraph(hidden_text))
                body = document.element.body
                body.insert(len(body) - 1, wrapper)
            return apply

        _assert_roundtrip_refusal(
            edited(
                "custom-xml",
                insert_block_wrapper("w:customXml", "hidden custom XML"),
            ),
            "DOCX-ROUNDTRIP-UNSUPPORTED-ELEMENT",
        )
        _assert_roundtrip_refusal(
            edited(
                "content-control",
                insert_block_wrapper("w:sdt", "hidden content control"),
            ),
            "DOCX-ROUNDTRIP-CONTENT-CONTROL",
        )

        def add_hyperlink(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            paragraph.add_run("before ")
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("w:anchor"), "bookmark")
            hyperlink.append(_word_text_run("linked text"))
            paragraph._p.append(hyperlink)
            paragraph.add_run(" after")

        _assert_roundtrip_refusal(
            edited("hyperlink", add_hyperlink),
            "DOCX-ROUNDTRIP-HYPERLINK",
        )

        def add_field(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), " DATE ")
            field.append(_word_text_run("field result"))
            paragraph._p.append(field)

        _assert_roundtrip_refusal(
            edited("field", add_field),
            "DOCX-ROUNDTRIP-FIELD",
        )

        def add_complex_field(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            for field_type in ("begin", "separate"):
                run = OxmlElement("w:r")
                field_char = OxmlElement("w:fldChar")
                field_char.set(qn("w:fldCharType"), field_type)
                run.append(field_char)
                paragraph._p.append(run)
                if field_type == "begin":
                    instruction_run = OxmlElement("w:r")
                    instruction = OxmlElement("w:instrText")
                    instruction.text = " DATE "
                    instruction_run.append(instruction)
                    paragraph._p.append(instruction_run)
            paragraph._p.append(_word_text_run("field result"))
            end_run = OxmlElement("w:r")
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            end_run.append(end)
            paragraph._p.append(end_run)

        _assert_roundtrip_refusal(
            edited("complex-field", add_complex_field),
            "DOCX-ROUNDTRIP-FIELD",
        )

        def add_break(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            paragraph.add_run("before")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            paragraph.add_run("after")

        _assert_roundtrip_refusal(
            edited("break", add_break),
            "DOCX-ROUNDTRIP-BREAK",
        )

        def add_tab(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            paragraph.add_run("before")
            tab_run = paragraph.add_run()
            tab_run._r.append(OxmlElement("w:tab"))
            paragraph.add_run("after")

        _assert_roundtrip_refusal(
            edited("tab", add_tab),
            "DOCX-ROUNDTRIP-BREAK",
        )

        def add_page_break_before(document):
            paragraph = document.paragraphs[0]
            paragraph._p.get_or_add_pPr().append(
                OxmlElement("w:pageBreakBefore"))

        _assert_roundtrip_refusal(
            edited("page-break-before", add_page_break_before),
            "DOCX-ROUNDTRIP-BREAK",
        )

        def add_last_rendered_page_break(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            paragraph.add_run("before")
            marker_run = paragraph.add_run()
            marker_run._r.append(
                OxmlElement("w:lastRenderedPageBreak"))
            paragraph.add_run(" after")

        pagination_artifact = edited(
            "last-rendered-page-break",
            add_last_rendered_page_break,
        )
        recovered, report = docx2md.convert_with_report(
            pagination_artifact, report_provenance=False)
        assert recovered == "before after\n"
        assert report["state"] == "exact-supported-profile"

        def add_header_revision(document):
            paragraph = document.sections[0].header.paragraphs[0]
            insertion = OxmlElement("w:ins")
            insertion.set(qn("w:id"), "1")
            insertion.set(qn("w:author"), "Reviewer")
            insertion.append(_word_text_run("tracked header"))
            paragraph._p.append(insertion)

        _assert_roundtrip_refusal(
            edited("header-revision", add_header_revision),
            "DOCX-ROUNDTRIP-TRACKED-CHANGES",
        )

        def add_header_shape(document):
            paragraph = document.sections[0].header.paragraphs[0]
            paragraph._p.append(parse_xml(
                '<w:r '
                'xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main" '
                'xmlns:v="urn:schemas-microsoft-com:vml">'
                '<w:pict><v:shape id="banner"><v:textbox>'
                '<w:txbxContent><w:p><w:r><w:t>shape banner</w:t>'
                '</w:r></w:p></w:txbxContent>'
                '</v:textbox></v:shape></w:pict></w:r>'
            ))

        _assert_roundtrip_refusal(
            edited("header-shape", add_header_shape),
            "DOCX-ROUNDTRIP-TEXTBOX",
        )

        def add_plain_header(document):
            document.sections[0].header.paragraphs[0].text = "arbitrary"

        _assert_roundtrip_refusal(
            edited("plain-header", add_plain_header),
            "DOCX-ROUNDTRIP-HEADER",
        )

        def add_outer_space_to_banner(document):
            text = " CUI//TEST "
            document.sections[0].header.paragraphs[0].text = text
            document.sections[0].footer.paragraphs[0].text = text

        _assert_roundtrip_refusal(
            edited("outer-space-banner", add_outer_space_to_banner),
            "DOCX-ROUNDTRIP-HEADER",
        )

        def add_bold_italic(document):
            paragraph = document.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run("both")
            run.bold = True
            run.italic = True

        _assert_roundtrip_refusal(
            edited("bold-italic", add_bold_italic),
            "DOCX-ROUNDTRIP-INLINE",
        )

        def add_underline(document):
            document.paragraphs[0].runs[0].underline = True

        _assert_roundtrip_refusal(
            edited("underline", add_underline),
            "DOCX-ROUNDTRIP-INLINE",
        )

        def add_plain_markdown(document):
            document.paragraphs[0].text = "**plain Word text**"

        _assert_roundtrip_refusal(
            edited("plain-markdown", add_plain_markdown),
            "DOCX-ROUNDTRIP-INLINE",
        )

        def add_heading_prefix(document):
            document.paragraphs[0].text = "# normal-style text"

        _assert_roundtrip_refusal(
            edited("normal-heading", add_heading_prefix),
            "DOCX-ROUNDTRIP-NONCANONICAL",
        )

        def add_trailing_space(document):
            document.paragraphs[0].text = "trailing "

        trailing_space = edited("trailing-space", add_trailing_space)
        recovered, report = docx2md.convert_with_report(
            trailing_space, report_provenance=False)
        assert recovered == "trailing\n"
        assert report["state"] == "normalized-supported-profile"
        assert report["semanticNormalization"]["state"] == "normalized"
        events = report["semanticNormalization"]["events"]
        assert len(events) == 1
        assert events[0] == {
            "code": "DOCX-ROUNDTRIP-TRAILING-ASCII-SPACE",
            "severity": "warning",
            "story": "word/document.xml",
            "paragraphIndex": 0,
            "styleId": "Normal",
            "styleName": "Normal",
            "edge": "trailing",
            "codePoint": "U+0020",
            "count": 1,
            "inputTextSha256": hashlib.sha256(
                b"trailing "
            ).hexdigest(),
            "message": (
                "removed 1 trailing U+0020 character from body paragraph 0"
            ),
        }
        assert events[0] in report["diagnostics"]

        def add_leading_space(document):
            document.paragraphs[0].text = " leading"

        _assert_roundtrip_refusal(
            edited("leading-space", add_leading_space),
            "DOCX-ROUNDTRIP-NONCANONICAL",
        )

        def add_trailing_nbsp(document):
            document.paragraphs[0].text = "trailing\u00a0"

        _assert_roundtrip_refusal(
            edited("trailing-nbsp", add_trailing_nbsp),
            "DOCX-ROUNDTRIP-NONCANONICAL",
        )

        def add_whitespace_only(document):
            document.paragraphs[0].text = " "

        _assert_roundtrip_refusal(
            edited("whitespace-only", add_whitespace_only),
            "DOCX-ROUNDTRIP-NONCANONICAL",
        )

        footnote_path = edited(
            "footnote-reference",
            lambda document: document.paragraphs[0]._p.append(
                parse_xml(
                    '<w:r xmlns:w="http://schemas.openxmlformats.org/'
                    'wordprocessingml/2006/main">'
                    '<w:footnoteReference w:id="1"/></w:r>'
                )
            ),
        )

        def add_footnote_graph(parts):
            rels = (
                '<Relationship Id="rId999" '
                'Type="http://schemas.openxmlformats.org/'
                'officeDocument/2006/relationships/footnotes" '
                'Target="footnotes.xml"/>'
            ).encode()
            parts["word/_rels/document.xml.rels"] = parts[
                "word/_rels/document.xml.rels"
            ].replace(b"</Relationships>", rels + b"</Relationships>")
            override = (
                '<Override PartName="/word/footnotes.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.footnotes+xml"/>'
            ).encode()
            parts["[Content_Types].xml"] = parts[
                "[Content_Types].xml"
            ].replace(b"</Types>", override + b"</Types>")
            parts["word/footnotes.xml"] = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main">'
                '<w:footnote w:id="1"><w:p><w:r>'
                '<w:t>footnote text</w:t></w:r></w:p></w:footnote>'
                '</w:footnotes>'
            ).encode()

        _rewrite_docx(
            footnote_path, footnote_path, add_footnote_graph)
        _assert_roundtrip_refusal(
            footnote_path, "DOCX-ROUNDTRIP-NOTE")

        table_source = td / "table.md"
        table_path = td / "table.docx"
        table_source.write_text(
            "| A | B |\n| --- | --- |\n| x | y |\n",
            encoding="utf8",
        )
        md2docx.Converter(md2docx.DEFAULTS).convert(
            table_source, table_path)
        table_document = docx2md.Document(str(table_path))
        table_document.tables[0].cell(1, 0).text = " cell "
        table_document.save(table_path)
        _assert_roundtrip_refusal(
            table_path, "DOCX-ROUNDTRIP-NONCANONICAL")

        md2docx.Converter(md2docx.DEFAULTS).convert(
            table_source, table_path)
        table_document = docx2md.Document(str(table_path))
        table_document.tables[0].cell(1, 0).text = "a|b"
        table_document.save(table_path)
        _assert_roundtrip_refusal(
            table_path, "DOCX-ROUNDTRIP-NONCANONICAL")

        code_source = td / "code.md"
        code_path = td / "code.docx"
        code_source.write_text("```\nalpha\n```\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(
            code_source, code_path)
        code_document = docx2md.Document(str(code_path))
        next(
            paragraph for paragraph in code_document.paragraphs
            if paragraph.text == "alpha"
        ).add_run("OMEGA")
        code_document.save(code_path)
        recovered, report = docx2md.convert_with_report(
            code_path, report_provenance=False)
        assert recovered == "```\nalphaOMEGA\n```\n"
        assert report["state"] == "exact-supported-profile"

        code_document = docx2md.Document(str(code_path))
        next(
            paragraph for paragraph in code_document.paragraphs
            if paragraph.text == "alphaOMEGA"
        ).add_run(" ")
        code_document.save(code_path)
        _assert_roundtrip_refusal(
            code_path, "DOCX-ROUNDTRIP-NONCANONICAL")


def test_docx_reverse_refuses_alternate_header_footer_content():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        original = td / "source.docx"
        source.write_text("# Source\n\nBody.\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(source, original)

        story_names = (
            "first_page_header",
            "even_page_header",
            "first_page_footer",
            "even_page_footer",
        )
        for story_name in story_names:
            edited = td / f"{story_name}.docx"
            document = docx2md.Document(str(original))
            section = document.sections[0]
            if story_name.startswith("first_page"):
                section.different_first_page_header_footer = True
            if story_name.startswith("even_page"):
                document.settings.odd_and_even_pages_header_footer = True
            story = getattr(section, story_name)
            story.paragraphs[0].text = f"unsupported {story_name}"
            document.save(edited)
            _assert_roundtrip_refusal(
                edited, "DOCX-ROUNDTRIP-HEADER")


def test_story_preflight_follows_renamed_header_relationship():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        original = td / "source.docx"
        renamed = td / "renamed-header.docx"
        source.write_text(
            "**CUI//TEST**\n\n# Source\n\nBody.\n",
            encoding="utf8",
        )
        md2docx.Converter(md2docx.DEFAULTS).convert(source, original)

        with zipfile.ZipFile(original, "r") as package:
            infos = package.infolist()
            parts = {
                info.filename: package.read(info)
                for info in infos
            }

        rels_name = "word/_rels/document.xml.rels"
        relationships = re.findall(
            br"<Relationship\b[^>]*/>", parts[rels_name])
        header_type = (
            b"http://schemas.openxmlformats.org/officeDocument/"
            b"2006/relationships/header"
        )
        header_relationship = next(
            relationship
            for relationship in relationships
            if b'Type="' + header_type + b'"' in relationship
        )
        target = re.search(
            br'Target="([^"]+)"', header_relationship).group(1).decode()
        assert "/" not in target and target.endswith(".xml")
        old_name = f"word/{target}"
        new_target = "stories/renamed-header.xml"
        new_name = f"word/{new_target}"
        assert docx2md.STORY_PART_PATTERN.fullmatch(new_name) is None

        replacement_relationship = header_relationship.replace(
            f'Target="{target}"'.encode(),
            f'Target="{new_target}"'.encode(),
        )
        parts[rels_name] = parts[rels_name].replace(
            header_relationship,
            replacement_relationship,
            1,
        )
        old_part_name = f'PartName="/{old_name}"'.encode()
        new_part_name = f'PartName="/{new_name}"'.encode()
        assert parts["[Content_Types].xml"].count(old_part_name) == 1
        parts["[Content_Types].xml"] = parts[
            "[Content_Types].xml"
        ].replace(old_part_name, new_part_name, 1)

        header_xml = parts.pop(old_name)
        insertion = (
            b'<w:ins w:id="1" w:author="Reviewer">'
            b"<w:r><w:t>tracked renamed header</w:t></w:r>"
            b"</w:ins>"
        )
        assert b"</w:p>" in header_xml
        parts[new_name] = header_xml.replace(
            b"</w:p>", insertion + b"</w:p>", 1)

        with zipfile.ZipFile(renamed, "w") as package:
            for info in infos:
                if info.filename == old_name:
                    continue
                package.writestr(info, parts[info.filename])
            package.writestr(new_name, parts[new_name])

        # The relationship and content type remain valid despite the
        # nonconventional part path; python-docx must still reopen the package.
        docx2md.Document(str(renamed))
        exc = _assert_roundtrip_refusal(
            renamed, "DOCX-ROUNDTRIP-TRACKED-CHANGES")
        assert new_name in exc.message


def test_docx_cli_refuses_aliases_and_stages_all_outputs():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        source = td / "source.md"
        docx_path = td / "source.docx"
        current = td / "current.md"
        source.write_text("# Source\n\nBody.\n", encoding="utf8")
        current.write_text("# Source\n\nBody.\n", encoding="utf8")
        md2docx.Converter(md2docx.DEFAULTS).convert(
            source, docx_path)
        original_docx = docx_path.read_bytes()
        command = [sys.executable, str(REPO_ROOT / "docx2md.py")]

        same = td / "same.md"
        other = td / "other.md"
        collision_args = [
            [str(docx_path), "-o", str(docx_path)],
            [
                str(docx_path), "-o", str(same),
                "--report", str(same),
            ],
            [
                str(docx_path), "-o", str(same),
                "--base-out", str(same),
            ],
            [
                str(docx_path),
                "--merge-current", str(current),
                "-o", str(current),
            ],
            [
                str(docx_path),
                "--merge-current", str(current),
                "-o", str(other),
                "--report", str(current),
            ],
            [
                str(docx_path), "-o", str(other),
                "--base-out", str(same),
                "--report", str(same),
            ],
            [
                str(docx_path), "-o", str(other),
                "--report", str(docx_path),
            ],
        ]
        for args in collision_args:
            failed = subprocess.run(
                [*command, *args],
                check=False,
                capture_output=True,
                text=True,
            )
            assert failed.returncode == 1, failed.stderr
            assert "DOCX-ROUNDTRIP-PATH-ALIAS" in failed.stderr
        assert docx_path.read_bytes() == original_docx
        assert current.read_text(encoding="utf8") == "# Source\n\nBody.\n"

        case_output = td / "Result.md"
        case_report = td / "result.md"
        assert md2docx._paths_alias(case_output, case_report)
        failed = subprocess.run(
            [
                *command,
                str(docx_path),
                "-o",
                str(case_output),
                "--report",
                str(case_report),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert failed.returncode == 1
        assert "DOCX-ROUNDTRIP-PATH-ALIAS" in failed.stderr
        assert not case_output.exists()
        assert not case_report.exists()

        composed_output = td / "Café.md"
        decomposed_report = td / "Cafe\u0301.md"
        assert md2docx._paths_alias(composed_output, decomposed_report)
        failed = subprocess.run(
            [
                *command,
                str(docx_path),
                "-o",
                str(composed_output),
                "--report",
                str(decomposed_report),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert failed.returncode == 1
        assert "DOCX-ROUNDTRIP-PATH-ALIAS" in failed.stderr
        assert not composed_output.exists()
        assert not decomposed_report.exists()

        staged_output = td / "staged.md"
        missing_report = td / "missing" / "report.json"
        failed = subprocess.run(
            [
                *command,
                str(docx_path),
                "-o",
                str(staged_output),
                "--report",
                str(missing_report),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert failed.returncode == 1
        assert "DOCX-ROUNDTRIP-IO" in failed.stderr
        assert not staged_output.exists()
        assert not missing_report.exists()


def test_multi_output_publication_rolls_back_replacement_failure():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        first = td / "first.md"
        second = td / "second.md"
        newly_created = td / "new.json"
        first.write_bytes(b"old first")
        second.write_bytes(b"old second")

        real_replace = docx2md._atomic_replace
        staged_replacements = 0

        def fail_second_staged_replace(source, destination):
            nonlocal staged_replacements
            if Path(source).suffix == ".tmp":
                staged_replacements += 1
                if staged_replacements == 2:
                    raise OSError("forced second replacement failure")
            real_replace(source, destination)

        docx2md._atomic_replace = fail_second_staged_replace
        try:
            try:
                docx2md._publish_outputs([
                    (first, b"new first"),
                    (second, b"new second"),
                    (newly_created, b"new third"),
                ])
                raise AssertionError("expected publication failure")
            except OSError as exc:
                assert "forced second replacement failure" in str(exc)
        finally:
            docx2md._atomic_replace = real_replace

        assert first.read_bytes() == b"old first"
        assert second.read_bytes() == b"old second"
        assert not newly_created.exists()
        assert not list(td.glob(".*.tmp"))
        assert not list(td.glob(".*.bak"))


def test_link_demotion():
    """CONTRACT:C3-ROUNDTRIP.1.2 retained link-demotion boundary:
    a relative link's target is dropped (label only survives); an
    absolute link keeps both label and URL."""
    assert md2docx.demote_links("[docs](ROADMAP.md)") == "docs"
    assert md2docx.demote_links("[docs](https://example.com/docs)") == "docs (https://example.com/docs)"
    # label == url: no duplication (label already IS the full URL string)
    assert md2docx.demote_links("[https://example.com](https://example.com)") == "https://example.com"
    assert md2docx.demote_links("plain text, no links") == "plain text, no links"


def test_no_footer_flag():
    """--no-footer must actually suppress footer text (footer_on threads
    through Converter.convert to the footer-emission branch)."""
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "doc.md"
        src.write_text("# Title\n\nBody text.\n", encoding="utf8")
        cfg = md2docx.deep_merge(md2docx.DEFAULTS, {"footer": {"text": "TEST FOOTER"}})

        with_footer = td / "with-footer.docx"
        md2docx.Converter(cfg).convert(src, with_footer, footer_on=True)
        doc = docx2md.Document(str(with_footer))
        assert doc.sections[0].footer.paragraphs[0].text == "TEST FOOTER"

        without_footer = td / "without-footer.docx"
        md2docx.Converter(cfg).convert(src, without_footer, footer_on=False)
        doc2 = docx2md.Document(str(without_footer))
        assert doc2.sections[0].footer.paragraphs[0].text == ""


# ---------------------------------------------------------------------------
# Runner — no test framework required.
# ---------------------------------------------------------------------------

def main():
    tests = [(name, fn) for name, fn in sorted(globals().items())
              if name.startswith("test_") and callable(fn)]
    failures = []
    for name, fn in tests:
        try:
            fn()
        except Exception:
            failures.append(name)
            print(f"FAIL {name}")
            traceback.print_exc()
        else:
            print(f"PASS {name}")

    print()
    print(f"{len(tests) - len(failures)}/{len(tests)} passed")
    if failures:
        print("Failed: " + ", ".join(failures))
        sys.exit(1)
    sys.exit(0)


if __name__ == "__main__":
    main()
