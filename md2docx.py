#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""md2docx — Markdown → styled DOCX, themed by a JSON template.

CONTRACT:C1-THEME-SCHEMA.1.0
CONTRACT:C2-PROVENANCE.2.0
CONTRACT:C3-ROUNDTRIP.1.1

Usage:
  md2docx.py file.md [more.md ...]            # writes <name>.docx next to each source
  md2docx.py -t theme.json file.md            # explicit template
  md2docx.py -o outdir file.md [more.md ...]  # write outputs into a directory
  md2docx.py -o out.docx file.md              # single input, explicit output path
  md2docx.py --no-footer file.md              # suppress footer text

Template resolution (see architecture/CONTRACT-C1-THEME-SCHEMA.1.0.md): the
--template flag wins; otherwise `md2docx-template.json` next to this script;
otherwise `themes/neutral.json` next to this script; otherwise the built-in
neutral defaults below. Any template key may be omitted — it deep-merges
over the defaults.

Markdown support: h1–h4; bullet lists; literal numbered lists; N-column pipe
tables (first row = shaded header); fenced code blocks; inline `code`, **bold**,
*italic*; links (relative links keep their label, absolute URLs appended in
parens); blockquotes; horizontal rules and HTML comments skipped; soft-wrapped
lines joined into one paragraph. A leading `**CUI...**`-style banner line is
promoted to the page header and footer (marking convention) and replaces the
footer text — see the theme's `cui_banner` block and `themes/marked-docs.json`
for an example.

Requires: python-docx  (pip install python-docx)
"""
import argparse
import base64
import datetime
import hashlib
import json
import re
import sys
import tempfile
import unicodedata
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree

TOOL_VERSION = "0.2.0"
EMBEDDED_SOURCE_SCHEMA = "office180-md-source/0.1"
EMBEDDED_SOURCE_NAMESPACE = "urn:office180:md-source:0.1"
CUSTOM_XML_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml"
)
CUSTOM_XML_PROPS_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXmlProps"
)
CUSTOM_XML_PROPS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.customXmlProperties+xml"
)
PACKAGE_RELATIONSHIPS_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)
CONTENT_TYPES_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/content-types"
)
MAX_EMBEDDED_SOURCE_BYTES = 8 * 1024 * 1024
MAX_DOCX_PARTS = 10_000
MAX_DOCX_UNCOMPRESSED_BYTES = 256 * 1024 * 1024

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEFAULTS = {
    "fonts": {"body": "Calibri", "mono": "Consolas"},
    "base": {"size_pt": 10.5, "color": "1A1A1A", "space_after_pt": 6},
    "headings": {
        "h1": {"size_pt": 18, "color": "000000"},
        "h2": {"size_pt": 14, "color": "000000"},
        "h3": {"size_pt": 11.5, "color": "000000"},
        "h4": {"size_pt": 10.5, "color": "333333"},
        "bold": True,
        "space_before_pt": 12,
        "space_after_pt": 5,
    },
    "table": {
        "style": "Table Grid",
        "header_fill": "EEEEEE",
        "header_color": "222222",
        "header_bold": True,
        "cell_size_pt": 9.5,
    },
    "code": {"fill": "F5F5F5", "color": "333333", "block_size_pt": 8.5},
    "blockquote": {
        "color": "555555",
        "size_pt": 9.5,
        "border_color": "888888",
        "border_size_eighths": 18,
        "indent_in": 0.3,
    },
    "footer": {"text": "", "color": "999999", "size_pt": 10, "bold": True},
    "cui_banner": {"detect": True, "size_pt": 10, "bold": True, "color": "000000"},
}


def deep_merge(base, over):
    out = dict(base)
    for k, v in (over or {}).items():
        out[k] = deep_merge(base[k], v) if isinstance(v, dict) and isinstance(base.get(k), dict) else v
    return out


def rgb(hexstr):
    return RGBColor(*(int(hexstr[i:i + 2], 16) for i in (0, 2, 4)))


def _shd(el_get, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    el_get().append(shd)


def shade_cell(cell, fill): _shd(cell._tc.get_or_add_tcPr, fill)
def shade_para(p, fill): _shd(p._p.get_or_add_pPr, fill)
def shade_run(r, fill): _shd(r._r.get_or_add_rPr, fill)


def left_border(p, color, size_eighths):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths))
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def demote_links(text):
    def repl(m):
        label, url = m.group(1), m.group(2)
        if url.startswith("http") and label != url:
            return f"{label} ({url})"
        return label
    return LINK_RE.sub(repl, text)


class MarkdownCanonicalizationError(ValueError):
    """A supported-profile Markdown normalization refusal."""

    def __init__(self, code, line, message):
        self.code = code
        self.line = line
        self.message = message
        super().__init__(f"{code} at line {line}: {message}")


@dataclass(frozen=True)
class MarkdownBlock:
    """Small block AST used only for deterministic canonicalization."""

    kind: str
    value: object
    line: int


def _find_single_star(text, start):
    offset = start
    while True:
        offset = text.find("*", offset)
        if offset < 0:
            return -1
        if (
            (offset == 0 or text[offset - 1] != "*")
            and (offset + 1 == len(text) or text[offset + 1] != "*")
        ):
            return offset
        offset += 1


def _validate_separate_inline_spans(text, line):
    """Reject nested/combined inline syntax the run writer cannot preserve."""
    offset = 0
    while offset < len(text):
        if text[offset] == "`":
            closing = text.find("`", offset + 1)
            if closing >= 0:
                offset = closing + 1
                continue
        if text.startswith("**", offset):
            closing = text.find("**", offset + 2)
            if closing >= 0:
                content = text[offset + 2:closing]
                if "*" in content or "`" in content:
                    raise MarkdownCanonicalizationError(
                        "MD-CANON-UNSUPPORTED-INLINE",
                        line,
                        "combined/nested inline spans are outside the "
                        "supported profile",
                    )
                offset = closing + 2
                continue
        if text[offset] == "*":
            closing = _find_single_star(text, offset + 1)
            if closing >= 0:
                content = text[offset + 1:closing]
                if "*" in content or "`" in content:
                    raise MarkdownCanonicalizationError(
                        "MD-CANON-UNSUPPORTED-INLINE",
                        line,
                        "combined/nested inline spans are outside the "
                        "supported profile",
                    )
                offset = closing + 1
                continue
        offset += 1


def _canonical_inline(text, line):
    if re.search(r"!\[", text) or re.search(
            r"<\s*img(?:\s|/?>)", text, flags=re.IGNORECASE):
        raise MarkdownCanonicalizationError(
            "MD-CANON-UNSUPPORTED-IMAGE", line,
            "images are outside the supported Markdown/DOCX profile")
    if "***" in text:
        raise MarkdownCanonicalizationError(
            "MD-CANON-UNSUPPORTED-INLINE", line,
            "combined/nested emphasis is not accepted by the forward parser")
    canonical = demote_links(text.strip())
    if "``" in canonical or canonical.count("`") % 2:
        raise MarkdownCanonicalizationError(
            "MD-CANON-UNSUPPORTED-CODE-SPAN",
            line,
            "multi-backtick and unterminated code spans are outside the "
            "supported profile",
        )
    if (
        re.search(r"(?<!\w)__(?=\S).*?(?<=\S)__(?!\w)", canonical)
        or re.search(r"(?<!\w)_(?=\S).*?(?<=\S)_(?!\w)", canonical)
    ):
        raise MarkdownCanonicalizationError(
            "MD-CANON-UNSUPPORTED-INLINE",
            line,
            "underscore emphasis is outside the supported profile",
        )
    _validate_separate_inline_spans(canonical, line)
    return canonical


def _table_cells(line, line_number):
    stripped = line.strip()
    if not stripped.endswith("|"):
        raise MarkdownCanonicalizationError(
            "MD-CANON-TABLE", line_number,
            "supported pipe-table rows must start and end with '|'")
    if "\\|" in stripped:
        raise MarkdownCanonicalizationError(
            "MD-CANON-ESCAPED-PIPE", line_number,
            "escaped pipes in table cells are not supported by md2docx")
    if re.search(r"`[^`]*\|[^`]*`", stripped):
        raise MarkdownCanonicalizationError(
            "MD-CANON-ESCAPED-PIPE", line_number,
            "pipes inside inline-code table cells are not supported")
    return [cell.strip() for cell in stripped[1:-1].split("|")]


def _starts_block(line):
    stripped = line.strip()
    return (
        not stripped
        or stripped.startswith(("```", "|", ">", "<!--"))
        or re.fullmatch(r"-{3,}", stripped) is not None
        or re.match(r"^#{1,}\s", stripped) is not None
        or re.match(r"^[-*]\s+", stripped) is not None
        or re.match(r"^\d+\.\s+", stripped) is not None
    )


def _preflight_markdown_source_lines(lines):
    """Refuse line syntax whose semantics are lost by the small block parser."""
    in_fence = False
    setext_underline = re.compile(r" {0,3}(?:=+|-+)[ \t]*")
    for index, raw in enumerate(lines):
        stripped = raw.strip()
        leading_spaces = len(raw) - len(raw.lstrip(" "))
        if stripped == "```":
            if raw.startswith("\t") or leading_spaces >= 4:
                raise MarkdownCanonicalizationError(
                    "MD-CANON-INDENTED-CODE",
                    index + 1,
                    "indented fences are outside the supported profile",
                )
            in_fence = not in_fence
            continue
        if in_fence or not stripped:
            continue
        if raw.startswith("\t") or leading_spaces >= 4:
            raise MarkdownCanonicalizationError(
                "MD-CANON-INDENTED-CODE",
                index + 1,
                "indented code is outside the fenced-code-only profile",
            )
        if raw.endswith("  ") or raw.endswith("\\"):
            raise MarkdownCanonicalizationError(
                "MD-CANON-HARD-BREAK",
                index + 1,
                "explicit Markdown hard breaks are outside the supported "
                "paragraph profile",
            )
        if (
            index + 1 < len(lines)
            and not _starts_block(raw)
            and setext_underline.fullmatch(lines[index + 1])
        ):
            raise MarkdownCanonicalizationError(
                "MD-CANON-SETEXT-HEADING",
                index + 1,
                "setext headings are outside the ATX-heading-only profile",
            )


def parse_markdown_blocks(text):
    """Parse the intentionally small supported Markdown profile.

    This is not a general Markdown parser. Unsupported constructs that the
    forward converter would otherwise flatten or lose are refused with stable
    codes so canonical equality never over-claims fidelity.
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    if normalized.startswith("\ufeff"):
        normalized = normalized[1:]
    lines = normalized.split("\n")
    _preflight_markdown_source_lines(lines)
    blocks = []
    i = 0

    first_nonblank = next(
        (idx for idx, line in enumerate(lines) if line.strip()), None)
    if first_nonblank is not None:
        match = re.fullmatch(
            r"\*\*(CUI[^*]*)\*\*", lines[first_nonblank].strip())
        if match:
            blocks.append(MarkdownBlock(
                "banner", match.group(1), first_nonblank + 1))
            i = first_nonblank + 1

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        line_number = i + 1
        if not stripped:
            i += 1
            continue
        if stripped.startswith("<!--"):
            if not stripped.endswith("-->"):
                raise MarkdownCanonicalizationError(
                    "MD-CANON-HTML-COMMENT", line_number,
                    "only complete single-line HTML comments are droppable")
            i += 1
            continue
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue
        if re.match(r"^#{5,}\s", stripped):
            raise MarkdownCanonicalizationError(
                "MD-CANON-HEADING-LEVEL", line_number,
                "only heading levels 1 through 4 are supported")
        if raw != raw.lstrip() and (
                re.match(r"^[-*]\s+", stripped)
                or re.match(r"^\d+\.\s+", stripped)):
            raise MarkdownCanonicalizationError(
                "MD-CANON-NESTED-LIST", line_number,
                "nested/indented lists are outside the supported profile")

        if stripped.startswith("```"):
            if stripped != "```":
                raise MarkdownCanonicalizationError(
                    "MD-CANON-FENCE-INFO", line_number,
                    "fence language/info strings are not preserved yet")
            start = line_number
            code_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "```":
                code_lines.append(lines[i].rstrip())
                i += 1
            if i == len(lines):
                raise MarkdownCanonicalizationError(
                    "MD-CANON-UNCLOSED-FENCE", start,
                    "fenced code block has no closing fence")
            blocks.append(MarkdownBlock("code", tuple(code_lines), start))
            i += 1
            continue

        if stripped.startswith("|"):
            start = line_number
            rows = []
            row_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_table_cells(lines[i], i + 1))
                row_lines.append(i + 1)
                i += 1
            if len(rows) < 2 or not all(
                    re.fullmatch(r":?-{3,}:?", cell)
                    for cell in rows[1]):
                raise MarkdownCanonicalizationError(
                    "MD-CANON-TABLE-SEPARATOR", start,
                    "a pipe table requires a separator row after its header")
            if len(rows[1]) != len(rows[0]):
                raise MarkdownCanonicalizationError(
                    "MD-CANON-TABLE-SEPARATOR", row_lines[1],
                    "table separator width must match the header width")
            if any(len(row) != len(rows[0]) for row in rows[2:]):
                raise MarkdownCanonicalizationError(
                    "MD-CANON-TABLE", start,
                    "every supported table row must have the header width")
            content_rows = [rows[0], *rows[2:]]
            canonical_rows = []
            for row_index, row in enumerate(content_rows):
                source_line = row_lines[0] if row_index == 0 else row_lines[row_index + 1]
                cells = [
                    _canonical_inline(cell, source_line)
                    for cell in row
                ]
                if row_index == 0:
                    cells = [cell.replace("**", "") for cell in cells]
                canonical_rows.append(tuple(cells))
            blocks.append(MarkdownBlock(
                "table", tuple(canonical_rows), start))
            continue

        if stripped.startswith(">"):
            if stripped[1:].lstrip().startswith(">"):
                raise MarkdownCanonicalizationError(
                    "MD-CANON-NESTED-BLOCKQUOTE", line_number,
                    "nested blockquotes are outside the supported profile")
            parts = []
            start = line_number
            while i < len(lines) and lines[i].strip().startswith(">"):
                current = lines[i].strip()
                if current[1:].lstrip().startswith(">"):
                    raise MarkdownCanonicalizationError(
                        "MD-CANON-NESTED-BLOCKQUOTE", i + 1,
                        "nested blockquotes are outside the supported profile")
                parts.append(current[1:].strip())
                i += 1
            blocks.append(MarkdownBlock(
                "blockquote",
                _canonical_inline(" ".join(parts), start),
                start))
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            blocks.append(MarkdownBlock(
                "heading",
                (len(heading.group(1)),
                 _canonical_inline(heading.group(2), line_number)),
                line_number))
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            blocks.append(MarkdownBlock(
                "bullet", _canonical_inline(bullet.group(1), line_number),
                line_number))
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            blocks.append(MarkdownBlock(
                "numbered",
                (numbered.group(1),
                 _canonical_inline(numbered.group(2), line_number)),
                line_number))
            i += 1
            continue

        start = line_number
        parts = [stripped]
        i += 1
        while i < len(lines) and not _starts_block(lines[i]):
            parts.append(lines[i].strip())
            i += 1
        blocks.append(MarkdownBlock(
            "paragraph", _canonical_inline(" ".join(parts), start), start))

    return tuple(blocks)


def canonicalize_markdown(text):
    """Return the single canonical spelling of supported-profile Markdown."""
    rendered = []
    for block in parse_markdown_blocks(text):
        if block.kind == "banner":
            rendered.append(f"**{block.value}**")
        elif block.kind == "heading":
            level, value = block.value
            rendered.append("#" * level + " " + value)
        elif block.kind == "bullet":
            rendered.append("- " + block.value)
        elif block.kind == "numbered":
            number, value = block.value
            rendered.append(f"{number}. {value}")
        elif block.kind == "blockquote":
            rendered.append("> " + block.value)
        elif block.kind == "code":
            rendered.append("```\n" + "\n".join(block.value) + "\n```")
        elif block.kind == "table":
            rows = block.value
            ncols = len(rows[0])
            table_lines = ["| " + " | ".join(rows[0]) + " |"]
            table_lines.append("| " + " | ".join(["---"] * ncols) + " |")
            table_lines.extend(
                "| " + " | ".join(row) + " |" for row in rows[1:])
            rendered.append("\n".join(table_lines))
        else:
            rendered.append(block.value)
    return "\n\n".join(rendered).rstrip() + "\n"


def check_canonical_roundtrip(text, cfg=None, tpl_path=None):
    """Return ``(canonical, recovered)`` after an isolated forward/reverse pass."""
    canonical = canonicalize_markdown(text)
    with tempfile.TemporaryDirectory() as td:
        scratch = Path(td)
        src = scratch / "canonical.md"
        out = scratch / "canonical.docx"
        src.write_text(canonical, encoding="utf8")
        Converter(cfg or DEFAULTS, tpl_path=tpl_path).convert(src, out)
        import docx2md
        recovered = docx2md.convert(out, report_provenance=False)
    return canonical, recovered


def _xml_attr(value):
    return (
        str(value)
        .replace("&", "&amp;")
        .replace('"', "&quot;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _xml_text(value):
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(
        ">", "&gt;")


def _next_numbered_part(names, pattern):
    used = []
    for name in names:
        match = re.fullmatch(pattern, name)
        if match:
            used.append(int(match.group(1)))
    return max(used, default=0) + 1


def _next_relationship_id(root):
    used = set()
    relationship_tag = f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationship"
    for relationship in root.findall(relationship_tag):
        match = re.fullmatch(r"rId([1-9]\d*)", relationship.get("Id", ""))
        if match:
            used.add(int(match.group(1)))
    candidate = 1
    while candidate in used:
        candidate += 1
    return f"rId{candidate}"


def _insert_before_closing(xml_bytes, closing, fragment):
    closing_bytes = closing.encode("utf8")
    offset = xml_bytes.rfind(closing_bytes)
    if offset < 0:
        raise ValueError(f"DOCX part is missing {closing}")
    return (
        xml_bytes[:offset]
        + fragment.encode("utf8")
        + xml_bytes[offset:]
    )


def _new_zip_info(name):
    info = zipfile.ZipInfo(name, date_time=(2000, 1, 1, 0, 0, 0))
    info.compress_type = zipfile.ZIP_DEFLATED
    info.create_system = 0
    info.external_attr = 0o600 << 16
    return info


def embed_source_snapshot(docx_path, original_bytes, canonical_text):
    """Atomically add the canonical Markdown merge base as a custom XML part."""
    path = Path(docx_path)
    if len(original_bytes) > MAX_EMBEDDED_SOURCE_BYTES:
        raise ValueError("original Markdown exceeds the embedded-source limit")
    try:
        original_text = original_bytes.decode("utf8")
    except UnicodeDecodeError as exc:
        raise ValueError("Markdown source must be valid UTF-8") from exc
    expected_canonical = canonicalize_markdown(original_text)
    if (
        canonical_text != expected_canonical
        or canonicalize_markdown(canonical_text) != canonical_text
    ):
        raise ValueError(
            "canonical Markdown must equal canonicalize(original)")
    canonical_bytes = canonical_text.encode("utf8")
    if len(canonical_bytes) > MAX_EMBEDDED_SOURCE_BYTES:
        raise ValueError("canonical Markdown exceeds the embedded-source limit")

    original_sha = hashlib.sha256(original_bytes).hexdigest()
    canonical_sha = hashlib.sha256(canonical_bytes).hexdigest()
    item_uuid = uuid.uuid5(
        uuid.NAMESPACE_URL,
        f"{EMBEDDED_SOURCE_SCHEMA}:{original_sha}:{canonical_sha}",
    )
    item_id = "{" + str(item_uuid).upper() + "}"

    with zipfile.ZipFile(path, "r") as source_zip:
        infos = source_zip.infolist()
        if len(infos) > MAX_DOCX_PARTS:
            raise ValueError("DOCX package exceeds the part-count limit")
        names = [info.filename for info in infos]
        if len(set(names)) != len(names):
            raise ValueError("DOCX package contains duplicate part names")
        if sum(info.file_size for info in infos) > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise ValueError("DOCX package exceeds the uncompressed-size limit")
        for info in infos:
            if (
                info.filename.startswith("/")
                or "\\" in info.filename
                or ".." in Path(info.filename).parts
                or info.flag_bits & 0x1
            ):
                raise ValueError(
                    f"DOCX package contains unsafe part {info.filename!r}")
        parts = {info.filename: source_zip.read(info) for info in infos}

    content_types_name = "[Content_Types].xml"
    document_rels_name = "word/_rels/document.xml.rels"
    if content_types_name not in parts or document_rels_name not in parts:
        raise ValueError("DOCX package lacks required OPC relationship parts")

    content_types_root = ElementTree.fromstring(parts[content_types_name])
    if content_types_root.tag != f"{{{CONTENT_TYPES_NAMESPACE}}}Types":
        raise ValueError("DOCX content-types root uses an unexpected namespace")
    document_rels_root = ElementTree.fromstring(parts[document_rels_name])
    if (
        document_rels_root.tag
        != f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationships"
    ):
        raise ValueError("DOCX document relationships use an unexpected namespace")

    item_number = _next_numbered_part(
        names, r"customXml/item([1-9]\d*)\.xml")
    item_name = f"customXml/item{item_number}.xml"
    props_name = f"customXml/itemProps{item_number}.xml"
    item_rels_name = f"customXml/_rels/item{item_number}.xml.rels"
    if any(name in parts for name in (item_name, props_name, item_rels_name)):
        raise ValueError("DOCX custom XML part allocation collided")

    document_rid = _next_relationship_id(document_rels_root)
    relationship_fragment = (
        f'<Relationship Id="{_xml_attr(document_rid)}" '
        f'Type="{CUSTOM_XML_RELATIONSHIP}" '
        f'Target="../{_xml_attr(item_name)}"/>'
    )
    parts[document_rels_name] = _insert_before_closing(
        parts[document_rels_name], "</Relationships>", relationship_fragment)
    override_fragment = (
        f'<Override PartName="/{_xml_attr(props_name)}" '
        f'ContentType="{CUSTOM_XML_PROPS_CONTENT_TYPE}"/>'
    )
    parts[content_types_name] = _insert_before_closing(
        parts[content_types_name], "</Types>", override_fragment)

    original_b64 = base64.b64encode(original_bytes).decode("ascii")
    canonical_b64 = base64.b64encode(canonical_bytes).decode("ascii")
    parts[item_name] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f'<o:source xmlns:o="{EMBEDDED_SOURCE_NAMESPACE}" '
        f'schema="{EMBEDDED_SOURCE_SCHEMA}" encoding="utf-8" '
        f'originalSha256="{original_sha}" canonicalSha256="{canonical_sha}">\n'
        f'  <o:original encoding="base64">{_xml_text(original_b64)}</o:original>\n'
        f'  <o:canonical encoding="base64">{_xml_text(canonical_b64)}</o:canonical>\n'
        "</o:source>\n"
    ).encode("utf8")
    parts[props_name] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<ds:datastoreItem '
        'xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml" '
        f'ds:itemID="{item_id}">\n'
        "  <ds:schemaRefs>\n"
        f'    <ds:schemaRef ds:uri="{EMBEDDED_SOURCE_NAMESPACE}"/>\n'
        "  </ds:schemaRefs>\n"
        "</ds:datastoreItem>\n"
    ).encode("utf8")
    parts[item_rels_name] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f'<Relationships xmlns="{PACKAGE_RELATIONSHIPS_NAMESPACE}">\n'
        f'  <Relationship Id="rId1" Type="{CUSTOM_XML_PROPS_RELATIONSHIP}" '
        f'Target="itemProps{item_number}.xml"/>\n'
        "</Relationships>\n"
    ).encode("utf8")

    with tempfile.NamedTemporaryFile(
        prefix=f".{path.name}.", suffix=".tmp", dir=path.parent, delete=False
    ) as handle:
        temporary_path = Path(handle.name)
    try:
        with zipfile.ZipFile(temporary_path, "w") as output_zip:
            for info in infos:
                output_zip.writestr(info, parts[info.filename])
            for name in (item_name, props_name, item_rels_name):
                output_zip.writestr(_new_zip_info(name), parts[name])
        with zipfile.ZipFile(temporary_path, "r") as check_zip:
            if check_zip.testzip() is not None:
                raise ValueError("embedded-source DOCX failed ZIP validation")
        try:
            import docx2md
        except ImportError as exc:
            raise ValueError(
                "embedded-source DOCX validator is unavailable") from exc
        try:
            snapshot = docx2md.read_embedded_source(
                temporary_path, required=True)
        except docx2md.RoundtripRefusalError as exc:
            raise ValueError(
                "embedded-source DOCX failed semantic validation: "
                f"{exc}") from exc
        if (
            snapshot.part != item_name
            or snapshot.original_sha256 != original_sha
            or snapshot.canonical_sha256 != canonical_sha
        ):
            raise ValueError(
                "embedded-source DOCX validation returned the wrong item")
        temporary_path.replace(path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()

    return {
        "schema": EMBEDDED_SOURCE_SCHEMA,
        "part": item_name,
        "originalSha256": original_sha,
        "canonicalSha256": canonical_sha,
    }


def stamp_provenance(doc, src, tpl_path, cfg, source_bytes=None):
    """Embed round-trip breadcrumbs in DOCX core properties.

    CONTRACT:C2-PROVENANCE.1.0 — see architecture/CONTRACT-C2-PROVENANCE.1.0.md
    for the field table and the 255-char / hash-truncation rules this
    function must honor.

    Survives Word saves (core props are standard OPC parts). docx2md reads
    `comments` to recover which template produced the formatting and whether
    the source MD has drifted since conversion (hash comparison).
    """
    src_bytes = (
        Path(src).read_bytes()
        if source_bytes is None
        else source_bytes
    )
    tpl_sha = (hashlib.sha256(Path(tpl_path).read_bytes()).hexdigest()[:16]
               if tpl_path and Path(tpl_path).exists() else None)
    prov = {
        "t": f"md2docx/{TOOL_VERSION}",
        "tpl": cfg.get("name") or "built-in defaults",
        "tplsha": tpl_sha,
        "srcsha": hashlib.sha256(src_bytes).hexdigest()[:16],
        "gen": datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%MZ"),
    }
    cp = doc.core_properties                      # each field capped at 255 chars
    cp.comments = json.dumps(prov, separators=(",", ":"))
    cp.subject = str(Path(src).resolve())[-255:]  # full source path, own field
    cp.keywords = "md2docx"
    cp.category = prov["tpl"]


class Converter:
    def __init__(self, cfg, tpl_path=None):
        self.cfg = cfg
        self.tpl_path = tpl_path

    def add_runs(self, p, text, color=None, size=None, base_bold=False):
        cfg = self.cfg
        text = demote_links(text)
        for tok in re.split(r"(`[^`]+`|\*\*.*?\*\*|\*[^*]+\*)", text):
            if not tok:
                continue
            if tok.startswith("`") and tok.endswith("`"):
                r = p.add_run(tok[1:-1])
                r.font.name = cfg["fonts"]["mono"]
                r.font.size = Pt((size or Pt(cfg["base"]["size_pt"])).pt - 0.5)
                r.font.color.rgb = rgb(cfg["code"]["color"])
                shade_run(r, cfg["code"]["fill"])
                if base_bold:
                    r.bold = True
                continue
            if tok.startswith("**") and tok.endswith("**"):
                r = p.add_run(tok[2:-2]); r.bold = True
            elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
                r = p.add_run(tok[1:-1]); r.italic = True
            else:
                r = p.add_run(tok)
                if base_bold:
                    r.bold = True
            r.font.name = cfg["fonts"]["body"]
            if color:
                r.font.color.rgb = color
            if size:
                r.font.size = size

    def convert(self, src, out_path, footer_on=True):
        cfg = self.cfg
        source_path = Path(src)
        original_bytes = source_path.read_bytes()
        try:
            source_text = original_bytes.decode("utf8")
        except UnicodeDecodeError as exc:
            raise MarkdownCanonicalizationError(
                "MD-CANON-INVALID-UTF8", 1,
                "Markdown source must be valid UTF-8") from exc
        canonical_text = canonicalize_markdown(source_text)
        if len(original_bytes) > MAX_EMBEDDED_SOURCE_BYTES:
            raise MarkdownCanonicalizationError(
                "MD-CANON-SOURCE-LIMIT", 1,
                "Markdown source exceeds the embedded-source limit")
        if len(canonical_text.encode("utf8")) > MAX_EMBEDDED_SOURCE_BYTES:
            raise MarkdownCanonicalizationError(
                "MD-CANON-SOURCE-LIMIT", 1,
                "canonical Markdown exceeds the embedded-source limit")
        lines = canonical_text.splitlines()
        doc = Document()

        n = doc.styles["Normal"]
        n.font.name = cfg["fonts"]["body"]
        n.font.size = Pt(cfg["base"]["size_pt"])
        n.font.color.rgb = rgb(cfg["base"]["color"])
        n.paragraph_format.space_after = Pt(cfg["base"]["space_after_pt"])
        for lvl in (1, 2, 3, 4):
            h = cfg["headings"][f"h{lvl}"]
            s = doc.styles[f"Heading {lvl}"]
            s.font.name = cfg["fonts"]["body"]
            s.font.size = Pt(h["size_pt"])
            s.font.color.rgb = rgb(h["color"])
            s.font.bold = cfg["headings"]["bold"]
            s.paragraph_format.space_before = Pt(cfg["headings"]["space_before_pt"])
            s.paragraph_format.space_after = Pt(cfg["headings"]["space_after_pt"])

        banner = None
        body_start = 0
        if cfg["cui_banner"]["detect"]:
            for idx, ln in enumerate(lines):
                if not ln.strip():
                    continue
                m = re.fullmatch(r"\*\*(CUI[^*]*)\*\*", ln.strip())
                if m:
                    banner = m.group(1)
                    body_start = idx + 1
                break

        i = body_start
        while i < len(lines):
            ln = lines[i].rstrip()
            stripped = ln.strip()
            if not stripped or stripped.startswith("<!--") or re.fullmatch(r"-{3,}", stripped):
                i += 1; continue

            if stripped.startswith("```"):
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Inches(0.2)
                    shade_para(p, cfg["code"]["fill"])
                    r = p.add_run(lines[i] if lines[i].strip() else " ")
                    r.font.name = cfg["fonts"]["mono"]
                    r.font.size = Pt(cfg["code"]["block_size_pt"])
                    r.font.color.rgb = rgb(cfg["code"]["color"])
                    i += 1
                i += 1
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue

            if stripped.startswith("|"):
                rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    if not all(re.fullmatch(r":?-+:?", c) for c in cells):
                        rows.append(cells)
                    i += 1
                if rows:
                    tc = cfg["table"]
                    ncols = max(len(r) for r in rows)
                    t = doc.add_table(rows=0, cols=ncols)
                    t.style = tc["style"]
                    t.autofit = True
                    for ri, row in enumerate(rows):
                        cells = t.add_row().cells
                        for ci in range(ncols):
                            txt = row[ci] if ci < len(row) else ""
                            p = cells[ci].paragraphs[0]
                            if ri == 0:
                                self.add_runs(p, re.sub(r"\*\*", "", txt),
                                              color=rgb(tc["header_color"]),
                                              size=Pt(tc["cell_size_pt"]),
                                              base_bold=tc["header_bold"])
                                shade_cell(cells[ci], tc["header_fill"])
                            else:
                                self.add_runs(p, txt, size=Pt(tc["cell_size_pt"]))
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue

            if stripped.startswith(">"):
                buf = []
                while i < len(lines) and lines[i].strip().startswith(">"):
                    buf.append(lines[i].strip().lstrip(">").strip())
                    i += 1
                bq = cfg["blockquote"]
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(bq["indent_in"])
                left_border(p, bq["border_color"], bq["border_size_eighths"])
                self.add_runs(p, " ".join(buf), color=rgb(bq["color"]), size=Pt(bq["size_pt"]))
                continue

            m = re.match(r"^(#{1,4}) ", stripped)
            if m:
                doc.add_heading(stripped[len(m.group(1)) + 1:], level=len(m.group(1)))
            elif re.match(r"^[-*] ", stripped):
                p = doc.add_paragraph(style="List Bullet")
                self.add_runs(p, stripped[2:])
            elif re.match(r"^\d+\. ", stripped):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                self.add_runs(p, stripped)
            else:
                buf = [stripped]
                while i + 1 < len(lines):
                    nxt = lines[i + 1].strip()
                    if (not nxt or nxt.startswith(("#", "|", ">", "```", "<!--"))
                            or re.match(r"^[-*] ", nxt) or re.match(r"^\d+\. ", nxt)
                            or re.fullmatch(r"-{3,}", nxt)):
                        break
                    buf.append(nxt); i += 1
                p = doc.add_paragraph()
                self.add_runs(p, " ".join(buf))
            i += 1

        sec = doc.sections[0]
        if banner:
            cb = cfg["cui_banner"]
            hp = sec.header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = hp.add_run(banner)
            hr.font.name = cfg["fonts"]["body"]
            hr.font.size = Pt(cb["size_pt"])
            hr.bold = cb["bold"]
            hr.font.color.rgb = rgb(cb["color"])
        ftext = banner if banner else (cfg["footer"]["text"] if footer_on else "")
        if ftext:
            fc = cfg["footer"]
            fp = sec.footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = fp.add_run(ftext)
            fr.font.name = cfg["fonts"]["body"]
            fr.font.size = Pt(fc["size_pt"])
            fr.bold = fc["bold"]
            if not banner:
                fr.font.color.rgb = rgb(fc["color"])

        stamp_provenance(
            doc, src, self.tpl_path, self.cfg,
            source_bytes=original_bytes)
        destination = Path(out_path)
        with tempfile.NamedTemporaryFile(
            prefix=f".{destination.name}.",
            suffix=".docx.tmp",
            dir=destination.parent,
            delete=False,
        ) as handle:
            staged_path = Path(handle.name)
        try:
            doc.save(staged_path)
            embed_source_snapshot(
                staged_path, original_bytes, canonical_text)
            staged_path.replace(destination)
        finally:
            if staged_path.exists():
                staged_path.unlink()
        return destination


def resolve_template(explicit, base_dir=None):
    """Template resolution order — CONTRACT:C1-THEME-SCHEMA.1.0.

    1. --template flag, if given.
    2. `md2docx-template.json` next to this script (a local override; not
       shipped in this repo).
    3. `themes/neutral.json` next to this script (the built-in-equivalent
       shipped theme, kept as a JSON file so it's documented and diffable).
    4. The DEFAULTS dict in this module, hard-coded.

    `base_dir` overrides "next to this script" (defaults to this module's
    directory); tests use it to exercise the resolution order against a
    scratch directory instead of the real repo.

    Returns (cfg, tpl_path, message) — tpl_path is None when no file was
    used (pure DEFAULTS), so provenance stamping records "built-in defaults".
    """
    here = Path(base_dir) if base_dir else Path(__file__).parent
    if explicit:
        tpl_path = Path(explicit)
        if not tpl_path.exists():
            sys.exit(f"template not found: {explicit}")
        cfg = deep_merge(DEFAULTS, json.loads(tpl_path.read_text(encoding="utf8")))
        return cfg, tpl_path, f"template: {tpl_path}" + (f" ({cfg.get('name')})" if cfg.get("name") else "")

    local_override = here / "md2docx-template.json"
    if local_override.exists():
        cfg = deep_merge(DEFAULTS, json.loads(local_override.read_text(encoding="utf8")))
        return cfg, local_override, f"template: {local_override}" + (f" ({cfg.get('name')})" if cfg.get("name") else "")

    shipped_neutral = here / "themes" / "neutral.json"
    if shipped_neutral.exists():
        cfg = deep_merge(DEFAULTS, json.loads(shipped_neutral.read_text(encoding="utf8")))
        return cfg, shipped_neutral, f"template: {shipped_neutral} ({cfg.get('name', 'Neutral')})"

    return DEFAULTS, None, "template: built-in defaults (no md2docx-template.json or themes/neutral.json found)"


def _read_markdown_utf8(path):
    try:
        return Path(path).read_bytes().decode("utf8")
    except UnicodeDecodeError as exc:
        raise MarkdownCanonicalizationError(
            "MD-CANON-INVALID-UTF8", 1,
            "Markdown source must be valid UTF-8") from exc


def _atomic_write_text(path, text):
    destination = Path(path)
    payload = text.encode("utf8")
    with tempfile.NamedTemporaryFile(
        prefix=f".{destination.name}.",
        suffix=".tmp",
        dir=destination.parent,
        delete=False,
    ) as handle:
        staged_path = Path(handle.name)
        handle.write(payload)
    try:
        staged_path.replace(destination)
    finally:
        if staged_path.exists():
            staged_path.unlink()


def _paths_alias(first, second):
    first_path = Path(first)
    second_path = Path(second)
    first_resolved = first_path.resolve()
    second_resolved = second_path.resolve()
    if first_resolved == second_resolved:
        return True
    first_key = unicodedata.normalize(
        "NFC", str(first_resolved)).casefold()
    second_key = unicodedata.normalize(
        "NFC", str(second_resolved)).casefold()
    if first_key == second_key:
        return True
    try:
        return first_path.exists() and second_path.exists() and (
            first_path.samefile(second_path)
        )
    except OSError:
        return False


def main():
    ap = argparse.ArgumentParser(description="Markdown → styled DOCX, themed by a JSON template.")
    ap.add_argument("files", nargs="+", help="Markdown source files")
    ap.add_argument("-t", "--template", help="Template JSON (default: md2docx-template.json next to this script, then themes/neutral.json)")
    ap.add_argument("-o", "--out", help="Output directory, or output .docx path for a single input")
    ap.add_argument("--no-footer", action="store_true", help="Suppress footer text")
    mode = ap.add_mutually_exclusive_group()
    mode.add_argument(
        "--normalize", action="store_true",
        help="print canonical supported-profile Markdown; write no DOCX")
    mode.add_argument(
        "--check", action="store_true",
        help="require canonical source and an exact canonical DOCX round trip")
    args = ap.parse_args()

    if args.normalize:
        if len(args.files) != 1:
            ap.error("--normalize accepts exactly one Markdown source")
        if args.template or args.no_footer:
            ap.error("--normalize does not use --template or --no-footer")
        src = Path(args.files[0])
        if not src.exists():
            sys.exit(f"not found: {src}")
        try:
            canonical = canonicalize_markdown(_read_markdown_utf8(src))
        except MarkdownCanonicalizationError as exc:
            sys.exit(str(exc))
        except OSError as exc:
            sys.exit(f"MD-CANON-IO: {exc}")
        if args.out:
            if _paths_alias(src, args.out):
                sys.exit(
                    "MD-CANON-PATH-ALIAS: --out must not overwrite the source")
            try:
                _atomic_write_text(args.out, canonical)
            except OSError as exc:
                sys.exit(f"MD-CANON-IO: {exc}")
            print(f"-> {args.out}")
        else:
            sys.stdout.write(canonical)
        return

    cfg, tpl_path, message = resolve_template(args.template)
    if not args.check:
        print(message)

    if args.check:
        if args.out or args.no_footer:
            ap.error("--check does not write output or use --no-footer")
        failed = False
        for name in args.files:
            src = Path(name)
            if not src.exists():
                print(f"{src}: not found", file=sys.stderr)
                failed = True
                continue
            try:
                source = _read_markdown_utf8(src)
                canonical, recovered = check_canonical_roundtrip(
                    source, cfg=cfg, tpl_path=tpl_path)
            except MarkdownCanonicalizationError as exc:
                print(f"{src}: {exc}", file=sys.stderr)
                failed = True
                continue
            except OSError as exc:
                print(f"{src}: MD-CANON-IO: {exc}", file=sys.stderr)
                failed = True
                continue
            if source != canonical:
                print(
                    f"{src}: not canonical; run md2docx.py --normalize",
                    file=sys.stderr)
                failed = True
            elif recovered != canonical:
                print(
                    f"{src}: canonical DOCX round-trip mismatch",
                    file=sys.stderr)
                failed = True
            else:
                print(f"{src}: canonical round-trip OK")
        if failed:
            raise SystemExit(1)
        return

    conv = Converter(cfg, tpl_path=tpl_path)
    single_file_out = args.out and args.out.lower().endswith(".docx")
    if single_file_out and len(args.files) > 1:
        sys.exit("-o <file.docx> only valid with a single input; use -o <dir> for many")

    planned = []
    for src in args.files:
        src_p = Path(src)
        if not src_p.exists():
            sys.exit(f"not found: {src}")
        if single_file_out:
            out = Path(args.out)
        elif args.out:
            Path(args.out).mkdir(parents=True, exist_ok=True)
            out = Path(args.out) / (src_p.stem + ".docx")
        else:
            out = src_p.with_suffix(".docx")
        if _paths_alias(src_p, out):
            sys.exit(
                f"MD-CANON-PATH-ALIAS: output must not overwrite {src_p}")
        if any(_paths_alias(out, existing_out)
               for _, existing_out in planned):
            sys.exit(
                f"MD-CANON-PATH-ALIAS: multiple inputs resolve to {out}")
        planned.append((src_p, out))

    protected_inputs = [source for source, _ in planned]
    if tpl_path is not None:
        protected_inputs.append(Path(tpl_path))
    for _, out in planned:
        for protected in protected_inputs:
            if _paths_alias(out, protected):
                sys.exit(
                    "MD-CANON-PATH-ALIAS: output must not overwrite "
                    f"{protected}")

    for src_p, out in planned:
        try:
            converted = conv.convert(
                src_p, out, footer_on=not args.no_footer)
        except MarkdownCanonicalizationError as exc:
            sys.exit(str(exc))
        except OSError as exc:
            sys.exit(f"MD-CANON-IO: {exc}")
        except (
            ValueError,
            zipfile.BadZipFile,
            ElementTree.ParseError,
        ) as exc:
            sys.exit(f"MD-CANON-PACKAGE: {exc}")
        print("→", converted)


if __name__ == "__main__":
    main()
