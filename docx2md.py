#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""docx2md — styled DOCX -> canonical Markdown (reverse of md2docx.py).

CONTRACT:C2-PROVENANCE.2.0
CONTRACT:C3-ROUNDTRIP.1.2

Style-driven-inversion converter: it walks the document body in order and
inverts the forward converter's construct->style mapping:

  Heading 1-4 style        -> #, ##, ###, ####
  List Bullet style        -> "- "
  Normal para "N. ..."     -> kept literally (forward writes literal numbers)
  table (row 0 shaded)     -> pipe table with a separator row
  left-bordered indent     -> "> " blockquote
  paragraph-shaded + mono  -> ``` fenced code block
  inline: mono/shaded run  -> `code`,  bold run -> **, italic run -> *

Adjacent runs with identical formatting are merged before emitting so words
are not fragmented. Output is canonical MD: "-" bullets, "**" bold, one blank
line between blocks, no trailing whitespace. C3 1.2 adds exact or explicitly
normalized canonical equality, explicit lossiness refusals, an embedded C2 2.0
merge base, semantic controlled-style evidence, fidelity reporting, and opt-in
three-way merge. Recovering link markup demoted by the forward tool remains
outside this profile.

Usage:
  docx2md.py file.docx [-o out.md]      # default output next to input
Requires: python-docx
"""
import argparse
import base64
import hashlib
import io
import json
import posixpath
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

MONO_FONTS = {"Consolas", "Courier New", "Courier", "Menlo", "Monaco",
              "DejaVu Sans Mono", "Cascadia Code", "Cascadia Mono"}
EMBEDDED_SOURCE_SCHEMA = "office180-md-source/0.1"
EMBEDDED_SOURCE_NAMESPACE = "urn:office180:md-source:0.1"
CUSTOM_XML_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml"
)
CUSTOM_XML_PROPS_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXmlProps"
)
CUSTOM_XML_PROPS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.customXmlProperties+xml"
)
CUSTOM_XML_DATASTORE_NAMESPACE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/customXml"
)
PACKAGE_RELATIONSHIPS_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)
CONTENT_TYPES_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/content-types"
)
CORE_PROPERTIES_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
)
DC_NAMESPACE = "http://purl.org/dc/elements/1.1/"
MAX_EMBEDDED_SOURCE_BYTES = 8 * 1024 * 1024
MAX_DOCX_PARTS = 10_000
MAX_DOCX_UNCOMPRESSED_BYTES = 256 * 1024 * 1024
SHA256_PATTERN = re.compile(r"^[a-f0-9]{64}$")
VISUAL_STYLE_PROJECTION_SCHEMA = (
    "office180-docx-visual-style-projection/0.2"
)
ROUNDTRIP_REPORT_SCHEMA = "office180-docx-roundtrip-report/0.2"
VISUAL_STYLE_DRIFT_CODE = "DOCX-ROUNDTRIP-VISUAL-STYLE-DRIFT"
VISUAL_STYLE_NATIVE_NORMALIZATION_CODE = (
    "DOCX-ROUNDTRIP-VISUAL-STYLE-NATIVE-NORMALIZATION"
)
EXPLICIT_FONT_ATTRIBUTES = ("ascii", "hAnsi", "eastAsia", "cs")
THEME_FONT_ATTRIBUTES = (
    "asciiTheme",
    "hAnsiTheme",
    "eastAsiaTheme",
    "cstheme",
)
CONTROLLED_STYLE_IDS = (
    "Normal",
    "Heading1",
    "Heading1Char",
    "Heading2",
    "Heading2Char",
    "Heading3",
    "Heading3Char",
    "Heading4",
    "Heading4Char",
)
SUPPORTED_BODY_STYLES = {
    "Normal",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "List Bullet",
}


class RoundtripRefusalError(ValueError):
    """A DOCX construct that cannot be inverted without silent fidelity loss."""

    def __init__(self, code, message):
        self.code = code
        self.message = message
        super().__init__(f"{code}: {message}")


@dataclass(frozen=True)
class EmbeddedSourceSnapshot:
    schema: str
    part: str
    original_text: str
    canonical_text: str
    original_sha256: str
    canonical_sha256: str


@dataclass(frozen=True)
class MergeResult:
    text: str
    conflicts: bool
    base_sha256: str
    current_sha256: str
    edited_sha256: str


# ---------- low-level XML probes ----------

def _pPr_child(el, tag):
    pPr = el.find(qn("w:pPr"))
    return None if pPr is None else pPr.find(qn(tag))


def para_shd_fill(p):
    """Paragraph-level shading fill (marks a code-block line), or None."""
    shd = _pPr_child(p._p, "w:shd")
    return shd.get(qn("w:fill")) if shd is not None else None


def para_has_left_border(p):
    """True if the paragraph carries a left border (marks a blockquote)."""
    pBdr = _pPr_child(p._p, "w:pBdr")
    return pBdr is not None and pBdr.find(qn("w:left")) is not None


def run_shd_fill(r):
    rPr = r._r.find(qn("w:rPr"))
    if rPr is None:
        return None
    shd = rPr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def run_is_code(r):
    return run_shd_fill(r) is not None or (r.font.name in MONO_FONTS)


def _run_complex_script_property(r, name):
    r_pr = r._r.find(qn("w:rPr"))
    if r_pr is None:
        return False
    element = r_pr.find(qn(f"w:{name}"))
    if element is None:
        return False
    return element.get(qn("w:val"), "1").lower() not in (
        "0", "false", "off")


# ---------- inline run inversion ----------

def _wrap(text, marker):
    """Wrap non-space core in marker, keeping surrounding whitespace outside."""
    core = text.strip()
    if not core:
        return text
    lead = text[:len(text) - len(text.lstrip())]
    trail = text[len(text.rstrip()):]
    return f"{lead}{marker}{core}{marker}{trail}"


def _append_inline_segment(segments, fmt, text):
    if not text:
        return
    if segments and segments[-1][0] == fmt:
        segments[-1] = (fmt, segments[-1][1] + text)
    else:
        segments.append((fmt, text))


def _word_inline_segments(runs, suppress_bold=False):
    segs = []  # (fmt, text) where fmt = (code, bold, italic)
    for r in runs:
        if not r.text:
            continue
        code = run_is_code(r)
        bold = (
            bool(r.bold) or _run_complex_script_property(r, "bCs")
        ) and not suppress_bold
        italic = (
            bool(r.italic) or _run_complex_script_property(r, "iCs")
        )
        if code:  # code is exclusive; bold/italic markers do not apply inside
            bold = italic = False
        fmt = (code, bold, italic)
        if bold and italic:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-INLINE",
                "combined bold+italic runs are outside the canonical profile")
        _append_inline_segment(segs, fmt, r.text)
    return segs


def _semantic_word_segments(run_segments):
    semantic = []
    for fmt, text in run_segments:
        code, bold, italic = fmt
        if code or not (bold or italic):
            _append_inline_segment(semantic, fmt, text)
            continue
        if not text.strip():
            _append_inline_segment(
                semantic, (False, False, False), text)
            continue
        lead = text[:len(text) - len(text.lstrip())]
        trail = text[len(text.rstrip()):]
        core = text.strip()
        _append_inline_segment(semantic, (False, False, False), lead)
        _append_inline_segment(semantic, fmt, core)
        _append_inline_segment(semantic, (False, False, False), trail)
    return semantic


INLINE_TOKEN_RE = re.compile(r"(`[^`]+`|\*\*.*?\*\*|\*[^*]+\*)")


def _markdown_inline_segments(text):
    semantic = []
    for token in INLINE_TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            fmt = (True, False, False)
            value = token[1:-1]
        elif token.startswith("**") and token.endswith("**"):
            fmt = (False, True, False)
            value = token[2:-2]
        elif token.startswith("*") and token.endswith("*"):
            fmt = (False, False, True)
            value = token[1:-1]
        else:
            fmt = (False, False, False)
            value = token
        _append_inline_segment(semantic, fmt, value)
    return semantic


def runs_to_md(runs, suppress_bold=False):
    """Invert runs only when forward inline parsing recreates their semantics."""
    segs = _word_inline_segments(runs, suppress_bold=suppress_bold)
    out = []
    for (code, bold, italic), text in segs:
        if code:
            out.append("`" + text + "`")
            continue
        if bold:
            out.append(_wrap(text, "**"))
        elif italic:
            out.append(_wrap(text, "*"))
        else:
            out.append(text)
    rendered = "".join(out)
    if (
        _markdown_inline_segments(rendered)
        != _semantic_word_segments(segs)
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-INLINE",
            "Word run semantics cannot be represented by canonical Markdown")
    return rendered


# ---------- cell / table helpers ----------

def cell_md(cell, header=False):
    if len(cell.paragraphs) != 1:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-TABLE",
            "table cells must contain exactly one supported paragraph")
    for paragraph in cell.paragraphs:
        style = paragraph.style.name if paragraph.style else "Normal"
        if style != "Normal":
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-STYLE",
                f"table-cell paragraph style {style!r} is unsupported")
    parts = [
        runs_to_md(p.runs, suppress_bold=header)
        for p in cell.paragraphs
    ]
    if any(part != part.strip() for part in parts):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            "table-cell text has unsupported outer whitespace")
    txt = " ".join(part for part in parts if part)
    return txt.replace("|", "\\|")


def table_to_md(tbl):
    rows = tbl.rows
    if not rows:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-TABLE",
            "empty Word tables are outside the supported profile")
    ncols = len(tbl.columns)
    if ncols == 0:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-TABLE",
            "Word tables must have at least one grid column")
    lines = []
    header = ["" for _ in range(ncols)]
    for ci, cell in enumerate(rows[0].cells[:ncols]):
        header[ci] = cell_md(cell, header=True)
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * ncols) + " |")
    for row in rows[1:]:
        vals = ["" for _ in range(ncols)]
        for ci, cell in enumerate(row.cells[:ncols]):
            vals[ci] = cell_md(cell)
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


# ---------- provenance ----------

def print_provenance(doc):
    """Recover the C2 1.0 core stamp retained by CONTRACT:C2-PROVENANCE.2.0."""
    cp = doc.core_properties
    print("--- provenance stamp ---", file=sys.stderr)
    if cp.comments:
        try:
            prov = json.loads(cp.comments)
            for k, v in prov.items():
                print(f"  {k}: {v}", file=sys.stderr)
        except (ValueError, TypeError):
            print(f"  comments: {cp.comments}", file=sys.stderr)
    else:
        print("  (no comments JSON found)", file=sys.stderr)
    if cp.subject:
        print(f"  subject (source path): {cp.subject}", file=sys.stderr)
    if cp.category:
        print(f"  category (template):   {cp.category}", file=sys.stderr)
    if cp.keywords:
        print(f"  keywords: {cp.keywords}", file=sys.stderr)
    print("------------------------", file=sys.stderr)


# ---------- embedded merge base ----------

def _read_safe_docx_parts_bytes(docx_bytes):
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as package:
            infos = package.infolist()
            if len(infos) > MAX_DOCX_PARTS:
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-PACKAGE",
                    "DOCX package exceeds the part-count limit")
            names = [info.filename for info in infos]
            if len(set(names)) != len(names):
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-PACKAGE",
                    "DOCX package contains duplicate part names")
            if sum(info.file_size for info in infos) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-PACKAGE",
                    "DOCX package exceeds the uncompressed-size limit")
            for info in infos:
                if (
                    info.filename.startswith("/")
                    or "\\" in info.filename
                    or ".." in Path(info.filename).parts
                    or info.flag_bits & 0x1
                ):
                    raise RoundtripRefusalError(
                        "DOCX-ROUNDTRIP-PACKAGE",
                        f"DOCX package contains unsafe part {info.filename!r}")
            return {info.filename: package.read(info) for info in infos}
    except (zipfile.BadZipFile, NotImplementedError, RuntimeError) as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PACKAGE",
            "input is not a valid ZIP/OPC package") from exc


def _read_safe_docx_parts(docx_path):
    return _read_safe_docx_parts_bytes(Path(docx_path).read_bytes())


def _parse_xml_part(parts, name, code="DOCX-ROUNDTRIP-PROVENANCE"):
    payload = parts.get(name)
    if payload is None:
        raise RoundtripRefusalError(code, f"required part {name!r} is missing")
    try:
        return ElementTree.fromstring(payload)
    except ElementTree.ParseError as exc:
        raise RoundtripRefusalError(
            code, f"part {name!r} is not well-formed XML") from exc


def _resolve_relationship_target(
    source_part,
    target,
    code="DOCX-ROUNDTRIP-PROVENANCE",
    description="custom XML",
):
    if (
        not target
        or target.startswith("/")
        or "\\" in target
        or ":" in target
        or "?" in target
        or "#" in target
    ):
        raise RoundtripRefusalError(
            code,
            f"unsafe {description} relationship target {target!r}")
    directory = posixpath.dirname(source_part)
    resolved = posixpath.normpath(posixpath.join(directory, target))
    if resolved == ".." or resolved.startswith("../"):
        raise RoundtripRefusalError(
            code,
            f"{description} relationship escapes the package: {target!r}")
    return resolved


def _decode_embedded_text(root, local_name, expected_sha):
    elements = root.findall(
        f"{{{EMBEDDED_SOURCE_NAMESPACE}}}{local_name}")
    if (
        len(elements) != 1
        or elements[0].get("encoding") != "base64"
        or set(elements[0].attrib) != {"encoding"}
        or len(elements[0]) != 0
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source must be one base64 text element")
    raw_encoded = elements[0].text or ""
    max_encoded = 4 * ((MAX_EMBEDDED_SOURCE_BYTES + 2) // 3)
    if len(raw_encoded) > max_encoded + 65_536:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source exceeds the encoded-size limit")
    encoded = "".join(raw_encoded.split())
    if len(encoded) > max_encoded:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source exceeds the encoded-size limit")
    try:
        payload = base64.b64decode(encoded, validate=True)
    except ValueError as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source is not valid base64") from exc
    if len(payload) > MAX_EMBEDDED_SOURCE_BYTES:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source exceeds the size limit")
    actual_sha = hashlib.sha256(payload).hexdigest()
    if actual_sha != expected_sha:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source hash does not match its manifest")
    try:
        return payload.decode("utf8")
    except UnicodeDecodeError as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded {local_name} source is not valid UTF-8") from exc


def _relationship_elements(
    root,
    relationship_tag,
    code="DOCX-ROUNDTRIP-PROVENANCE",
):
    relationships = root.findall(relationship_tag)
    ids = [relationship.get("Id", "") for relationship in relationships]
    if any(not relationship_id for relationship_id in ids):
        raise RoundtripRefusalError(
            code,
            "OPC relationships must have non-empty IDs")
    if len(set(ids)) != len(ids):
        raise RoundtripRefusalError(
            code,
            "OPC relationship IDs must be unique within a part")
    return relationships


def _require_internal_relationship(
    relationship,
    description,
    code="DOCX-ROUNDTRIP-PROVENANCE",
):
    target_mode = relationship.get("TargetMode")
    if target_mode not in (None, "Internal"):
        raise RoundtripRefusalError(
            code,
            f"{description} relationship must be internal")


def _validate_core_source_hash(parts, original_sha):
    core = _parse_xml_part(parts, "docProps/core.xml")
    if core.tag != f"{{{CORE_PROPERTIES_NAMESPACE}}}coreProperties":
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "core properties use an unexpected namespace")
    description = core.find(f"{{{DC_NAMESPACE}}}description")
    if description is None or not description.text:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source lacks its C2 core-property breadcrumb")
    try:
        provenance = json.loads(description.text)
    except (ValueError, TypeError) as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "C2 core-property provenance is not valid JSON") from exc
    srcsha = provenance.get("srcsha") if isinstance(provenance, dict) else None
    if (
        not isinstance(srcsha, str)
        or re.fullmatch(r"[a-f0-9]{16}", srcsha) is None
        or srcsha != original_sha[:16]
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "core-property source hash contradicts the embedded source")


def _validate_embedded_canonical(original_text, canonical_text):
    try:
        import md2docx
        expected = md2docx.canonicalize_markdown(original_text)
        idempotent = md2docx.canonicalize_markdown(canonical_text)
    except (ImportError, ValueError) as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded canonical source is outside the C3 profile: {exc}") from exc
    if expected != canonical_text or idempotent != canonical_text:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded canonical source is not canonicalize(original)")


def _read_embedded_source_parts(parts, required=False):
    rels_name = "word/_rels/document.xml.rels"
    rels_root = _parse_xml_part(parts, rels_name)
    if (
        rels_root.tag
        != f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationships"
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "document relationships use an unexpected namespace")
    relationship_tag = f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationship"
    custom_parts = []
    for relationship in _relationship_elements(
            rels_root, relationship_tag):
        if relationship.get("Type") != CUSTOM_XML_RELATIONSHIP:
            continue
        _require_internal_relationship(
            relationship, "document custom XML")
        target = relationship.get("Target", "")
        part_name = _resolve_relationship_target("word/document.xml", target)
        root = _parse_xml_part(parts, part_name)
        if root.tag == f"{{{EMBEDDED_SOURCE_NAMESPACE}}}source":
            custom_parts.append((part_name, root))

    if not custom_parts:
        if required:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-NO-MERGE-BASE",
                "DOCX has no office180 embedded Markdown merge base")
        return None
    if len(custom_parts) != 1:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "DOCX contains more than one office180 embedded source")

    part_name, root = custom_parts[0]
    if (
        root.get("schema") != EMBEDDED_SOURCE_SCHEMA
        or root.get("encoding") != "utf-8"
        or set(root.attrib) != {
            "schema",
            "encoding",
            "originalSha256",
            "canonicalSha256",
        }
        or [child.tag for child in root] != [
            f"{{{EMBEDDED_SOURCE_NAMESPACE}}}original",
            f"{{{EMBEDDED_SOURCE_NAMESPACE}}}canonical",
        ]
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source schema or encoding is unsupported")
    original_sha = root.get("originalSha256", "")
    canonical_sha = root.get("canonicalSha256", "")
    if (
        SHA256_PATTERN.fullmatch(original_sha) is None
        or SHA256_PATTERN.fullmatch(canonical_sha) is None
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source digests must be lowercase full SHA-256 values")

    match = re.fullmatch(r"customXml/item([1-9]\d*)\.xml", part_name)
    if match is None:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            f"embedded source uses unexpected part name {part_name!r}")
    item_number = match.group(1)
    props_name = f"customXml/itemProps{item_number}.xml"
    item_rels_name = f"customXml/_rels/item{item_number}.xml.rels"
    props_root = _parse_xml_part(parts, props_name)
    item_id = props_root.get(
        f"{{{CUSTOM_XML_DATASTORE_NAMESPACE}}}itemID", "")
    if (
        props_root.tag
        != f"{{{CUSTOM_XML_DATASTORE_NAMESPACE}}}datastoreItem"
        or re.fullmatch(
            r"\{[A-F0-9]{8}-[A-F0-9]{4}-[A-F0-9]{4}-"
            r"[A-F0-9]{4}-[A-F0-9]{12}\}",
            item_id,
        )
        is None
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source properties have invalid datastore identity")
    expected_uuid = uuid.uuid5(
        uuid.NAMESPACE_URL,
        f"{EMBEDDED_SOURCE_SCHEMA}:{original_sha}:{canonical_sha}",
    )
    if item_id != "{" + str(expected_uuid).upper() + "}":
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source datastore identity does not match its digests")
    schema_refs = props_root.findall(
        f".//{{{CUSTOM_XML_DATASTORE_NAMESPACE}}}schemaRef")
    if (
        len(schema_refs) != 1
        or schema_refs[0].get(
            f"{{{CUSTOM_XML_DATASTORE_NAMESPACE}}}uri")
        != EMBEDDED_SOURCE_NAMESPACE
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source properties do not declare the source namespace")
    item_rels = _parse_xml_part(parts, item_rels_name)
    if (
        item_rels.tag
        != f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationships"
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source properties relationships use an unexpected namespace")
    props_relationships = [
        relationship
        for relationship in _relationship_elements(
            item_rels, relationship_tag)
        if relationship.get("Type") == CUSTOM_XML_PROPS_RELATIONSHIP
    ]
    if len(props_relationships) != 1:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source must have one custom XML properties relationship")
    _require_internal_relationship(
        props_relationships[0], "custom XML properties")
    resolved_props = _resolve_relationship_target(
        part_name, props_relationships[0].get("Target", ""))
    if resolved_props != props_name:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source properties relationship targets the wrong part")

    content_types = _parse_xml_part(parts, "[Content_Types].xml")
    if content_types.tag != f"{{{CONTENT_TYPES_NAMESPACE}}}Types":
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "content types use an unexpected namespace")
    override_tag = f"{{{CONTENT_TYPES_NAMESPACE}}}Override"
    overrides = content_types.findall(override_tag)
    override_names = [
        override.get("PartName", "") for override in overrides
    ]
    if any(not name for name in override_names) or (
            len(set(override_names)) != len(override_names)):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "content-type overrides must have unique non-empty part names")
    props_overrides = [
        override
        for override in overrides
        if (
            override.get("PartName") == f"/{props_name}"
            and override.get("ContentType") == CUSTOM_XML_PROPS_CONTENT_TYPE
        )
    ]
    if len(props_overrides) != 1:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source properties lack their content-type declaration")
    default_tag = f"{{{CONTENT_TYPES_NAMESPACE}}}Default"
    xml_defaults = [
        default
        for default in content_types.findall(default_tag)
        if default.get("Extension", "").lower() == "xml"
    ]
    item_overrides = [
        override for override in overrides
        if override.get("PartName") == f"/{part_name}"
    ]
    item_has_xml_type = (
        len(item_overrides) == 1
        and item_overrides[0].get("ContentType") == "application/xml"
    ) or (
        not item_overrides
        and len(xml_defaults) == 1
        and xml_defaults[0].get("ContentType") == "application/xml"
    )
    if not item_has_xml_type:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PROVENANCE",
            "embedded source lacks an unambiguous XML content type")

    original_text = _decode_embedded_text(root, "original", original_sha)
    canonical_text = _decode_embedded_text(root, "canonical", canonical_sha)
    _validate_core_source_hash(parts, original_sha)
    _validate_embedded_canonical(original_text, canonical_text)
    return EmbeddedSourceSnapshot(
        schema=EMBEDDED_SOURCE_SCHEMA,
        part=part_name,
        original_text=original_text,
        canonical_text=canonical_text,
        original_sha256=original_sha,
        canonical_sha256=canonical_sha,
    )


def read_embedded_source(docx_path, required=False):
    """Read a package-internally-consistent custom XML Markdown snapshot."""
    return _read_embedded_source_parts(
        _read_safe_docx_parts(docx_path), required=required)


# ---------- controlled visual-style projection ----------

def _on_off_projection(r_pr, name):
    if r_pr is None:
        return None
    element = r_pr.find(qn(f"w:{name}"))
    if element is None:
        return None
    return element.get(qn("w:val"), "1").lower() not in (
        "0", "false", "off")


def _style_reference(style, name):
    element = style.find(qn(f"w:{name}"))
    return None if element is None else element.get(qn("w:val"))


def _project_controlled_style(style):
    r_pr = style.find(qn("w:rPr"))
    r_fonts = None if r_pr is None else r_pr.find(qn("w:rFonts"))
    color = None if r_pr is None else r_pr.find(qn("w:color"))
    size = None if r_pr is None else r_pr.find(qn("w:sz"))
    size_cs = None if r_pr is None else r_pr.find(qn("w:szCs"))
    return {
        "styleId": style.get(qn("w:styleId")),
        "styleType": style.get(qn("w:type")),
        "basedOn": _style_reference(style, "basedOn"),
        "link": _style_reference(style, "link"),
        "explicitFonts": {
            name: (
                None
                if r_fonts is None
                else r_fonts.get(qn(f"w:{name}"))
            )
            for name in EXPLICIT_FONT_ATTRIBUTES
        },
        "themeFonts": {
            name: (
                None
                if r_fonts is None
                else r_fonts.get(qn(f"w:{name}"))
            )
            for name in THEME_FONT_ATTRIBUTES
        },
        "sizeHalfPoints": {
            "latin": None if size is None else size.get(qn("w:val")),
            "complex": (
                None if size_cs is None else size_cs.get(qn("w:val"))
            ),
        },
        "bold": {
            "latin": _on_off_projection(r_pr, "b"),
            "complex": _on_off_projection(r_pr, "bCs"),
        },
        "italic": {
            "latin": _on_off_projection(r_pr, "i"),
            "complex": _on_off_projection(r_pr, "iCs"),
        },
        "color": None if color is None else color.get(qn("w:val")),
    }


def _theme_font_context(parts):
    theme_name = "word/theme/theme1.xml"
    if theme_name not in parts:
        return {"majorLatin": None, "minorLatin": None}
    root = _parse_xml_part(
        parts, theme_name, code="DOCX-ROUNDTRIP-PACKAGE")
    drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"

    def latin(kind):
        element = root.find(
            f".//{{{drawing_ns}}}{kind}Font/{{{drawing_ns}}}latin")
        return None if element is None else element.get("typeface")

    return {"majorLatin": latin("major"), "minorLatin": latin("minor")}


def _explicit_family(style_projection):
    values = list(style_projection["explicitFonts"].values())
    if any(value is None or not value for value in values):
        return None
    unique = set(values)
    return values[0] if len(unique) == 1 else None


def _all_none(values):
    return all(value is None for value in values.values())


def _default_character_style_is_neutral(style_projection):
    if style_projection is None:
        return False
    return (
        style_projection["styleType"] == "character"
        and style_projection["basedOn"] is None
        and style_projection["link"] is None
        and _all_none(style_projection["explicitFonts"])
        and _all_none(style_projection["themeFonts"])
        and _all_none(style_projection["sizeHalfPoints"])
        and _all_none(style_projection["bold"])
        and _all_none(style_projection["italic"])
        and style_projection["color"] is None
    )


def _effective_false_toggle(style_projection, document_defaults, name):
    values = {}
    for script in ("latin", "complex"):
        direct = style_projection[name][script]
        inherited = document_defaults[name][script]
        values[script] = (
            direct
            if direct is not None
            else inherited
            if inherited is not None
            else False
        )
    return values == {"latin": False, "complex": False}


def _project_visual_styles(parts):
    styles_root = _parse_xml_part(
        parts, "word/styles.xml", code="DOCX-ROUNDTRIP-PACKAGE")
    styles_by_id = {
        style.get(qn("w:styleId")): style
        for style in styles_root.findall(qn("w:style"))
        if style.get(qn("w:styleId"))
    }
    projected = []
    projected_by_id = {}
    diagnostics = []
    normalizations = []

    defaults = styles_root.find(
        f"{qn('w:docDefaults')}/{qn('w:rPrDefault')}/{qn('w:rPr')}"
    )
    document_defaults = {
        "italic": {
            "latin": _on_off_projection(defaults, "i"),
            "complex": _on_off_projection(defaults, "iCs"),
        },
    }
    default_character_style = styles_by_id.get("DefaultParagraphFont")
    default_character_projection = (
        None
        if default_character_style is None
        else _project_controlled_style(default_character_style)
    )
    default_character_neutral = _default_character_style_is_neutral(
        default_character_projection
    )

    def drift(style_id, property_name, expected, actual):
        diagnostics.append({
            "code": VISUAL_STYLE_DRIFT_CODE,
            "severity": "warning",
            "styleId": style_id,
            "property": property_name,
            "expected": expected,
            "actual": actual,
            "message": (
                f"controlled style {style_id!r} has non-materialized "
                f"{property_name}"
            ),
        })

    def normalized(style_id, property_name, linked_style_id, proof):
        normalizations.append({
            "code": VISUAL_STYLE_NATIVE_NORMALIZATION_CODE,
            "severity": "info",
            "styleId": style_id,
            "property": property_name,
            "effectiveState": "materialized-equivalent",
            "proof": {
                "baseStyleId": "Normal",
                "linkedStyleId": linked_style_id,
                **proof,
            },
            "message": (
                f"controlled style {style_id!r} has Word-native omitted "
                f"{property_name} proven equivalent through its exact "
                "Normal/linked-style cascade"
            ),
        })

    for style_id in CONTROLLED_STYLE_IDS:
        style = styles_by_id.get(style_id)
        if style is None:
            drift(style_id, "presence", True, False)
            continue
        item = _project_controlled_style(style)
        projected.append(item)
        projected_by_id[style_id] = item

    normal = projected_by_id.get("Normal")
    body_family = _explicit_family(normal) if normal is not None else None
    if normal is not None:
        if normal["styleType"] != "paragraph":
            drift("Normal", "styleType", "paragraph", normal["styleType"])
        if normal["basedOn"] is not None:
            drift("Normal", "basedOn", None, normal["basedOn"])
        if normal["link"] is not None:
            drift("Normal", "link", None, normal["link"])
        if body_family is None:
            drift(
                "Normal",
                "explicitFonts",
                "one non-empty family in all four script slots",
                normal["explicitFonts"],
            )
        if any(normal["themeFonts"].values()):
            drift("Normal", "themeFonts", {}, normal["themeFonts"])
        sizes = normal["sizeHalfPoints"]
        if (
            sizes["latin"] is None
            or sizes["latin"] != sizes["complex"]
        ):
            drift(
                "Normal",
                "sizeHalfPoints",
                "equal explicit latin and complex sizes",
                sizes,
            )

    for level in range(1, 5):
        paragraph_id = f"Heading{level}"
        character_id = f"Heading{level}Char"
        paragraph = projected_by_id.get(paragraph_id)
        character = projected_by_id.get(character_id)
        if paragraph is not None:
            if paragraph["styleType"] != "paragraph":
                drift(
                    paragraph_id,
                    "styleType",
                    "paragraph",
                    paragraph["styleType"],
                )
            if paragraph["basedOn"] != "Normal":
                drift(
                    paragraph_id,
                    "basedOn",
                    "Normal",
                    paragraph["basedOn"],
                )
        if character is not None:
            if character["styleType"] != "character":
                drift(
                    character_id,
                    "styleType",
                    "character",
                    character["styleType"],
                )
            if character["basedOn"] != "DefaultParagraphFont":
                drift(
                    character_id,
                    "basedOn",
                    "DefaultParagraphFont",
                    character["basedOn"],
                )

        if paragraph is not None and character is not None:
            links_exact = (
                paragraph["link"] == character_id
                and character["link"] == paragraph_id
            )
            normal_is_exact_base = (
                normal is not None
                and normal["styleType"] == "paragraph"
                and normal["basedOn"] is None
                and normal["link"] is None
                and body_family is not None
                and not any(normal["themeFonts"].values())
            )
            character_is_exact_link = (
                character["styleType"] == "character"
                and character["basedOn"] == "DefaultParagraphFont"
                and _explicit_family(character) == body_family
                and not any(character["themeFonts"].values())
                and default_character_neutral
            )
            paragraph_is_exact_child = (
                paragraph["styleType"] == "paragraph"
                and paragraph["basedOn"] == "Normal"
                and not any(paragraph["themeFonts"].values())
            )

            paragraph_family = _explicit_family(paragraph)
            paragraph_fonts_native = (
                _all_none(paragraph["explicitFonts"])
                and paragraph_is_exact_child
                and normal_is_exact_base
                and character_is_exact_link
                and links_exact
            )
            if paragraph_fonts_native:
                normalized(
                    paragraph_id,
                    "explicitFonts",
                    character_id,
                    {
                        "omittedDirectValues": paragraph["explicitFonts"],
                        "baseValues": normal["explicitFonts"],
                        "linkedValues": character["explicitFonts"],
                        "effectiveValues": normal["explicitFonts"],
                    },
                )
            elif (
                paragraph_family is None
                or (
                    body_family is not None
                    and paragraph_family != body_family
                )
            ):
                drift(
                    paragraph_id,
                    "explicitFonts",
                    (
                        body_family
                        if body_family is not None
                        else "one matching explicit family"
                    ),
                    paragraph["explicitFonts"],
                )

            character_family = _explicit_family(character)
            if character_family is None or (
                body_family is not None
                and character_family != body_family
            ):
                drift(
                    character_id,
                    "explicitFonts",
                    (
                        body_family
                        if body_family is not None
                        else "one matching explicit family"
                    ),
                    character["explicitFonts"],
                )

            paragraph_italic_native = (
                _all_none(paragraph["italic"])
                and paragraph_is_exact_child
                and normal_is_exact_base
                and character_is_exact_link
                and links_exact
                and character["italic"]
                == {"latin": False, "complex": False}
                and _effective_false_toggle(
                    normal,
                    document_defaults,
                    "italic",
                )
            )
            if paragraph_italic_native:
                normalized(
                    paragraph_id,
                    "italic",
                    character_id,
                    {
                        "omittedDirectValues": paragraph["italic"],
                        "baseValues": normal["italic"],
                        "documentDefaultValues": (
                            document_defaults["italic"]
                        ),
                        "linkedValues": character["italic"],
                        "effectiveValues": {
                            "latin": False,
                            "complex": False,
                        },
                    },
                )
            elif paragraph["italic"] != {
                "latin": False,
                "complex": False,
            }:
                drift(
                    paragraph_id,
                    "italic",
                    {"latin": False, "complex": False},
                    paragraph["italic"],
                )

            if character["italic"] != {
                "latin": False,
                "complex": False,
            }:
                drift(
                    character_id,
                    "italic",
                    {"latin": False, "complex": False},
                    character["italic"],
                )

            for style_id, item in (
                (paragraph_id, paragraph),
                (character_id, character),
            ):
                if any(item["themeFonts"].values()):
                    drift(
                        style_id,
                        "themeFonts",
                        {},
                        item["themeFonts"],
                    )
                sizes = item["sizeHalfPoints"]
                if (
                    sizes["latin"] is None
                    or sizes["latin"] != sizes["complex"]
                ):
                    drift(
                        style_id,
                        "sizeHalfPoints",
                        "equal explicit latin and complex sizes",
                        sizes,
                    )
                bold = item["bold"]
                if (
                    bold["latin"] is None
                    or bold["latin"] != bold["complex"]
                ):
                    drift(
                        style_id,
                        "bold",
                        "equal explicit latin and complex values",
                        bold,
                    )

            paragraph_for_compare = {
                **paragraph,
                "explicitFonts": (
                    normal["explicitFonts"]
                    if paragraph_fonts_native
                    else paragraph["explicitFonts"]
                ),
                "italic": (
                    {"latin": False, "complex": False}
                    if paragraph_italic_native
                    else paragraph["italic"]
                ),
            }
            comparable = (
                "explicitFonts",
                "themeFonts",
                "sizeHalfPoints",
                "bold",
                "italic",
                "color",
            )
            differences = {
                name: {
                    "paragraph": paragraph_for_compare[name],
                    "character": character[name],
                }
                for name in comparable
                if paragraph_for_compare[name] != character[name]
            }
            if differences:
                drift(
                    paragraph_id,
                    "linkedCharacterAgreement",
                    "matching controlled run properties",
                    differences,
                )
            if paragraph["link"] != character_id:
                drift(
                    paragraph_id,
                    "link",
                    character_id,
                    paragraph["link"],
                )
            if character["link"] != paragraph_id:
                drift(
                    character_id,
                    "link",
                    paragraph_id,
                    character["link"],
                )
        else:
            for style_id, item in (
                (paragraph_id, paragraph),
                (character_id, character),
            ):
                if item is None:
                    continue
                family = _explicit_family(item)
                if family is None or (
                    body_family is not None and family != body_family
                ):
                    drift(
                        style_id,
                        "explicitFonts",
                        (
                            body_family
                            if body_family is not None
                            else "one matching explicit family"
                        ),
                        item["explicitFonts"],
                    )
                if item["italic"] != {
                    "latin": False,
                    "complex": False,
                }:
                    drift(
                        style_id,
                        "italic",
                        {"latin": False, "complex": False},
                        item["italic"],
                    )
                if any(item["themeFonts"].values()):
                    drift(
                        style_id,
                        "themeFonts",
                        {},
                        item["themeFonts"],
                    )
                sizes = item["sizeHalfPoints"]
                if (
                    sizes["latin"] is None
                    or sizes["latin"] != sizes["complex"]
                ):
                    drift(
                        style_id,
                        "sizeHalfPoints",
                        "equal explicit latin and complex sizes",
                        sizes,
                    )
                bold = item["bold"]
                if (
                    bold["latin"] is None
                    or bold["latin"] != bold["complex"]
                ):
                    drift(
                        style_id,
                        "bold",
                        "equal explicit latin and complex values",
                        bold,
                    )

    return {
        "schema": VISUAL_STYLE_PROJECTION_SCHEMA,
        "state": (
            "drifted"
            if diagnostics
            else "native-normalized-materialized-equivalent"
            if normalizations
            else "materialized"
        ),
        "themeFonts": _theme_font_context(parts),
        "bodyFont": body_family,
        "documentDefaults": document_defaults,
        "cascadeStyles": {
            "DefaultParagraphFont": default_character_projection,
        },
        "styles": projected,
        "normalizations": normalizations,
        "diagnostics": diagnostics,
    }


# ---------- main body walk ----------

def heading_level(style_name):
    m = re.fullmatch(r"Heading (\d+)", style_name or "")
    return int(m.group(1)) if m else None


STORY_PART_PATTERN = re.compile(
    r"^word/(?:document|header\d+|footer\d+|"
    r"footnotes|endnotes|comments)\.xml$")
STORY_RELATIONSHIP_TYPES = {
    (
        "http://schemas.openxmlformats.org/officeDocument/"
        f"2006/relationships/{name}"
    )
    for name in ("header", "footer", "footnotes", "endnotes", "comments")
}
REVISION_TAGS = {
    qn(f"w:{name}") for name in (
        "ins",
        "del",
        "moveFrom",
        "moveTo",
        "pPrChange",
        "rPrChange",
        "tblPrChange",
        "tblGridChange",
        "trPrChange",
        "tcPrChange",
        "sectPrChange",
        "numberingChange",
        "moveFromRangeStart",
        "moveFromRangeEnd",
        "moveToRangeStart",
        "moveToRangeEnd",
        "customXmlInsRangeStart",
        "customXmlInsRangeEnd",
        "customXmlDelRangeStart",
        "customXmlDelRangeEnd",
        "customXmlMoveFromRangeStart",
        "customXmlMoveFromRangeEnd",
        "customXmlMoveToRangeStart",
        "customXmlMoveToRangeEnd",
        "cellIns",
        "cellDel",
        "cellMerge",
        "conflictIns",
        "conflictDel",
    )
}
PARAGRAPH_CHILD_TAGS = {
    qn(f"w:{name}") for name in (
        "pPr",
        "r",
        "bookmarkStart",
        "bookmarkEnd",
        "proofErr",
        "permStart",
        "permEnd",
        "commentRangeStart",
        "commentRangeEnd",
    )
}
RUN_CHILD_TAGS = {
    qn("w:rPr"),
    qn("w:t"),
    # Pagination-only layout marker emitted by Word. It carries no text and
    # does not alter the supported Markdown semantics.
    qn("w:lastRenderedPageBreak"),
}
UNSUPPORTED_RUN_PROPERTY_TAGS = {
    qn(f"w:{name}") for name in (
        "u",
        "strike",
        "dstrike",
        "vertAlign",
        "smallCaps",
        "caps",
        "emboss",
        "imprint",
        "outline",
        "shadow",
        "vanish",
        "webHidden",
        "specVanish",
        "highlight",
        "effect",
        "bdr",
        "fitText",
        "rStyle",
    )
}


def _contains_any(root, tags):
    return any(element.tag in tags for element in root.iter())


def _related_story_part_names(parts):
    rels_name = "word/_rels/document.xml.rels"
    rels_root = _parse_xml_part(
        parts, rels_name, code="DOCX-ROUNDTRIP-PACKAGE")
    if (
        rels_root.tag
        != f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationships"
    ):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PACKAGE",
            "document relationships use an unexpected namespace")
    relationship_tag = f"{{{PACKAGE_RELATIONSHIPS_NAMESPACE}}}Relationship"
    names = set()
    for relationship in _relationship_elements(
        rels_root,
        relationship_tag,
        code="DOCX-ROUNDTRIP-PACKAGE",
    ):
        if relationship.get("Type") not in STORY_RELATIONSHIP_TYPES:
            continue
        _require_internal_relationship(
            relationship,
            "Word story",
            code="DOCX-ROUNDTRIP-PACKAGE",
        )
        names.add(_resolve_relationship_target(
            "word/document.xml",
            relationship.get("Target", ""),
            code="DOCX-ROUNDTRIP-PACKAGE",
            description="Word story",
        ))
    return names


def _story_roots(parts):
    roots = {}
    story_names = {
        name for name in parts if STORY_PART_PATTERN.fullmatch(name)
    }
    story_names.update(_related_story_part_names(parts))
    for name in sorted(story_names):
        roots[name] = _parse_xml_part(
            parts, name, code="DOCX-ROUNDTRIP-PACKAGE")
    if "word/document.xml" not in roots:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PACKAGE",
            "DOCX package lacks word/document.xml")
    return roots


def _preflight_story_parts(parts):
    roots = _story_roots(parts)
    for part_name, root in roots.items():
        if _contains_any(root, REVISION_TAGS):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-TRACKED-CHANGES",
                f"accept or reject tracked changes in {part_name}")
        if next(root.iter(qn("w:txbxContent")), None) is not None:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-TEXTBOX",
                f"text boxes in {part_name} are unsupported")
        if (
            next(root.iter(qn("w:drawing")), None) is not None
            or next(root.iter(qn("w:pict")), None) is not None
            or next(root.iter(qn("w:object")), None) is not None
        ):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-IMAGE",
                f"images/drawings in {part_name} are unsupported")
        if next(root.iter(qn("w:hyperlink")), None) is not None:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-HYPERLINK",
                f"native Word hyperlinks in {part_name} are unsupported")
        if any(
            next(root.iter(qn(f"w:{name}")), None) is not None
            for name in ("fldSimple", "fldChar", "instrText")
        ):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-FIELD",
                f"Word fields in {part_name} are unsupported")
        if any(
            next(root.iter(qn(f"w:{name}")), None) is not None
            for name in ("footnoteReference", "endnoteReference")
        ):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-NOTE",
                f"footnotes/endnotes in {part_name} are unsupported")
        if next(root.iter(qn("w:sdt")), None) is not None:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-CONTENT-CONTROL",
                f"content controls in {part_name} are unsupported")
        if any(
            next(root.iter(qn(f"w:{name}")), None) is not None
            for name in ("customXml", "smartTag", "altChunk", "subDoc")
        ):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-UNSUPPORTED-ELEMENT",
                f"unsupported content wrapper in {part_name}")
        if any(
            next(root.iter(qn(f"w:{name}")), None) is not None
            for name in ("br", "cr", "tab", "pageBreakBefore")
        ):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-BREAK",
                f"tabs and hard/page breaks in {part_name} are unsupported")
        for r_pr in root.iter(qn("w:rPr")):
            if any(child.tag in UNSUPPORTED_RUN_PROPERTY_TAGS
                   for child in r_pr):
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-INLINE",
                    f"unsupported direct run formatting in {part_name}")
        for paragraph in root.iter(qn("w:p")):
            if any(child.tag not in PARAGRAPH_CHILD_TAGS
                   for child in paragraph):
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-UNSUPPORTED-ELEMENT",
                    f"unsupported paragraph content in {part_name}")
        for run in root.iter(qn("w:r")):
            if any(child.tag not in RUN_CHILD_TAGS for child in run):
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-UNSUPPORTED-ELEMENT",
                    f"unsupported run content in {part_name}")
        if next(root.iter(qn("w:gridSpan")), None) is not None or (
                next(root.iter(qn("w:vMerge")), None) is not None):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-TABLE",
                f"merged table cells in {part_name} are unsupported")
        for cell in root.iter(qn("w:tc")):
            if next(cell.iter(qn("w:tbl")), None) is not None:
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-TABLE",
                    f"nested tables in {part_name} are unsupported")

    document_root = roots["word/document.xml"]
    body = document_root.find(qn("w:body"))
    if body is None:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PACKAGE",
            "word/document.xml lacks a body")
    allowed_body_children = {qn("w:p"), qn("w:tbl"), qn("w:sectPr")}
    if any(child.tag not in allowed_body_children for child in body):
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-UNSUPPORTED-ELEMENT",
            "document body contains an unsupported block element")


def _preflight_supported_docx(doc):
    if len(doc.sections) != 1:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-SECTION",
            "the supported profile requires exactly one Word section")

    section = doc.sections[0]
    alternate_stories = (
        ("first-page header", section.first_page_header),
        ("even-page header", section.even_page_header),
        ("first-page footer", section.first_page_footer),
        ("even-page footer", section.even_page_footer),
    )
    for label, story in alternate_stories:
        if story.tables or any(
                paragraph.text.strip() for paragraph in story.paragraphs):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-HEADER",
                f"{label} content is outside the supported banner profile")

    header = section.header
    if header.tables:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-HEADER",
            "header tables are outside the supported banner profile")
    nonempty_header = [
        index for index, paragraph in enumerate(header.paragraphs)
        if paragraph.text.strip()
    ]
    if nonempty_header:
        if nonempty_header != [0]:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-HEADER",
                "the banner must be the sole first header paragraph")
        header_text = header.paragraphs[0].text.strip()
        if header.paragraphs[0].text != header_text:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-HEADER",
                "banner text must not have outer whitespace")
        if re.fullmatch(r"CUI[^*\r\n]*", header_text) is None:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-HEADER",
                "the only supported non-empty header is a CUI-shaped banner")
        footer = section.footer
        nonempty_footer = [
            index for index, paragraph in enumerate(footer.paragraphs)
            if paragraph.text.strip()
        ]
        if (
            footer.tables
            or nonempty_footer != [0]
            or footer.paragraphs[0].text != header_text
        ):
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-HEADER",
                "the banner header/footer copies must agree")


def _validate_canonical_inverse(markdown, expected_kinds):
    try:
        import md2docx
        canonical = md2docx.canonicalize_markdown(markdown)
        actual_kinds = tuple(
            block.kind for block in md2docx.parse_markdown_blocks(markdown))
    except (ImportError, ValueError) as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            f"inverted Markdown is outside the canonical profile: {exc}") from exc
    if canonical != markdown:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            "inverted Markdown is not byte-exact canonical Markdown")
    if tuple(expected_kinds) != actual_kinds:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            "Word styles and emitted Markdown block semantics disagree")


def _document_from_bytes(docx_bytes):
    try:
        return Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-PACKAGE",
            "input is not a readable WordprocessingML package",
        ) from exc


def _normalize_supported_body_paragraph(
    text,
    paragraph,
    paragraph_index,
):
    """Apply C3's one diagnosed Word whitespace normalization."""
    if not text:
        return text, None
    if not text.strip():
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            "non-code body paragraphs containing only whitespace are "
            "outside the supported profile",
        )
    if text != text.lstrip():
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            "supported Word paragraphs must not have leading whitespace",
        )

    trailing_spaces = len(text) - len(text.rstrip(" "))
    if trailing_spaces:
        normalized = text[:-trailing_spaces]
        if normalized != normalized.rstrip():
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-NONCANONICAL",
                "supported Word paragraphs must not have trailing "
                "non-U+0020 whitespace",
            )
        style = paragraph.style
        event = {
            "code": "DOCX-ROUNDTRIP-TRAILING-ASCII-SPACE",
            "severity": "warning",
            "story": "word/document.xml",
            "paragraphIndex": paragraph_index,
            "styleId": style.style_id if style is not None else "Normal",
            "styleName": style.name if style is not None else "Normal",
            "edge": "trailing",
            "codePoint": "U+0020",
            "count": trailing_spaces,
            "inputTextSha256": hashlib.sha256(
                text.encode("utf8")
            ).hexdigest(),
            "message": (
                f"removed {trailing_spaces} trailing U+0020 "
                f"character{'s' if trailing_spaces != 1 else ''} from "
                f"body paragraph {paragraph_index}"
            ),
        }
        return normalized, event

    if text != text.rstrip():
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-NONCANONICAL",
            "supported Word paragraphs must not have trailing "
            "non-U+0020 whitespace",
        )
    return text, None


def _convert_docx_bytes_with_report(
    docx_bytes,
    input_name,
    report_provenance=True,
    require_embedded=False,
):
    parts = _read_safe_docx_parts_bytes(docx_bytes)
    snapshot = _read_embedded_source_parts(
        parts, required=require_embedded)
    _preflight_story_parts(parts)
    visual_style_projection = _project_visual_styles(parts)
    doc = _document_from_bytes(docx_bytes)
    _preflight_supported_docx(doc)
    if report_provenance:
        print_provenance(doc)

    blocks = []
    expected_kinds = []
    normalizations = []
    body_paragraph_index = -1

    # Page-header banner (e.g. a CUI-style marking) -> leading **BANNER** line.
    sec = doc.sections[0]
    hdr = sec.header.paragraphs[0].text.strip() if sec.header.paragraphs else ""
    if hdr:
        blocks.append(f"**{hdr}**")
        expected_kinds.append("banner")

    code_buf = None  # accumulating fenced-code lines

    def flush_code():
        nonlocal code_buf
        if code_buf is not None:
            body = "\n".join(l.rstrip() for l in code_buf)
            blocks.append("```\n" + body + "\n```")
            expected_kinds.append("code")
            code_buf = None

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]

        if tag == "tbl":
            flush_code()
            md = table_to_md(Table(child, doc))
            if md:
                blocks.append(md)
                expected_kinds.append("table")
            continue

        if tag != "p":
            continue

        body_paragraph_index += 1
        p = Paragraph(child, doc)

        # Code-block line: paragraph shading.
        if para_shd_fill(p) is not None:
            line = "".join(run.text for run in p.runs)
            if line.strip() and line != line.rstrip():
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-NONCANONICAL",
                    "code lines must not have trailing whitespace")
            if not line.strip():
                line = ""
            if code_buf is None:
                code_buf = []
            code_buf.append(line)
            continue

        flush_code()

        style = p.style.name if p.style else "Normal"
        if style.startswith("List Bullet") and style != "List Bullet":
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-NESTED-LIST",
                f"paragraph style {style!r} is not a supported flat bullet")
        if style not in SUPPORTED_BODY_STYLES:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-STYLE",
                f"paragraph style {style!r} is outside the supported profile")
        if _pPr_child(p._p, "w:numPr") is not None:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-NUMBERING",
                "native Word numbering is not the literal-number profile")
        text = runs_to_md(p.runs)

        # Blank spacer paragraphs the forward tool inserts between blocks.
        if not text:
            continue
        text, normalization = _normalize_supported_body_paragraph(
            text,
            p,
            body_paragraph_index,
        )
        if normalization is not None:
            normalizations.append(normalization)

        lvl = heading_level(style)
        if lvl:
            blocks.append("#" * lvl + " " + text.strip())
            expected_kinds.append("heading")
        elif style == "List Bullet":
            blocks.append("- " + text.strip())
            expected_kinds.append("bullet")
        elif para_has_left_border(p):
            blocks.append("> " + text.strip())
            expected_kinds.append("blockquote")
        else:
            blocks.append(text.rstrip())
            expected_kinds.append(
                "numbered"
                if re.match(r"^\d+\.\s+", text.strip())
                else "paragraph"
            )

    flush_code()

    md = "\n\n".join(b.rstrip() for b in blocks)
    md = "\n".join(line.rstrip() for line in md.split("\n"))
    md = md.rstrip() + "\n"
    _validate_canonical_inverse(md, expected_kinds)
    input_sha = hashlib.sha256(docx_bytes).hexdigest()
    diagnostics = []
    if snapshot is None:
        diagnostics.append({
            "code": "DOCX-ROUNDTRIP-NO-MERGE-BASE",
            "severity": "warning",
            "message": (
                "style inversion succeeded, but baseline-aware "
                "three-way merge is unavailable"
            ),
        })
    diagnostics.extend(normalizations)
    diagnostics.extend(visual_style_projection["normalizations"])
    diagnostics.extend(visual_style_projection["diagnostics"])
    report = {
        "schema": ROUNDTRIP_REPORT_SCHEMA,
        "state": (
            "normalized-supported-profile"
            if normalizations
            else "exact-supported-profile"
        ),
        "input": {
            "path": input_name,
            "sha256": input_sha,
        },
        "output": {
            "sha256": hashlib.sha256(md.encode("utf8")).hexdigest(),
            "bytes": len(md.encode("utf8")),
        },
        "embeddedSource": (
            {
                "state": "internally-consistent",
                "schema": snapshot.schema,
                "part": snapshot.part,
                "originalSha256": snapshot.original_sha256,
                "canonicalSha256": snapshot.canonical_sha256,
            }
            if snapshot is not None
            else {"state": "missing"}
        ),
        "semanticNormalization": {
            "state": "normalized" if normalizations else "exact",
            "events": normalizations,
        },
        "visualStyleProjection": visual_style_projection,
        "diagnostics": diagnostics,
    }
    return md, report, snapshot


def convert_with_report(docx_path, report_provenance=True):
    path = Path(docx_path)
    md, report, _ = _convert_docx_bytes_with_report(
        path.read_bytes(),
        path.name,
        report_provenance=report_provenance,
    )
    return md, report


def convert(docx_path, report_provenance=True):
    md, _ = convert_with_report(
        docx_path, report_provenance=report_provenance)
    return md


def _merge_docx_bytes_with_current(
    docx_bytes,
    input_name,
    current_path,
    report_provenance=True,
):
    """Merge canonical Word edits into a separately evolved canonical MD file."""
    edited_text, report, snapshot = _convert_docx_bytes_with_report(
        docx_bytes,
        input_name,
        report_provenance=report_provenance,
        require_embedded=True,
    )
    current_path = Path(current_path)
    try:
        current_text = current_path.read_bytes().decode("utf8")
    except UnicodeDecodeError as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-CURRENT-UTF8",
            "current Markdown source is not valid UTF-8") from exc

    try:
        import md2docx
        canonical_current = md2docx.canonicalize_markdown(current_text)
    except (ImportError, ValueError) as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-CURRENT-PROFILE",
            f"current Markdown is outside the supported profile: {exc}") from exc
    if canonical_current != current_text:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-CURRENT-NONCANONICAL",
            "normalize the current Markdown before three-way merge")

    with tempfile.TemporaryDirectory() as td:
        scratch = Path(td)
        current_file = scratch / "current.md"
        base_file = scratch / "base.md"
        edited_file = scratch / "edited.md"
        current_file.write_text(current_text, encoding="utf8")
        base_file.write_text(snapshot.canonical_text, encoding="utf8")
        edited_file.write_text(edited_text, encoding="utf8")
        try:
            merged = subprocess.run(
                [
                    "git",
                    "merge-file",
                    "-p",
                    "--diff3",
                    str(current_file),
                    str(base_file),
                    str(edited_file),
                ],
                check=False,
                capture_output=True,
            )
        except FileNotFoundError as exc:
            raise RoundtripRefusalError(
                "DOCX-ROUNDTRIP-MERGE-UNAVAILABLE",
                "git merge-file is unavailable") from exc
    conflicts = _merge_exit_has_conflicts(merged.returncode)
    if merged.returncode != 0 and not conflicts:
        diagnostic = merged.stderr.decode("utf8", errors="replace").strip()
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-MERGE",
            diagnostic or f"git merge-file failed with {merged.returncode}")
    try:
        merged_text = merged.stdout.decode("utf8")
    except UnicodeDecodeError as exc:
        raise RoundtripRefusalError(
            "DOCX-ROUNDTRIP-MERGE",
            "git merge-file returned non-UTF-8 output") from exc
    result = MergeResult(
        text=merged_text,
        conflicts=conflicts,
        base_sha256=snapshot.canonical_sha256,
        current_sha256=hashlib.sha256(current_text.encode("utf8")).hexdigest(),
        edited_sha256=hashlib.sha256(edited_text.encode("utf8")).hexdigest(),
    )
    return result, report, snapshot


def merge_with_current(docx_path, current_path, report_provenance=True):
    path = Path(docx_path)
    result, _, _ = _merge_docx_bytes_with_current(
        path.read_bytes(),
        path.name,
        current_path,
        report_provenance=report_provenance,
    )
    return result


def _merge_exit_has_conflicts(returncode):
    """Interpret git merge-file's documented conflict-count exit value."""
    return 1 <= returncode <= 127


def _paths_alias(first, second):
    first_path = Path(first)
    second_path = Path(second)
    first_resolved = first_path.resolve()
    second_resolved = second_path.resolve()
    if first_resolved == second_resolved:
        return True
    # Refuse prospective case/normalization aliases conservatively. This closes
    # the gap where two nonexistent spellings identify one destination on the
    # default macOS filesystem, before samefile() has an inode to compare.
    first_key = unicodedata.normalize(
        "NFC", str(first_resolved)).casefold()
    second_key = unicodedata.normalize(
        "NFC", str(second_resolved)).casefold()
    if first_key == second_key:
        return True
    try:
        return first_path.exists() and second_path.exists() and (
            first_path.samefile(second_path)
        )
    except OSError:
        return False


def _validate_cli_paths(roles):
    items = [(role, Path(path)) for role, path in roles if path is not None]
    for index, (first_role, first_path) in enumerate(items):
        for second_role, second_path in items[index + 1:]:
            if _paths_alias(first_path, second_path):
                raise RoundtripRefusalError(
                    "DOCX-ROUNDTRIP-PATH-ALIAS",
                    f"{first_role} and {second_role} resolve to the same file")


def _print_visual_style_warnings(report):
    projection = report.get("visualStyleProjection", {})
    for diagnostic in projection.get("diagnostics", ()):
        if diagnostic.get("code") != VISUAL_STYLE_DRIFT_CODE:
            continue
        actual = json.dumps(
            diagnostic.get("actual"),
            sort_keys=True,
            separators=(",", ":"),
        )
        print(
            f"{VISUAL_STYLE_DRIFT_CODE}: "
            f"{diagnostic.get('styleId')}.{diagnostic.get('property')} "
            f"has {actual}",
            file=sys.stderr,
        )


def _atomic_replace(source, destination):
    Path(source).replace(destination)


def _path_exists_without_following(path):
    try:
        Path(path).lstat()
        return True
    except FileNotFoundError:
        return False


def _publish_outputs(outputs):
    """Stage every output and restore destinations on replacement failure.

    Each individual replacement is atomic. Backup/rollback gives callers
    all-old or all-new results for ordinary in-process I/O failures; it cannot
    provide a filesystem-wide transaction across process or power loss.
    """
    staged = []
    backups = []
    published = []
    retained_backups = set()
    try:
        for destination, payload in outputs:
            destination = Path(destination)
            with tempfile.NamedTemporaryFile(
                prefix=f".{destination.name}.",
                suffix=".tmp",
                dir=destination.parent,
                delete=False,
            ) as handle:
                staged_path = Path(handle.name)
                handle.write(payload)
            staged.append((destination, staged_path))

        # Snapshot every old destination before replacing any output. Preserve
        # an output symlink as a symlink if rollback becomes necessary.
        for destination, _ in staged:
            backup_path = None
            if _path_exists_without_following(destination):
                with tempfile.NamedTemporaryFile(
                    prefix=f".{destination.name}.",
                    suffix=".bak",
                    dir=destination.parent,
                    delete=False,
                ) as handle:
                    backup_path = Path(handle.name)
                backup_path.unlink()
                backups.append((destination, backup_path))
                shutil.copy2(
                    destination,
                    backup_path,
                    follow_symlinks=False,
                )
            else:
                backups.append((destination, None))

        for destination, staged_path in staged:
            _atomic_replace(staged_path, destination)
            published.append(destination)
    except Exception as exc:
        rollback_errors = []
        backup_by_destination = dict(backups)
        for destination in reversed(published):
            backup_path = backup_by_destination.get(destination)
            try:
                if backup_path is None:
                    destination.unlink()
                else:
                    _atomic_replace(backup_path, destination)
            except OSError as rollback_exc:
                if (
                    backup_path is not None
                    and _path_exists_without_following(backup_path)
                ):
                    retained_backups.add(backup_path)
                rollback_errors.append(
                    (
                        f"{destination}: {rollback_exc}"
                        + (
                            f" (old bytes retained at {backup_path})"
                            if backup_path in retained_backups
                            else ""
                        )
                    ))
        if rollback_errors:
            raise OSError(
                "output publication failed and rollback was incomplete: "
                + "; ".join(rollback_errors)
            ) from exc
        raise
    finally:
        for _, staged_path in staged:
            try:
                staged_path.unlink()
            except OSError:
                pass
        for _, backup_path in backups:
            if backup_path is None or backup_path in retained_backups:
                continue
            try:
                backup_path.unlink()
            except OSError:
                pass


def main():
    ap = argparse.ArgumentParser(description="Styled DOCX -> canonical Markdown.")
    ap.add_argument("file", help="Input .docx")
    ap.add_argument("-o", "--out", help="Output .md (default: next to input)")
    ap.add_argument(
        "--base-out",
        help=(
            "Extract the internally consistent embedded canonical merge base "
            "to this path"
        ))
    ap.add_argument(
        "--merge-current",
        help="Three-way merge Word edits into this canonical Markdown source")
    ap.add_argument(
        "--report",
        help="Write a hash-bound JSON fidelity/reconciliation report")
    args = ap.parse_args()

    src = Path(args.file)
    if not src.exists():
        sys.exit(f"not found: {src}")
    if args.merge_current and not args.out:
        ap.error("--merge-current requires an explicit --out path")
    out = Path(args.out) if args.out else src.with_suffix(".md")

    try:
        _validate_cli_paths([
            ("input DOCX", src),
            ("current Markdown", args.merge_current),
            ("Markdown output", out),
            ("base output", args.base_out),
            ("report output", args.report),
        ])
        docx_bytes = src.read_bytes()
        if args.merge_current:
            merge, report, snapshot = _merge_docx_bytes_with_current(
                docx_bytes,
                src.name,
                args.merge_current,
            )
            md = merge.text
            report["merge"] = {
                "state": "conflicted" if merge.conflicts else "merged",
                "baseSha256": merge.base_sha256,
                "currentSha256": merge.current_sha256,
                "editedSha256": merge.edited_sha256,
                "outputSha256": hashlib.sha256(
                    merge.text.encode("utf8")).hexdigest(),
            }
        else:
            md, report, snapshot = _convert_docx_bytes_with_report(
                docx_bytes,
                src.name,
                require_embedded=bool(args.base_out),
            )
            merge = None
        _print_visual_style_warnings(report)
        outputs = [(out, md.encode("utf8"))]
        if args.base_out:
            outputs.append((
                Path(args.base_out),
                snapshot.canonical_text.encode("utf8"),
            ))
        if args.report:
            outputs.append((
                Path(args.report),
                (
                    json.dumps(report, indent=2, sort_keys=True) + "\n"
                ).encode("utf8"),
            ))
        _publish_outputs(outputs)
    except (RoundtripRefusalError, OSError) as exc:
        if isinstance(exc, RoundtripRefusalError):
            sys.exit(str(exc))
        sys.exit(f"DOCX-ROUNDTRIP-IO: {exc}")

    print(f"-> {out}")
    if args.base_out:
        print(f"-> {args.base_out}")
    if args.report:
        print(f"-> {args.report}")
    if merge is not None and merge.conflicts:
        print(
            "DOCX-ROUNDTRIP-MERGE-CONFLICT: review conflict markers",
            file=sys.stderr)
        raise SystemExit(1)


if __name__ == "__main__":
    main()
