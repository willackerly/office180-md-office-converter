#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""Generate privacy-safe, content-bound DOCX round-trip evidence.

CONTRACT:C2-PROVENANCE.2.0
CONTRACT:C3-ROUNDTRIP.1.2
CONTRACT:C11-OFFICE-VISUAL-EVIDENCE.1.1

The generator refuses every existing destination, builds under a sibling
staging directory, and publishes only after structural checks, Quick Look
captures, an exact same-renderer comparison, manifest validation, and privacy
scans all pass. Native Microsoft Word remains an explicit manual/unavailable
gate; the supported edit in this automated lane is a python-docx plain-text
save that preserves the target paragraph and run formatting.
"""

from __future__ import annotations

import argparse
import getpass
import hashlib
import importlib.metadata
import importlib.util
import json
import os
import platform
import plistlib
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

import docx2md  # noqa: E402
import md2docx  # noqa: E402


VISUAL_SCRIPT = REPO_ROOT / "scripts" / "visual-evidence.py"
VISUAL_SPEC = importlib.util.spec_from_file_location(
    "office180_visual_evidence_generator", VISUAL_SCRIPT
)
if VISUAL_SPEC is None or VISUAL_SPEC.loader is None:
    raise RuntimeError(f"cannot load visual evidence harness: {VISUAL_SCRIPT}")
visual = importlib.util.module_from_spec(VISUAL_SPEC)
sys.modules[VISUAL_SPEC.name] = visual
VISUAL_SPEC.loader.exec_module(visual)


DEFAULT_DESTINATION = (
    REPO_ROOT / "tests" / "fixtures" / "roundtrip-evidence" / "docx"
)
SOURCE_FIXTURE = REPO_ROOT / "tests" / "fixtures" / "kitchen-sink.md"
THEME = REPO_ROOT / "themes" / "neutral.json"
FIXED_SCRATCH_ROOT = Path("/private/tmp/office180-evidence")
FIXED_SOURCE_DIR = FIXED_SCRATCH_ROOT / "source"
EDIT_BEFORE = (
    "This paragraph soft-wraps across several source lines that should join "
    "into one paragraph in the output."
)
EDIT_AFTER = (
    "This reviewed paragraph remains one supported plain-text paragraph "
    "after the Word-like edit."
)
# Conservative full-page Quick Look crop containing only the edited paragraph
# line plus surrounding whitespace: [x, y, width, height].
EXPECTED_EDIT_CROP = [170, 220, 900, 60]
SECRET_PATTERNS = {
    "private-key": re.compile(
        rb"-----BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY-----"
    ),
    "aws-access-key": re.compile(rb"\b(?:AKIA|ASIA)[0-9A-Z]{16}\b"),
    "github-token": re.compile(
        rb"\b(?:gh[opsu]_[A-Za-z0-9]{30,}|github_pat_[A-Za-z0-9_]{30,})\b"
    ),
    "openai-token": re.compile(rb"\bsk-[A-Za-z0-9_-]{20,}\b"),
    "slack-token": re.compile(rb"\bxox[baprs]-[A-Za-z0-9-]{20,}\b"),
}
SENSITIVE_ENV_NAME = re.compile(
    r"(?:TOKEN|SECRET|PASSWORD|PASSWD|API_KEY|ACCESS_KEY|PRIVATE_KEY)",
    re.IGNORECASE,
)


class EvidenceGenerationError(RuntimeError):
    """The evidence workflow cannot publish a complete trusted result."""


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _lexists(path: Path) -> bool:
    try:
        path.lstat()
        return True
    except FileNotFoundError:
        return False


def _relative(path: Path, root: Path) -> str:
    return path.resolve().relative_to(root.resolve()).as_posix()


def _resolve_destination(value: Path) -> Path:
    destination = value if value.is_absolute() else REPO_ROOT / value
    destination = destination.resolve()
    try:
        destination.relative_to(REPO_ROOT.resolve())
    except ValueError as exc:
        raise EvidenceGenerationError(
            "destination must stay inside the repository"
        ) from exc
    if _lexists(destination):
        raise EvidenceGenerationError(
            f"refusing existing destination: {destination}"
        )
    if not destination.parent.is_dir():
        raise EvidenceGenerationError(
            f"destination parent does not exist: {destination.parent}"
        )
    return destination


def _write_new_bytes(path: Path, payload: bytes) -> None:
    if _lexists(path):
        raise EvidenceGenerationError(f"refusing to overwrite: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("xb") as stream:
        stream.write(payload)


def _write_new_text(path: Path, text: str) -> None:
    _write_new_bytes(path, text.encode("utf-8"))


def _write_new_json(path: Path, value: Any) -> None:
    _write_new_text(
        path,
        json.dumps(
            value,
            ensure_ascii=False,
            allow_nan=False,
            indent=2,
            sort_keys=True,
        )
        + "\n",
    )


def _load_theme() -> dict[str, Any]:
    raw = json.loads(THEME.read_text(encoding="utf-8"))
    return md2docx.deep_merge(md2docx.DEFAULTS, raw)


def _canonical_source() -> str:
    raw = SOURCE_FIXTURE.read_bytes().decode("utf-8")
    canonical = md2docx.canonicalize_markdown(raw)
    if md2docx.canonicalize_markdown(canonical) != canonical:
        raise EvidenceGenerationError("canonical source is not idempotent")
    return canonical


def _generate_docx(source: Path, output: Path) -> None:
    if _lexists(output):
        raise EvidenceGenerationError(f"refusing to overwrite: {output}")
    md2docx.Converter(_load_theme(), tpl_path=THEME).convert(source, output)
    if not output.is_file() or output.stat().st_size == 0:
        raise EvidenceGenerationError(f"DOCX generation was empty: {output}")


def _run_format_signature(run: Any) -> dict[str, Any]:
    color = run.font.color.rgb
    size = run.font.size
    return {
        "bold": run.bold,
        "italic": run.italic,
        "font": run.font.name,
        "size_pt": size.pt if size is not None else None,
        "color": str(color) if color is not None else None,
    }


def _apply_plain_text_edit(source: Path, destination: Path) -> dict[str, Any]:
    if _lexists(destination):
        raise EvidenceGenerationError(f"refusing to overwrite: {destination}")
    document = Document(str(source))
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == EDIT_BEFORE
    ]
    if len(matches) != 1 or len(matches[0].runs) != 1:
        raise EvidenceGenerationError(
            "plain-text edit target is not one unambiguous single-run paragraph"
        )
    paragraph = matches[0]
    run = paragraph.runs[0]
    style_before = paragraph.style.name
    formatting_before = _run_format_signature(run)
    run.text = EDIT_AFTER
    if (
        paragraph.style.name != style_before
        or _run_format_signature(run) != formatting_before
    ):
        raise EvidenceGenerationError(
            "plain-text edit changed paragraph or run formatting"
        )
    document.save(destination)

    reopened = Document(str(destination))
    reopened_matches = [
        paragraph
        for paragraph in reopened.paragraphs
        if paragraph.text == EDIT_AFTER
    ]
    if len(reopened_matches) != 1 or len(reopened_matches[0].runs) != 1:
        raise EvidenceGenerationError(
            "edited paragraph did not survive package save unambiguously"
        )
    reopened_paragraph = reopened_matches[0]
    formatting_after = _run_format_signature(reopened_paragraph.runs[0])
    if (
        reopened_paragraph.style.name != style_before
        or formatting_after != formatting_before
    ):
        raise EvidenceGenerationError(
            "edited paragraph formatting changed after package reopen"
        )
    return {
        "method": "python-docx single-run text replacement and package save",
        "before": EDIT_BEFORE,
        "after": EDIT_AFTER,
        "paragraph_style": style_before,
        "run_formatting": formatting_after,
    }


def _check_docx(path: Path, expected_subject: Path) -> None:
    with zipfile.ZipFile(path, "r") as package:
        if package.testzip() is not None:
            raise EvidenceGenerationError(f"ZIP validation failed: {path.name}")
    package = Document(str(path))
    expected = str(expected_subject.resolve())[-255:]
    if package.core_properties.subject != expected:
        raise EvidenceGenerationError(
            f"unexpected C2 subject in {path.name}"
        )
    snapshot = docx2md.read_embedded_source(path, required=True)
    if snapshot is None:
        raise EvidenceGenerationError(
            f"embedded source missing from {path.name}"
        )


def _recover_edited(
    edited_docx: Path,
    evidence_root: Path,
    canonical: str,
) -> tuple[str, dict[str, Any]]:
    recovered, report = docx2md.convert_with_report(
        edited_docx, report_provenance=False
    )
    if md2docx.canonicalize_markdown(recovered) != recovered:
        raise EvidenceGenerationError("recovered Markdown is not canonical")
    if EDIT_AFTER not in recovered or EDIT_BEFORE in recovered:
        raise EvidenceGenerationError("supported edit is absent after recovery")
    snapshot = docx2md.read_embedded_source(edited_docx, required=True)
    if snapshot is None or snapshot.canonical_text != canonical:
        raise EvidenceGenerationError(
            "edited DOCX no longer binds the generated canonical base"
        )
    if (
        report.get("state") != "exact-supported-profile"
        or report.get("embeddedSource", {}).get("state")
        != "internally-consistent"
    ):
        raise EvidenceGenerationError(
            "round-trip report did not establish exact supported-profile state"
        )
    _write_new_text(evidence_root / "supported-recovered.md", recovered)
    _write_new_text(
        evidence_root / "supported-embedded-base.md",
        snapshot.canonical_text,
    )
    _write_new_json(evidence_root / "supported-edited.report.json", report)
    return recovered, report


def _capture_quicklook(
    artifact: Path,
    evidence_root: Path,
    stem: str,
    checkpoint: str,
) -> dict[str, Any]:
    image = evidence_root / f"{stem}.quicklook.png"
    manifest_path = evidence_root / f"{stem}.quicklook.evidence.json"
    manifest = visual.capture_quicklook(
        artifact,
        image,
        root=evidence_root,
        lane="markdown-docx",
        checkpoint=checkpoint,
        trusted=True,
        pixel_size=1600,
        timeout_seconds=30.0,
    )
    if manifest["capture"]["status"] != "passed":
        raise EvidenceGenerationError(
            f"Quick Look capture did not pass for {artifact.name}: "
            f"{manifest['capture']['diagnostic']}"
        )
    visual.write_manifest(manifest_path, manifest)
    return manifest


def _word_identity() -> tuple[str, str, str, str]:
    application = Path("/Applications/Microsoft Word.app")
    info = application / "Contents" / "Info.plist"
    if not application.is_dir():
        return (
            "unavailable",
            "unavailable",
            "Microsoft Word is not installed at the standard application path; "
            "no native lifecycle was attempted.",
            "unavailable",
        )
    version = "installed-version-unavailable"
    try:
        with info.open("rb") as stream:
            plist = plistlib.load(stream)
        candidate = plist.get("CFBundleShortVersionString")
        if isinstance(candidate, str) and candidate:
            version = candidate
    except (OSError, plistlib.InvalidFileException):
        pass
    return (
        "manual-required",
        version,
        "Native Microsoft Word open, representative edit, save, and reopen "
        "were not automated by this generator; manual lifecycle validation "
        "remains required.",
        application.as_posix(),
    )


def _record_native_status(
    edited_docx: Path,
    evidence_root: Path,
) -> dict[str, Any]:
    status, version, diagnostic, executable = _word_identity()
    manifest = visual.record_status(
        edited_docx,
        root=evidence_root,
        lane="markdown-docx",
        checkpoint="native-word-edit-save-reopen",
        renderer_class="native-word",
        status=status,
        diagnostic=diagnostic,
        product="Microsoft Word",
        version=version,
        executable_path=executable,
    )
    visual.write_manifest(
        evidence_root / "supported-edited.native-word.evidence.json",
        manifest,
    )
    return manifest


def _same_renderer(
    baseline: dict[str, Any],
    candidate: dict[str, Any],
) -> bool:
    baseline_capture = baseline["capture"]
    candidate_capture = candidate["capture"]
    return all(
        baseline_capture[key] == candidate_capture[key]
        for key in ("renderer_class", "renderer", "environment", "fonts", "input")
    )


def _compare_edited_and_regenerated(
    edited: dict[str, Any],
    regenerated: dict[str, Any],
    evidence_root: Path,
) -> dict[str, Any]:
    if not _same_renderer(edited, regenerated):
        raise EvidenceGenerationError(
            "edited and regenerated captures do not identify the same renderer"
        )
    comparison = visual.compare_evidence(
        edited,
        regenerated,
        root=evidence_root,
        antialias_tolerance=0,
        max_changed_fraction=0.0,
        max_mean_absolute_error=0.0,
        max_channel_delta=0,
        max_masked_fraction=0.0,
    )
    if comparison.get("comparison", {}).get("status") != "passed":
        raise EvidenceGenerationError(
            "edited and regenerated Quick Look captures are not exact"
        )
    visual.write_manifest(
        evidence_root
        / "supported-edited-vs-regenerated.comparison.evidence.json",
        comparison,
    )
    return comparison


def _bounds_within_crop(
    bounds: list[int],
    crop: list[int],
) -> bool:
    x, y, width, height = bounds
    crop_x, crop_y, crop_width, crop_height = crop
    return (
        x >= crop_x
        and y >= crop_y
        and x + width <= crop_x + crop_width
        and y + height <= crop_y + crop_height
    )


def _compare_canonical_and_edited(
    canonical: dict[str, Any],
    edited: dict[str, Any],
    evidence_root: Path,
) -> dict[str, Any]:
    if not _same_renderer(canonical, edited):
        raise EvidenceGenerationError(
            "canonical and edited captures do not identify the same renderer"
        )
    comparison = visual.compare_evidence(
        canonical,
        edited,
        root=evidence_root,
        antialias_tolerance=0,
        max_changed_fraction=0.01,
        max_mean_absolute_error=1.0,
        max_channel_delta=255,
        max_masked_fraction=0.0,
    )
    comparison_data = comparison.get("comparison", {})
    metrics = comparison_data.get("metrics")
    changed_bounds = comparison_data.get("changed_bounds")
    if (
        comparison_data.get("status") != "passed"
        or not isinstance(metrics, dict)
        or metrics.get("changed_pixels", 0) <= 0
        or not isinstance(changed_bounds, list)
        or not _bounds_within_crop(changed_bounds, EXPECTED_EDIT_CROP)
    ):
        raise EvidenceGenerationError(
            "canonical-to-edited Quick Look change is zero, exceeds its "
            "bounded thresholds, or escapes the expected paragraph crop"
        )
    comparison_data["crop_regions"] = [EXPECTED_EDIT_CROP]
    comparison = visual.finalize_evidence(comparison)
    visual.write_manifest(
        evidence_root
        / "supported-canonical-vs-edited.comparison.evidence.json",
        comparison,
    )
    return comparison


def _validate_evidence_manifests(evidence_root: Path) -> None:
    for path in sorted(evidence_root.glob("*.evidence.json")):
        manifest = visual.load_manifest(path)
        errors = visual.validate_manifest_data(
            manifest, evidence_root, check_files=True
        )
        if errors:
            raise EvidenceGenerationError(
                f"invalid evidence manifest {path.name}: {'; '.join(errors)}"
            )


def _sensitive_environment_values() -> list[tuple[str, bytes]]:
    values: list[tuple[str, bytes]] = []
    for name, value in os.environ.items():
        if (
            SENSITIVE_ENV_NAME.search(name)
            and isinstance(value, str)
            and len(value) >= 8
        ):
            values.append((name, value.encode("utf-8")))
    return values


def _privacy_findings(label: str, payload: bytes) -> list[str]:
    findings: list[str] = []
    username = getpass.getuser().encode("utf-8")
    hostname = platform.node().encode("utf-8")
    home = str(Path.home()).encode("utf-8")
    checks: list[tuple[str, bytes]] = [
        ("macOS-home-prefix", b"/Users/"),
        ("current-home-path", home),
        ("current-hostname", hostname),
    ]
    for name, token in checks:
        if token and token != b"/" and token in payload:
            findings.append(f"{label}: contains {name}")
    if username:
        username_pattern = re.compile(
            rb"(?<![A-Za-z0-9_])"
            + re.escape(username)
            + rb"(?![A-Za-z0-9_])",
            re.IGNORECASE,
        )
        if username_pattern.search(payload):
            findings.append(f"{label}: contains current-account-name")
    for name, pattern in SECRET_PATTERNS.items():
        if pattern.search(payload):
            findings.append(f"{label}: contains {name} pattern")
    for env_name, secret in _sensitive_environment_values():
        if secret in payload:
            findings.append(
                f"{label}: contains a sensitive environment value from "
                f"{env_name}"
            )
    return findings


def _iter_package_payloads(path: Path) -> Iterable[tuple[str, bytes]]:
    if path.suffix.lower() != ".docx":
        return
    with zipfile.ZipFile(path, "r") as package:
        for info in package.infolist():
            yield f"{path.name}!/{info.filename}", package.read(info)


def _scan_privacy(
    evidence_root: Path,
    extra_files: Iterable[Path] = (),
) -> dict[str, int]:
    findings: list[str] = []
    files = [
        path
        for path in sorted(evidence_root.rglob("*"))
        if path.is_file()
    ]
    package_members = 0
    for path in [*files, *extra_files]:
        label = (
            _relative(path, evidence_root)
            if path.is_relative_to(evidence_root)
            else f"$FIXED_SOURCE/{path.name}"
        )
        findings.extend(_privacy_findings(label, path.read_bytes()))
        for member_label, payload in _iter_package_payloads(path):
            package_members += 1
            findings.extend(_privacy_findings(member_label, payload))
    if findings:
        raise EvidenceGenerationError(
            "privacy scan failed: " + "; ".join(findings)
        )
    return {"files": len(files), "docx_package_members": package_members}


def _artifact_inventory(
    evidence_root: Path,
    excluded: set[str] | None = None,
) -> list[dict[str, Any]]:
    excluded = excluded or set()
    return [
        {
            "path": _relative(path, evidence_root),
            "sha256": sha256_file(path),
            "bytes": path.stat().st_size,
        }
        for path in sorted(evidence_root.rglob("*"))
        if path.is_file() and _relative(path, evidence_root) not in excluded
    ]


def _generation_manifest(
    evidence_root: Path,
    edit: dict[str, Any],
    native: dict[str, Any],
    edit_comparison: dict[str, Any],
    exact_comparison: dict[str, Any],
    scan: dict[str, int],
) -> dict[str, Any]:
    edit_comparison_data = edit_comparison["comparison"]
    edit_metrics = edit_comparison_data["metrics"]
    exact_metrics = exact_comparison["comparison"]["metrics"]
    return {
        "schema": "office180-docx-roundtrip-evidence-generation/0.1",
        "generator": {
            "path": "scripts/generate-docx-roundtrip-evidence.py",
            "sha256": sha256_file(Path(__file__)),
            "runtime": {
                "python": platform.python_version(),
                "python_docx": importlib.metadata.version("python-docx"),
            },
            "dependencies": [
                {
                    "path": "md2docx.py",
                    "sha256": sha256_file(REPO_ROOT / "md2docx.py"),
                },
                {
                    "path": "docx2md.py",
                    "sha256": sha256_file(REPO_ROOT / "docx2md.py"),
                },
                {
                    "path": "scripts/visual-evidence.py",
                    "sha256": sha256_file(VISUAL_SCRIPT),
                },
            ],
        },
        "source": {
            "fixture_path": "tests/fixtures/kitchen-sink.md",
            "fixture_sha256": sha256_file(SOURCE_FIXTURE),
            "fixed_generation_path": (
                "/private/tmp/office180-evidence/source/"
                "supported-canonical.md"
            ),
            "theme_path": "themes/neutral.json",
            "theme_sha256": sha256_file(THEME),
        },
        "edit": edit,
        "quicklook_comparisons": {
            "canonical_to_edited": {
                "status": edit_comparison_data["status"],
                "changed_pixels": edit_metrics["changed_pixels"],
                "changed_fraction": edit_metrics["changed_fraction"],
                "mean_absolute_error": edit_metrics[
                    "mean_absolute_error"
                ],
                "max_channel_delta": edit_metrics["max_channel_delta"],
                "changed_bounds": edit_comparison_data["changed_bounds"],
                "expected_crop": EXPECTED_EDIT_CROP,
            },
            "edited_to_regenerated": {
                "status": exact_comparison["comparison"]["status"],
                "changed_pixels": exact_metrics["changed_pixels"],
                "changed_fraction": exact_metrics["changed_fraction"],
                "mean_absolute_error": exact_metrics[
                    "mean_absolute_error"
                ],
                "max_channel_delta": exact_metrics["max_channel_delta"],
            },
        },
        "native_word": {
            "status": native["native_lifecycle"]["status"],
            "diagnostic": native["native_lifecycle"]["diagnostic"],
        },
        "privacy_scan": {
            "status": "passed",
            "scope": (
                "Every durable file, every DOCX member, both fixed temporary "
                "Markdown sources, and sensitive environment values"
            ),
            **scan,
        },
        "commands": [
            [
                ".venv/bin/python",
                "scripts/generate-docx-roundtrip-evidence.py",
                "--destination",
                "tests/fixtures/roundtrip-evidence/docx",
            ],
            [
                ".venv/bin/python",
                "scripts/visual-evidence.py",
                "validate",
                "<manifest>",
                "--root",
                "tests/fixtures/roundtrip-evidence/docx",
            ],
            [
                "shasum",
                "-a",
                "256",
                "-c",
                "SHA256SUMS",
            ],
        ],
        "limitations": [
            "Quick Look is a first-page automated preview, not native Word.",
            "The automated edit uses python-docx and preserves one supported "
            "paragraph/run formatting signature.",
            "Native Word open, edit, save, and reopen remains manual or "
            "unavailable as recorded in its evidence envelope.",
            "Quick Look uses host fonts; no exact font-byte claim is made.",
            "DOCX package bytes vary across runs because provenance and ZIP "
            "metadata include generation-time values.",
            "No C11 human-review envelope is claimed by this generator.",
        ],
        "artifacts": _artifact_inventory(evidence_root),
    }


def _readme(
    evidence_root: Path,
    edit_comparison: dict[str, Any],
    exact_comparison: dict[str, Any],
    native: dict[str, Any],
) -> str:
    inventory = _artifact_inventory(
        evidence_root,
        excluded={"README.md", "SHA256SUMS"},
    )
    hashes = "\n".join(
        f"| `{entry['path']}` | `{entry['sha256']}` | {entry['bytes']} |"
        for entry in inventory
    )
    edit_comparison_data = edit_comparison["comparison"]
    edit_metrics = edit_comparison_data["metrics"]
    exact_metrics = exact_comparison["comparison"]["metrics"]
    native_status = native["native_lifecycle"]["status"]
    native_diagnostic = native["native_lifecycle"]["diagnostic"]
    return f"""# DOCX round-trip evidence

This isolated C2/C3/C11 fixture is generated from the public kitchen-sink
Markdown profile. The source used for both DOCX generations is staged beneath
`/private/tmp/office180-evidence/source/`, keeping the legacy C2 absolute
subject free of workstation account paths.

## Reproduce and verify

The generator refuses any existing destination:

```bash
.venv/bin/python scripts/generate-docx-roundtrip-evidence.py \\
  --destination tests/fixtures/roundtrip-evidence/docx
```

Validate each C11 envelope with:

```bash
for manifest in tests/fixtures/roundtrip-evidence/docx/*.evidence.json; do
  .venv/bin/python scripts/visual-evidence.py validate "$manifest" \\
    --root tests/fixtures/roundtrip-evidence/docx
done
```

From this directory, verify all durable file hashes with:

```bash
shasum -a 256 -c SHA256SUMS
```

## Automated observations

- Canonical, plain-text-edited, and regenerated DOCX files passed the same
  Quick Look Office-generator capture profile.
- Canonical versus edited comparison found
  `{edit_metrics['changed_pixels']}` changed pixels with exact bounds
  `{edit_comparison_data['changed_bounds']}`. Those bounds are contained by
  the declared paragraph crop `{EXPECTED_EDIT_CROP}`; a zero change or any
  escaped pixel fails generation.
- Edited versus regenerated preview comparison passed with
  `{exact_metrics['changed_pixels']}` changed pixels, changed fraction
  `{exact_metrics['changed_fraction']}`, mean absolute error
  `{exact_metrics['mean_absolute_error']}`, and maximum channel delta
  `{exact_metrics['max_channel_delta']}`.
- Native Word state is `{native_status}`: {native_diagnostic}
- The final privacy pass covers every durable file, every uncompressed DOCX
  member, both fixed temporary Markdown sources, account/home/host tokens,
  high-confidence credential patterns, and sensitive environment values.

## Exact artifact hashes

| Artifact | SHA-256 | Bytes |
|----------|---------|------:|
{hashes}

## Limitations

- Quick Look is a first-page automated preview, not native Word lifecycle
  evidence.
- The supported plain-text edit is a python-docx simulation; native Word
  open/edit/save/reopen remains the separately recorded gate.
- Quick Look relies on host fonts, so no exact-font or cross-host equivalence
  claim is made.
- DOCX bytes are content-bound for this run but not byte-reproducible across
  runs because provenance and ZIP metadata include generation-time values.
- No mask, tolerance, native pass, or C11 human-review envelope is claimed.
"""


def _write_sha256s(evidence_root: Path) -> None:
    path = evidence_root / "SHA256SUMS"
    lines = [
        f"{sha256_file(candidate)}  {_relative(candidate, evidence_root)}"
        for candidate in sorted(evidence_root.rglob("*"))
        if candidate.is_file() and candidate != path
    ]
    _write_new_text(path, "\n".join(lines) + "\n")


def generate(destination: Path) -> Path:
    destination = _resolve_destination(destination)
    if _lexists(FIXED_SCRATCH_ROOT):
        raise EvidenceGenerationError(
            f"refusing existing fixed scratch root: {FIXED_SCRATCH_ROOT}"
        )

    stage = Path(
        tempfile.mkdtemp(
            prefix=".docx-evidence.",
            suffix=".tmp",
            dir=destination.parent,
        )
    )
    scratch_created = False
    published = False
    try:
        FIXED_SOURCE_DIR.mkdir(parents=True, exist_ok=False)
        scratch_created = True
        canonical = _canonical_source()
        fixed_canonical = FIXED_SOURCE_DIR / "supported-canonical.md"
        fixed_recovered = FIXED_SOURCE_DIR / "supported-recovered.md"
        _write_new_text(fixed_canonical, canonical)
        _write_new_text(stage / "supported-canonical.md", canonical)

        canonical_docx = stage / "supported-canonical.docx"
        edited_docx = stage / "supported-edited.docx"
        regenerated_docx = stage / "supported-regenerated.docx"
        _generate_docx(fixed_canonical, canonical_docx)
        _check_docx(canonical_docx, fixed_canonical)

        edit = _apply_plain_text_edit(canonical_docx, edited_docx)
        _check_docx(edited_docx, fixed_canonical)
        recovered, _ = _recover_edited(edited_docx, stage, canonical)
        _write_new_text(fixed_recovered, recovered)

        _generate_docx(fixed_recovered, regenerated_docx)
        _check_docx(regenerated_docx, fixed_recovered)
        regenerated_markdown = docx2md.convert(
            regenerated_docx, report_provenance=False
        )
        if regenerated_markdown != recovered:
            raise EvidenceGenerationError(
                "regenerated DOCX does not invert to recovered Markdown"
            )

        canonical_capture = _capture_quicklook(
            canonical_docx,
            stage,
            "supported-canonical",
            "generated-canonical-docx",
        )
        edited_capture = _capture_quicklook(
            edited_docx,
            stage,
            "supported-edited",
            "python-docx-plain-text-edit",
        )
        regenerated_capture = _capture_quicklook(
            regenerated_docx,
            stage,
            "supported-regenerated",
            "regenerated-from-recovered-markdown",
        )
        edit_comparison = _compare_canonical_and_edited(
            canonical_capture, edited_capture, stage
        )
        exact_comparison = _compare_edited_and_regenerated(
            edited_capture, regenerated_capture, stage
        )
        native = _record_native_status(edited_docx, stage)
        _validate_evidence_manifests(stage)

        initial_scan = _scan_privacy(
            stage, extra_files=(fixed_canonical, fixed_recovered)
        )
        # The manifest describes the durable published bundle. Three files are
        # intentionally written after this first content scan: this manifest,
        # the bundle README, and SHA256SUMS. Account for them explicitly so the
        # recorded file count covers the final tree rather than the pre-manifest
        # staging state. Later scans still inspect the actual final files.
        published_scan = dict(initial_scan)
        published_scan["files"] = initial_scan["files"] + 3
        generation_manifest = _generation_manifest(
            stage,
            edit,
            native,
            edit_comparison,
            exact_comparison,
            published_scan,
        )
        _write_new_json(
            stage / "generation-manifest.json",
            generation_manifest,
        )
        _write_new_text(
            stage / "README.md",
            _readme(
                stage,
                edit_comparison,
                exact_comparison,
                native,
            ),
        )
        _scan_privacy(
            stage, extra_files=(fixed_canonical, fixed_recovered)
        )
        _write_sha256s(stage)
        final_scan = _scan_privacy(
            stage, extra_files=(fixed_canonical, fixed_recovered)
        )
        if final_scan != published_scan:
            raise EvidenceGenerationError(
                "published privacy inventory differs from the manifest: "
                f"expected {published_scan}, got {final_scan}"
            )

        if _lexists(destination):
            raise EvidenceGenerationError(
                f"destination appeared during generation: {destination}"
            )
        stage.rename(destination)
        published = True
        return destination
    finally:
        if scratch_created and _lexists(FIXED_SCRATCH_ROOT):
            shutil.rmtree(FIXED_SCRATCH_ROOT)
        if not published and _lexists(stage):
            shutil.rmtree(stage)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description=(
            "Generate a new privacy-safe DOCX round-trip evidence subtree."
        )
    )
    parser.add_argument(
        "--destination",
        type=Path,
        default=DEFAULT_DESTINATION,
        help=(
            "new repository-contained evidence directory; every existing "
            "destination is refused"
        ),
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _build_parser().parse_args(argv)
    try:
        destination = generate(args.destination)
    except (
        EvidenceGenerationError,
        OSError,
        ValueError,
        zipfile.BadZipFile,
    ) as exc:
        print(
            json.dumps(
                {
                    "ok": False,
                    "code": "DOCX-EVIDENCE-GENERATION",
                    "message": str(exc),
                },
                sort_keys=True,
            ),
            file=sys.stderr,
        )
        return 1
    print(
        json.dumps(
            {
                "ok": True,
                "destination": _relative(destination, REPO_ROOT),
                "files": sum(
                    path.is_file() for path in destination.rglob("*")
                ),
            },
            sort_keys=True,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
